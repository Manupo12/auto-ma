#!/usr/bin/env python3
"""
Script de verificación mensual de backup.
Restaura el backup más reciente en un entorno temporal, verifica que
los documentos se abren correctamente, y notifica por Telegram.

Uso:
  python backend/verificar_backup.py
  # O como cron job (primer día de cada mes a las 3 AM):
  # 0 3 1 * * cd /root/fisioterapia && python backend/verificar_backup.py
"""

import os
import sys
import json
import shutil
import tempfile
import subprocess
from datetime import datetime
from pathlib import Path
from typing import List, Tuple

# Configuración
BACKUP_DIR = "/root/fisioterapia/storage/backups"
BACKUP_PATTERN = "backup_*.tar.gz"
TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN", "")
TELEGRAM_CHAT_ID = "8424650300"

# Extensiones a verificar
EXTENSIONES_VERIFICAR = [".docx", ".pdf", ".json"]


def encontrar_backup_mas_reciente() -> str:
    """Encuentra el archivo de backup más reciente."""
    if not os.path.exists(BACKUP_DIR):
        raise FileNotFoundError(f"Directorio de backup no existe: {BACKUP_DIR}")
    
    backups = sorted(
        [f for f in os.listdir(BACKUP_DIR) if f.endswith(".tar.gz")],
        reverse=True
    )
    
    if not backups:
        raise FileNotFoundError("No se encontraron archivos de backup")
    
    return os.path.join(BACKUP_DIR, backups[0])


def restaurar_backup(backup_path: str, destino: str) -> bool:
    """Restaura el backup en un directorio temporal."""
    try:
        subprocess.run(
            ["tar", "-xzf", backup_path, "-C", destino],
            check=True, capture_output=True, timeout=120
        )
        return True
    except subprocess.CalledProcessError as e:
        print(f"❌ Error al restaurar backup: {e.stderr}")
        return False


def verificar_archivos(directorio: str) -> Tuple[int, int, List[str]]:
    """
    Verifica que los archivos del backup se puedan abrir.
    
    Returns:
        (total, fallidos, lista_errores)
    """
    total = 0
    fallidos = 0
    errores = []
    
    for root, dirs, files in os.walk(directorio):
        for file in files:
            ext = os.path.splitext(file)[1].lower()
            if ext not in EXTENSIONES_VERIFICAR:
                continue
            
            filepath = os.path.join(root, file)
            total += 1
            
            try:
                if ext == ".docx":
                    verificar_docx(filepath)
                elif ext == ".pdf":
                    verificar_pdf(filepath)
                elif ext == ".json":
                    verificar_json(filepath)
            except Exception as e:
                fallidos += 1
                errores.append(f"  ❌ {file}: {str(e)[:100]}")
    
    return total, fallidos, errores


def verificar_docx(path: str):
    """Verifica que un .docx se puede abrir con python-docx."""
    from docx import Document
    doc = Document(path)
    # Verificar que tiene al menos 1 párrafo
    if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
        raise ValueError("Documento vacío")
    # Verificar que no está corrupto accediendo al texto
    _ = doc.paragraphs[0].text if doc.paragraphs else ""


def verificar_pdf(path: str):
    """Verifica que un PDF es válido (firma %PDF-)."""
    with open(path, "rb") as f:
        header = f.read(5)
        if header != b"%PDF-":
            raise ValueError("No es un PDF válido (falta firma %PDF-)")


def verificar_json(path: str):
    """Verifica que un JSON es válido."""
    with open(path, "r", encoding="utf-8") as f:
        json.load(f)


def enviar_telegram(mensaje: str):
    """Envía notificación por Telegram."""
    if not TELEGRAM_BOT_TOKEN:
        print("⚠️ TELEGRAM_BOT_TOKEN no configurado")
        return
    
    import requests
    url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage"
    data = {"chat_id": TELEGRAM_CHAT_ID, "text": mensaje, "parse_mode": "Markdown"}
    
    try:
        requests.post(url, json=data, timeout=10)
    except Exception as e:
        print(f"⚠️ Error enviando Telegram: {e}")


def main():
    TEMP_DIR = tempfile.mkdtemp(prefix="backup_verify_")
    print(f"🔍 Verificación de backup — {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    print(f"   Directorio temporal: {TEMP_DIR}")
    
    try:
        # 1. Encontrar backup
        backup_path = encontrar_backup_mas_reciente()
        backup_name = os.path.basename(backup_path)
        backup_size_mb = os.path.getsize(backup_path) / (1024 * 1024)
        print(f"📦 Backup encontrado: {backup_name} ({backup_size_mb:.1f} MB)")
        
        # 2. Restaurar
        print("📂 Restaurando...")
        if not restaurar_backup(backup_path, TEMP_DIR):
            enviar_telegram(
                "⚠️ *Verificación de Backup FALLIDA*\n\n"
                f"Backup: `{backup_name}`\n"
                "Error: No se pudo restaurar el archivo.\n"
                "Revisar manualmente."
            )
            return
        
        # 3. Verificar archivos
        print("🔎 Verificando archivos...")
        total, fallidos, errores = verificar_archivos(TEMP_DIR)
        
        # 4. Reportar
        print(f"\n📊 Resultados: {total} archivos, {fallidos} fallidos")
        
        if fallidos == 0:
            mensaje = (
                f"✅ *Backup Mensual Verificado*\n\n"
                f"📦 Backup: `{backup_name}`\n"
                f"📏 Tamaño: {backup_size_mb:.1f} MB\n"
                f"📄 Archivos verificados: *{total}*\n"
                f"❌ Fallidos: *0*\n\n"
                f"_Verificación del {datetime.now().strftime('%d/%m/%Y')}_"
            )
            print("✅ BACKUP OK")
        else:
            mensaje = (
                f"⚠️ *Backup Mensual — {fallidos} ERRORES*\n\n"
                f"📦 Backup: `{backup_name}`\n"
                f"📄 Total archivos: {total}\n"
                f"❌ Fallidos: *{fallidos}*\n\n"
                f"Errores:\n" + "\n".join(errores[:5]) + "\n\n"
                f"_Revisar manualmente._"
            )
            print(f"⚠️ {fallidos} archivos fallaron")
        
        enviar_telegram(mensaje)
        
    except FileNotFoundError as e:
        print(f"❌ {e}")
        enviar_telegram(
            "⚠️ *Verificación de Backup*\n\n"
            f"❌ No se encontraron backups.\n"
            f"Directorio: `{BACKUP_DIR}`\n\n"
            "_Configurar backup automático._"
        )
    except Exception as e:
        print(f"❌ Error inesperado: {e}")
        enviar_telegram(
            f"⚠️ *Verificación de Backup — ERROR*\n\n"
            f"❌ Error inesperado: `{str(e)[:200]}`\n\n"
            f"_Revisar logs del sistema._"
        )
    finally:
        # Limpiar directorio temporal
        try:
            shutil.rmtree(TEMP_DIR)
            print("🧹 Directorio temporal limpiado")
        except:
            pass


if __name__ == "__main__":
    main()
