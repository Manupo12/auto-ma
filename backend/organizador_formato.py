"""
Organiza un formato crudo (DOCX desorganizado/incompleto que Sandra subió)
cruzándolo con datos verificados (si existen) y devuelve estructura limpia.

Caso de uso: Sandra encuentra un formato antiguo de un paciente con info
suelta y le pide a Tomy "organízame este formato".
"""
import json
import os
import sys
import subprocess
import tempfile
from pathlib import Path
from datetime import datetime
from typing import Dict


def _log(msg: str):
    print(f"[ORGANIZA {datetime.now().strftime('%H:%M:%S')}] {msg}", file=sys.stderr, flush=True)


def _leer_docx_completo(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    lines = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text)
    for tabla in doc.tables:
        for row in tabla.rows:
            celdas = [c.text.strip()[:80] for c in row.cells if c.text.strip()]
            if celdas:
                lines.append(" | ".join(celdas))
    return "\n".join(lines)


PROMPT_ORGANIZADOR = """Eres Tomy. Sandra te pasó este formato antiguo de un paciente. Está
desorganizado y le falta información. Tu tarea:

1. EXTRAER toda la información estructurada que puedas (datos del paciente, diagnóstico, etc.).
2. CRUZAR con los datos verificados de portales que te paso (si están).
3. DETECTAR campos faltantes críticos.
4. REPORTAR contradicciones (si dice algo distinto al portal).

Devuelve SOLO un JSON entre ```json ... ``` con esta estructura:

```json
{
  "datos_organizados": {
    "paciente": {...},
    "empresa": {...},
    "siniestro": {...},
    "diagnostico": "...",
    "secciones_libres": {"metodologia": "...", "observaciones": "..."}
  },
  "campos_faltantes": ["lista de paths críticos faltantes"],
  "contradicciones": [{"campo": "...", "formato_dice": "...", "portal_dice": "..."}],
  "confianza_global": 0.0-1.0
}
```"""


def _llamar_llm_organizador(contexto: str, paciente_cc: str) -> Dict:
    """Llama lote_worker --tipo sintetizar para organizar."""
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as tmp:
        tmp.write(contexto)
        ctx_file = tmp.name

    modelo = os.getenv("LLM_MODEL_SINTESIS", "deepseek-v4-pro")
    try:
        proc = subprocess.run(
            [sys.executable, "-m", "backend.lote_worker",
             "--tipo", "sintetizar",
             "--modelo", modelo,
             "--mensaje", PROMPT_ORGANIZADOR,
             "--contexto-file", ctx_file,
             "--cc", paciente_cc],
            capture_output=True, text=True, timeout=300,
            cwd=str(Path(__file__).parent.parent),
        )
    finally:
        try: os.unlink(ctx_file)
        except: pass

    if proc.returncode != 0:
        return {"ok": False, "error": proc.stderr[:500]}

    resultado = json.loads(proc.stdout)
    respuesta = resultado.get("respuesta", "")

    import re
    match = re.search(r'```json\s*(\{.*?\})\s*```', respuesta, re.DOTALL)
    if not match:
        return {"ok": False, "error": "LLM no devolvió JSON parseable"}

    datos = json.loads(match.group(1))
    return {
        "ok": True,
        "datos_organizados": datos.get("datos_organizados", {}),
        "campos_faltantes": datos.get("campos_faltantes", []),
        "contradicciones": datos.get("contradicciones", []),
        "confianza_global": datos.get("confianza_global", 0.5),
    }


def organizar_formato(docx_path: str, paciente_cc: str) -> Dict:
    """Función principal. Devuelve dict con datos organizados + reportes."""
    path = Path(docx_path)
    if not path.exists():
        return {"ok": False, "error": f"Archivo no existe: {docx_path}"}

    contenido_crudo = _leer_docx_completo(path)
    _log(f"Crudo: {len(contenido_crudo)} chars de {path.name}")

    storage = Path(os.getenv("STORAGE_DIR", "./storage"))
    json_files = sorted((storage / "data").glob(f"*{paciente_cc}*completo.json"),
                        key=lambda p: p.stat().st_mtime, reverse=True)
    datos_verificados = {}
    if json_files:
        with open(json_files[0], encoding="utf-8") as f:
            datos_verificados = json.load(f)

    contexto = f"""═══ FORMATO CRUDO DE SANDRA ({path.name}) ═══
{contenido_crudo}

═══ DATOS VERIFICADOS DE PORTALES (CC {paciente_cc}) ═══
{json.dumps(datos_verificados, ensure_ascii=False, indent=2) if datos_verificados else '[NO DISPONIBLES]'}
"""

    return _llamar_llm_organizador(contexto, paciente_cc)
