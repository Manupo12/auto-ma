"""
Paso 3: Leer notas crudas que Sandra escribió en su workspace.

Busca archivos .txt, .md, .docx que contengan la CC del paciente en el
nombre o en su contenido (primeros 500 chars). Devuelve hasta 5 archivos
con contenido (~2000 chars c/u).
"""
import os
import sys
from pathlib import Path
from datetime import datetime
from typing import List, Dict


def _log(msg: str):
    print(f"[NOTAS {datetime.now().strftime('%H:%M:%S')}] {msg}", file=sys.stderr, flush=True)


def _leer_texto(path: Path, max_chars: int = 16000) -> str:
    """Lee primeros N chars de .txt/.md/.docx."""
    try:
        if path.suffix.lower() in (".txt", ".md"):
            return path.read_text(encoding="utf-8", errors="ignore")[:max_chars]
        if path.suffix.lower() == ".docx":
            from docx import Document
            doc = Document(str(path))
            partes = []
            seen_parrafos = set()
            for p in doc.paragraphs:
                t = p.text.strip()
                if t and t not in seen_parrafos:
                    seen_parrafos.add(t)
                    partes.append(t)
            for table in doc.tables:
                for row in table.rows:
                    seen_tcs = set()
                    row_texts = []
                    for cell in row.cells:
                        tc = cell._tc
                        if tc not in seen_tcs:
                            seen_tcs.add(tc)
                            row_texts.append(cell.text.strip())
                    
                    row_texts_merged = []
                    temp_digits = []
                    for t in row_texts:
                        if t.isdigit() and len(t) == 1:
                            temp_digits.append(t)
                        else:
                            if temp_digits:
                                digits_str = "".join(temp_digits)
                                if len(digits_str) == 8:
                                    row_texts_merged.append(f"{digits_str[:2]}/{digits_str[2:4]}/{digits_str[4:]}")
                                else:
                                    row_texts_merged.append(digits_str)
                                temp_digits = []
                            row_texts_merged.append(t)
                    if temp_digits:
                        digits_str = "".join(temp_digits)
                        if len(digits_str) == 8:
                            row_texts_merged.append(f"{digits_str[:2]}/{digits_str[2:4]}/{digits_str[4:]}")
                        else:
                            row_texts_merged.append(digits_str)
                    
                    non_empty = [t for t in row_texts_merged if t]
                    if non_empty:
                        partes.append(" | ".join(non_empty))
            return "\n".join(partes)[:max_chars]
    except Exception as e:
        _log(f"Error leyendo {path.name}: {e}")
    return ""


def _tiene_cc(path: Path, cc: str) -> bool:
    """Verifica si un archivo contiene la CC (nombre o contenido completo)."""
    cc_norm = cc.replace(".","").replace("-","").replace("'","")
    if cc_norm in path.name.replace(".","").replace("-",""):
        return True
    contenido = _leer_texto(path, max_chars=5000)
    contenido_limpio = (
        contenido
        .replace(".","").replace("-","").replace("'","")
        .replace(" ","").replace("\n","").replace("\r","").replace("\t","")
    )
    return cc_norm in contenido_limpio


def _fast_scan(directory: Path, max_depth: int = 2, timeout_s: float = 5.0) -> list:
    """Escanea directorio rapido con scandir, max 2 niveles, sin rglob."""
    import time as _time
    results = []
    start = _time.time()

    def _scan(dirpath: Path, depth: int):
        if depth > max_depth or _time.time() - start > timeout_s:
            return
        try:
            for entry in os.scandir(dirpath):
                if _time.time() - start > timeout_s:
                    return
                if entry.name.startswith('.'):
                    continue
                if entry.is_file():
                    results.append(Path(entry.path))
                elif entry.is_dir() and depth < max_depth:
                    _scan(Path(entry.path), depth + 1)
        except (PermissionError, OSError):
            pass

    _scan(directory, 0)
    return results


def _mtime_safe(ruta: str) -> float:
    """Retorna st_mtime de forma segura, 0 si falla."""
    try:
        return Path(ruta).stat().st_mtime
    except OSError:
        return 0.0


def leer_notas_crudas(cc: str, max_archivos: int = 5) -> tuple:
    """Busca archivos del workspace con el CC. Retorna (lista, completo)."""
    workspace = Path(os.getenv("WORKSPACE_DIR", str(Path.home() / "rilo-workspace")))
    if not workspace.exists():
        _log(f"Workspace no existe: {workspace}")
        return [], False

    candidatos = []
    extensiones = (".txt", ".md", ".docx")
    completo = True
    try:
        archivos = _fast_scan(workspace, max_depth=2, timeout_s=8.0)
        completo = len(archivos) < 500  # si se corto por timeout, incomplete
        for path in archivos:
            if path.suffix.lower() not in extensiones:
                continue
            try:
                if path.stat().st_size > 5_000_000:
                    continue
            except OSError:
                continue

            if _tiene_cc(path, cc):
                contenido_completo = _leer_texto(path, max_chars=16000)
                if contenido_completo:
                    nombre = path.name
                    # Detectar tipo de documento por prefijo del nombre
                    nombre_upper = nombre.upper()
                    if nombre_upper.startswith("VOI_") or nombre_upper.startswith("VALORACION_"):
                        tipo_doc = "voi_referencia"  # VOI completado por Sandra = GROUND TRUTH
                    elif nombre_upper.startswith("CC_"):
                        tipo_doc = "cc_paciente"     # Cierre de Caso
                    elif nombre_upper.startswith("CR_"):
                        tipo_doc = "cr_paciente"     # Carta de Recomendaciones
                    elif nombre_upper.startswith("PT_") or nombre_upper.startswith("PRUEBA"):
                        tipo_doc = "prueba_trabajo"
                    else:
                        tipo_doc = "nota_workspace"
                    candidatos.append({
                        "nombre": nombre,
                        "ruta": str(path),
                        "contenido": contenido_completo,
                        "tipo": tipo_doc,
                    })
    except Exception as e:
        _log(f"Error escaneando: {e}")

    # Ordenar: voi_referencia primero, luego por mtime (más reciente primero)
    tipo_prioridad = {"voi_referencia": 0, "cc_paciente": 1, "cr_paciente": 2, "prueba_trabajo": 3, "nota_workspace": 4}
    candidatos.sort(key=lambda d: (tipo_prioridad.get(d.get("tipo",""), 9), -_mtime_safe(d["ruta"])))
    return candidatos[:max_archivos], completo


def ejecutar(paciente_cc: str) -> dict:
    notas, completo = leer_notas_crudas(paciente_cc)
    return {"notas_crudas": notas, "total": len(notas), "escaneo_completo": completo}
