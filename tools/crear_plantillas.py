"""
Convierte documentos de ejemplo en plantillas con marcadores {{...}}.
Reemplaza datos personales reales con placeholders genéricos.
"""
from docx import Document
import shutil, os, re

BASE = "/root/fisioterapia/templates/formatos"
OUT = BASE  # Sobrescribir los mismos archivos como plantillas

# Patrones de reemplazo (regex → marcador)
REEMPLAZOS = [
    # Fechas: 12/12/2024, 12-12-2024, etc.
    (r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', '{{FECHA_COMPLETA}}'),
    (r'\b\d{1,2} de (enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre)( de)? \d{4}\b', '{{FECHA_LARGA}}'),
    # Cédulas/NIT (números de 6-12 dígitos con posibles separadores)
    (r'\b\d{1,3}\.\d{3}\.\d{3}[-.]?\d?\b', '{{NUMERO_DOCUMENTO}}'),
    (r'\b\d{7,12}\b', '{{NUMERO_DOCUMENTO}}'),
    # Teléfonos
    (r'\b\d{7,10}\b', '{{TELEFONO}}'),
    # Correos
    (r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '{{CORREO}}'),
    # URLs
    (r'https?://[^\s]+', '{{URL}}'),
]

# Nombres de personas conocidas (de los ejemplos) → reemplazar
NOMBRES_CONOCIDOS = [
    "Diana Cristina Sánchez Pama",
    "Juan Carlos Duran Narvaez",
    "HEBERTH ARMANDO RUIZ PAVA",
    "Ingrid Johan Vargas Reyes",
    "Sandra Patricia Polania Osorio",
    "SANDRA PATRICIA POLANIA OSORIO",
    "Sandra Patricia Polanía Osorio",
]

EMPRESAS_CONOCIDAS = [
    "FISCALÍA GENERAL DE LA NACIÓN",
    "BOLIVARIANA DE MINERALES Y CIA LTDA",
    "Seccional Rama Judicial Neiva",
]


def reemplazar_en_runs(parrafo, viejo, nuevo):
    """Reemplaza texto en todos los runs de un párrafo."""
    for run in parrafo.runs:
        if viejo in run.text:
            run.text = run.text.replace(viejo, nuevo)
            return True
    return False


def procesar_documento(archivo):
    """Convierte un documento ejemplo en plantilla."""
    path = os.path.join(BASE, archivo)
    if not archivo.endswith('.docx'):
        return
    
    # Crear backup
    backup = path + ".bak"
    if not os.path.exists(backup):
        shutil.copy(path, backup)
        print(f"  Backup: {backup}")
    
    doc = Document(path)
    cambios = 0
    
    # Procesar párrafos
    for parrafo in doc.paragraphs:
        texto = parrafo.text
        
        # Reemplazar nombres conocidos
        for nombre in NOMBRES_CONOCIDOS:
            if nombre in texto:
                if "SANDRA" in nombre or "Sandra" in nombre:
                    if reemplazar_en_runs(parrafo, nombre, "{{NOMBRE_PROFESIONAL}}"):
                        cambios += 1
                else:
                    if reemplazar_en_runs(parrafo, nombre, "{{NOMBRE_AFILIADO}}"):
                        cambios += 1
        
        # Reemplazar empresas conocidas
        for empresa in EMPRESAS_CONOCIDAS:
            if empresa in texto:
                if reemplazar_en_runs(parrafo, empresa, "{{EMPRESA}}"):
                    cambios += 1
    
    # Procesar tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    texto = parrafo.text
                    
                    # Nombres
                    for nombre in NOMBRES_CONOCIDOS:
                        if nombre in texto:
                            if "SANDRA" in nombre or "Sandra" in nombre:
                                if reemplazar_en_runs(parrafo, nombre, "{{NOMBRE_PROFESIONAL}}"):
                                    cambios += 1
                            else:
                                if reemplazar_en_runs(parrafo, nombre, "{{NOMBRE_AFILIADO}}"):
                                    cambios += 1
                    
                    # Empresas
                    for empresa in EMPRESAS_CONOCIDAS:
                        if empresa in texto:
                            if reemplazar_en_runs(parrafo, empresa, "{{EMPRESA}}"):
                                cambios += 1
                    
                    # Patrones regex
                    for patron, marcador in REEMPLAZOS:
                        match = re.search(patron, texto)
                        if match:
                            valor = match.group(0)
                            if reemplazar_en_runs(parrafo, valor, marcador):
                                cambios += 1
    
    # Guardar
    doc.save(path)
    return cambios


# Procesar todos los .docx
print("Creando plantillas desde ejemplos...\n")
for archivo in sorted(os.listdir(BASE)):
    if archivo.endswith('.docx') and not archivo.endswith('.bak'):
        print(f"📄 {archivo}")
        cambios = procesar_documento(archivo)
        print(f"  → {cambios} reemplazos\n")

print("✅ Plantillas generadas.")
print("⚠️ Revisa los archivos — algunos campos pueden necesitar ajuste manual.")
print("   Los backups están como .docx.bak por si necesitas restaurar.")
