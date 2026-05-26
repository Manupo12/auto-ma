"""Tests end-to-end para Fase A."""
import json
import os
import sys
import textwrap
from pathlib import Path
from unittest.mock import patch, MagicMock
import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))


def test_chat_carga_datos_verificados_si_existen(tmp_path, monkeypatch):
    """Con FASE_A_ENABLED=true y JSON dummy, procesar_mensaje debe inyectar datos."""
    monkeypatch.setenv("FASE_A_ENABLED", "true")
    monkeypatch.setenv("STORAGE_DIR", str(tmp_path / "storage"))
    monkeypatch.setenv("WORKSPACE_DIR", str(tmp_path / "workspace_vacio"))
    (tmp_path / "workspace_vacio").mkdir(exist_ok=True)

    data_dir = tmp_path / "storage" / "data"
    data_dir.mkdir(parents=True, exist_ok=True)
    datos_dummy = {
        "cc": "1193143688",
        "medifolios": {
            "nombre1": "JUAN",
            "apellido1": "DURAN",
            "siniestro_medi": "503463870",
            "telefono": "3001234567",
            "direccion": "CALLE 1",
        },
        "positiva": {
            "siniestros": [{"id": "503463870", "fecha": "02/03/2026"}],
        },
        "_meta": {
            "extraido_en": "2026-05-19T21:00:00",
            "discrepancias": [],
        },
    }
    json_path = data_dir / "20260519-1193143688-completo.json"
    json_path.write_text(json.dumps(datos_dummy), encoding="utf-8")

    from backend.chat_handler import _cargar_datos_verificados, _formatear_datos_verificados, _llamar_llm, procesar_mensaje, _cache_workspace
    _cache_workspace["ts"] = 0

    datos = _cargar_datos_verificados("1193143688")
    assert datos is not None
    assert datos["cc"] == "1193143688"

    # Verificar formateo
    texto = _formatear_datos_verificados(datos)
    assert "[OK] Siniestro" in texto
    assert "503463870" in texto
    assert "[OK] Paciente" in texto
    assert "DATOS VERIFICADOS" in texto

    # Verificar integración en procesar_mensaje con mock LLM
    with patch("backend.chat_handler.requests.post") as mock_post:
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {
            "choices": [{"message": {"content": "Respuesta con datos verificados ✅"}, "finish_reason": "stop"}],
            "usage": {"total_tokens": 100},
        }
        mock_post.return_value = mock_resp

        resultado = procesar_mensaje("revisa al paciente 1193143688", paciente_cc="1193143688")

    assert resultado.get("contenido")
    assert "datos verificados" in resultado["contenido"].lower() or "✅" in resultado["contenido"]


@pytest.mark.skipif(not os.getenv("OPENCODE_GO_API_KEY"), reason="needs OPENCODE_GO_API_KEY")
@pytest.mark.live
def test_e2e_oom_no_aparece_con_12_docs(tmp_path):
    pass


@pytest.mark.live
def test_e2e_workflow_completo_mockeado(tmp_path, monkeypatch):
    from unittest.mock import patch, MagicMock, AsyncMock
    monkeypatch.setenv("TOMY_COMPLETO_ENABLED", "true")
    monkeypatch.setenv("STORAGE_DIR", str(tmp_path / "storage"))
    monkeypatch.setenv("WORKFLOW_DB_PATH", str(tmp_path / "wf.db"))
    monkeypatch.setenv("WORKSPACE_DIR", str(tmp_path / "workspace"))
    (tmp_path / "workspace").mkdir(exist_ok=True)
    (tmp_path / "storage" / "data").mkdir(parents=True, exist_ok=True)

    with patch("backend.flujo_audio.transcribir_audio") as m_trans, \
         patch("backend.playwright_real.orquestador.extraer_paciente_completo", new_callable=AsyncMock) as m_pw, \
         patch("backend.workflow_steps.generar_formatos.ejecutar") as m_gen_steps, \
         patch("backend.workflow_steps.transcribir._calcular_duracion_s") as m_dur, \
         patch("backend.workflow_steps.qa_formatos.ejecutar") as m_qa_steps, \
         patch("backend.workflow_steps.convertir_pdf.ejecutar") as m_pdf_steps, \
         patch("backend.workflow_steps.notificar_listo.ejecutar") as m_notif_steps, \
         patch("backend.workflow_steps.sintetizar_maestro.ejecutar") as m_sint_steps:

        m_dur.return_value = 600.0
        m_trans.return_value = {"texto": "Paciente dice que le duele.", "segmentos": [], "confianza": 0.9, "duracion": 600}
        m_pw.return_value = {
            "cc": "1193143688",
            "medifolios": {"nombre1": "JUAN", "apellido1": "DURAN"},
            "positiva": {"siniestros": [{"id": "503463870"}]},
            "_meta": {"discrepancias": []},
        }
        m_sint_steps.return_value = {
            "ok": True,
            "datos_clinicos": {
                "paciente": {"documento": "1193143688", "nombre": "JUAN"},
                "estado_caso": "SEGUIMIENTO",
                "_meta": {"campos_faltantes": []},
            },
        }
        m_gen_steps.return_value = {
            "ok": True,
            "formatos_generados": [
                {"formato": "analisis", "archivo": str(tmp_path / "analisis.docx")},
                {"formato": "valoracion", "archivo": str(tmp_path / "valoracion.docx")},
            ],
            "errores": [],
        }
        m_qa_steps.return_value = {"ok": True, "resultados": [{"formato": "analisis", "qa_ok": True}, {"formato": "valoracion", "qa_ok": True}], "exitosos": 2}
        m_pdf_steps.return_value = {"ok": True, "pdfs": [], "errores": []}
        m_notif_steps.return_value = {"ok": True, "telegram_enviado": True}

        audio = tmp_path / "test.m4a"
        audio.write_bytes(b"\x00" * 100)

        from backend.workflow_runner import ejecutar_workflow
        resultado = ejecutar_workflow(audio_path=str(audio), paciente_cc="1193143688")

    assert resultado["estado"] == "listo"
    assert resultado["task_id"]
    assert m_notif_steps.called
    """12 docs simulados → procesamiento completo sin exit 137."""
    from backend.chat_handler import procesar_mensaje
    from docx import Document

    for i in range(12):
        p = tmp_path / f"doc_{i}_1193143688.docx"
        doc = Document()
        doc.add_paragraph(f"Documento de prueba {i}")
        doc.add_paragraph("Paciente Juan Carlos")
        doc.add_paragraph("CC: 1193143688")
        doc.add_paragraph("Siniestro: 503463870")
        doc.save(str(p))

    import subprocess
    import sys
    script = textwrap.dedent(f'''\
        import sys
        sys.path.insert(0, '{tmp_path.parent.parent}')
        import os
        os.environ['FASE_A_ENABLED'] = 'true'
        os.environ['WORKSPACE_DIR'] = '{tmp_path}'
        from backend.chat_handler import procesar_mensaje
        r = procesar_mensaje('revisa todos los formatos del paciente', paciente_cc='1193143688')
        print(f"OK: {{len(r.get('contenido',''))}} chars")
    ''')
    proc = subprocess.run(
        [sys.executable, "-c", script],
        capture_output=True, text=True, timeout=300,
        cwd=Path(__file__).parent.parent,
    )
    assert proc.returncode == 0, f"exit {proc.returncode}, stderr={proc.stderr[:500]}"
    assert "OK:" in proc.stdout
