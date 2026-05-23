"""
lote_worker — subprocess CLI que ejecuta 1 llamada LLM y muere.

Diseño: el padre (FastAPI) acumula reasoning_content de DeepSeek v4 y crashea
por OOM. Aislar cada llamada en un subprocess hace que el OS libere la RAM
al terminar el proceso, eliminando el leak.

Uso:
  python -m backend.lote_worker \
      --tipo extraer \
      --archivos a.docx,b.docx \
      --modelo deepseek-chat \
      --cc 1193143688

Salida: JSON por stdout. Exit 0 si OK, exit 1 si error (stacktrace en stderr).
"""
import argparse
import json
import os
import sys
import time
import traceback
from pathlib import Path
from datetime import datetime

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

import requests


API_KEY = os.getenv("OPENCODE_GO_API_KEY", "")
BASE_URL = os.getenv("OPENCODE_GO_BASE_URL", "https://opencode.ai/zen/go/v1")
MODEL_EXTRACCION = os.getenv("LLM_MODEL_EXTRACCION", "deepseek-v4-flash")
MODEL_SINTESIS = os.getenv("LLM_MODEL_SINTESIS", "deepseek-v4-pro")

PROMPT_EXTRACCION = """Extrae los siguientes campos del documento. Formato EXACTO (una línea por campo):

PACIENTE: <nombre completo>
CC: <número>
SINIESTRO: <id>
FECHA_EVENTO: <dd/mm/aaaa>
DIAGNOSTICO: <código CIE-10 + descripción>
EMPRESA: <nombre>
SEGMENTO: <parte del cuerpo>
CAMPOS_VACIOS: <lista separada por comas>
OBSERVACIONES: <notas relevantes>

Si un campo no aparece, escribe FALTA."""


def _log(msg: str):
    print(f"[LOTE_WORKER {datetime.now().strftime('%H:%M:%S')}] {msg}", file=sys.stderr, flush=True)


def _leer_docx(ruta: Path) -> str:
    """Lee .docx (primeros 60 párrafos + 5 tablas) o .txt."""
    if not ruta.exists():
        return ""
    if ruta.suffix.lower() == ".txt":
        return ruta.read_text(encoding="utf-8", errors="ignore")[:6000]
    if ruta.suffix.lower() != ".docx":
        return ""
    from docx import Document
    doc = Document(str(ruta))
    lines = [f"📄 {ruta.name}\n"]
    for p in doc.paragraphs[:60]:
        t = p.text.strip()
        if t:
            lines.append(t[:200])
    for ti, tabla in enumerate(doc.tables[:5]):
        lines.append(f"\nTabla {ti+1}:")
        for row in tabla.rows[:5]:
            celdas = [c.text.strip()[:50] for c in row.cells[:6] if c.text.strip()]
            if celdas:
                lines.append(" | ".join(celdas))
    return "\n".join(lines)[:6000]


def _llamar_llm(messages: list, modelo: str, max_tokens: int) -> dict:
    """Llamada HTTP al LLM. Devuelve dict con content, reasoning, tokens, finish."""
    t0 = time.time()
    last_error = None
    for intento in range(3):
        try:
            resp = requests.post(
                f"{BASE_URL}/chat/completions",
                headers={
                    "Authorization": f"Bearer {API_KEY}",
                    "Content-Type": "application/json",
                    "User-Agent": "Mozilla/5.0",
                },
                json={
                    "model": modelo,
                    "messages": messages,
                    "temperature": 0.3 if "flash" in modelo else 0.7,
                    "max_tokens": max_tokens,
                },
                timeout=(30, 300),  # (connect, read) timeouts
            )
            resp.raise_for_status()
            break
        except Exception as e:
            last_error = e
            if intento < 2:
                wait = (intento + 1) * 10
                print(f"[LOTE] Reintento {intento+2}/3 en {wait}s: {e}", file=sys.stderr)
                time.sleep(wait)
    else:
        print(json.dumps({"ok": False, "error": f"HTTP error after 3 retries: {last_error}"}))
        sys.exit(1)

    elapsed = time.time() - t0
    body = resp.json()
    choice = body.get("choices", [{}])[0]
    msg = choice.get("message", {})
    usage = body.get("usage", {})
    try:
        from backend.costos_tracker import registrar_uso
        registrar_uso(modelo, {
            "input": usage.get("prompt_tokens", 0),
            "output": usage.get("completion_tokens", 0),
            "reasoning": usage.get("reasoning_tokens", 0),
        })
    except Exception:
        pass
    return {
        "content": (msg.get("content") or "").strip(),
        "reasoning": (msg.get("reasoning_content") or "").strip(),
        "finish": choice.get("finish_reason"),
        "tokens": {
            "input": usage.get("prompt_tokens", 0),
            "output": usage.get("completion_tokens", 0),
            "reasoning": usage.get("reasoning_tokens", 0),
            "total": usage.get("total_tokens", 0),
        },
        "tiempo_s": round(elapsed, 2),
    }


def ejecutar_extraccion(archivos: list[Path], cc: str, modelo: str) -> dict:
    """Lee archivos, manda prompt mínimo de extracción, devuelve datos."""
    contenidos = []
    for ruta in archivos:
        c = _leer_docx(ruta)
        if c:
            contenidos.append(c)
    if not contenidos:
        return {"ok": False, "error": "ningún archivo legible"}

    contexto = "\n\n---\n\n".join(contenidos)[:8000]
    messages = [
        {"role": "system", "content": PROMPT_EXTRACCION},
        {"role": "user", "content": f"CC: {cc}\n\n{contexto}"},
    ]
    resultado = _llamar_llm(messages, modelo, max_tokens=4000)
    return {
        "ok": True,
        "tipo": "extraer",
        "modelo": modelo,
        "tokens": resultado["tokens"],
        "tiempo_s": resultado["tiempo_s"],
        "datos_raw": resultado["content"],
        "finish": resultado["finish"],
    }


def ejecutar_sintesis(mensaje: str, contexto: str, modelo: str) -> dict:
    """Llamada de síntesis: razonamiento completo, prompt completo."""
    from backend.chat_handler import SYSTEM_PROMPT
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": f"{mensaje}\n\n[CONTEXTO]:\n{contexto}"},
    ]
    resultado = _llamar_llm(messages, modelo, max_tokens=16000)
    if not resultado["content"] and resultado["reasoning"]:
        fallback = "🤔 " + " ".join(resultado["reasoning"].split("\n")[-3:])[:300]
        return {"ok": True, "tipo": "sintetizar", "modelo": modelo,
                "tokens": resultado["tokens"], "tiempo_s": resultado["tiempo_s"],
                "respuesta": fallback, "finish": resultado["finish"], "fallback": True}
    return {
        "ok": True,
        "tipo": "sintetizar",
        "modelo": modelo,
        "tokens": resultado["tokens"],
        "tiempo_s": resultado["tiempo_s"],
        "respuesta": resultado["content"],
        "finish": resultado["finish"],
    }


def main():
    parser = argparse.ArgumentParser(description="lote_worker — subprocess LLM aislado")
    parser.add_argument("--tipo", choices=["extraer", "sintetizar"], required=True)
    parser.add_argument("--archivos", default="", help="rutas separadas por coma")
    parser.add_argument("--modelo", default="")
    parser.add_argument("--cc", default="")
    parser.add_argument("--mensaje", default="")
    parser.add_argument("--contexto-file", default="", help="ruta a archivo con contexto largo")
    args = parser.parse_args()

    if not API_KEY:
        print("ERROR: OPENCODE_GO_API_KEY no configurada", file=sys.stderr)
        sys.exit(1)

    try:
        if args.tipo == "extraer":
            archivos = [Path(p.strip()) for p in args.archivos.split(",") if p.strip()]
            if not archivos:
                print("ERROR: --archivos vacío", file=sys.stderr)
                sys.exit(1)
            modelo = args.modelo or MODEL_EXTRACCION
            resultado = ejecutar_extraccion(archivos, args.cc, modelo)
        else:  # sintetizar
            contexto = ""
            if args.contexto_file:
                contexto = Path(args.contexto_file).read_text(encoding="utf-8")
            modelo = args.modelo or MODEL_SINTESIS
            resultado = ejecutar_sintesis(args.mensaje, contexto, modelo)

        print(json.dumps(resultado, ensure_ascii=False))
        sys.exit(0 if resultado.get("ok") else 1)
    except Exception as e:
        traceback.print_exc(file=sys.stderr)
        print(json.dumps({"ok": False, "error": str(e), "tipo_error": type(e).__name__}), file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
