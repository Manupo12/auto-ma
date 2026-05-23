"""
Chat Tomy Dashboard v5 -- Robusto, rápido, tolerante a fallos.

Arreglos de raíz (Mayo 2026):
- Sin rglob (causaba hangs de 30s+ en /mnt/c/)
- Manejo de respuesta LLM vacía (0 chars -> mensaje de fallback)
- Búsqueda de archivos por nombre parcial (VOI_JOHN -> encuentra VOI_JOHN DEIVER...)
- Workspace alternativo: si el principal no existe, usa storage/docs/
- Timeouts agresivos en todas las operaciones de archivo
- Endpoint /api/workspace para diagnóstico
"""
import os, json, re, requests, time, sys, traceback, subprocess
from pathlib import Path
from datetime import datetime
from typing import Optional

try:
    from dotenv import load_dotenv; load_dotenv()
except ImportError: pass

API_KEY = os.getenv("OPENCODE_GO_API_KEY", "")
BASE_URL = os.getenv("OPENCODE_GO_BASE_URL", "https://opencode.ai/zen/go/v1")
MODEL = os.getenv("LLM_MODEL", "deepseek-v4-pro")

# Workspace principal (configurado en .env)
WORKSPACE_DIR = Path(os.getenv("WORKSPACE_DIR", Path.home() / "rilo-workspace"))
# Fallback: storage/docs del proyecto (siempre existe en Docker/WSL)
FALLBACK_WORKSPACE = Path(__file__).parent.parent / "storage" / "docs"

# Caché (5 min)
_cache_workspace = {"files": "", "ts": 0}

def _log(msg):
    print(f"[TOMY {datetime.now().strftime('%H:%M:%S')}] {msg}", file=sys.stderr, flush=True)

SYSTEM_PROMPT_SIMPLE = """Eres Tomy, asistente de Sandra (RILO SAS, fisioterapeuta Colombia).
Espanol colombiano, calido y directo. Responde de forma breve y concreta.
NUNCA inventes datos clinicos. Si no sabes algo, dilo sin rodeos.
Los 7 formatos: Analisis de Exigencias, Carta Medidas Preventivas, Carta Recomendaciones,
Cierre de Caso, Citacion a Empresas, Prueba de Trabajo, VOI."""

SYSTEM_PROMPT = """Eres Tomy, asistente EXPERTO de Sandra (RILO SAS, fisioterapeuta ARL Positiva Colombia).
Espanol colombiano, calido pero profesional. NUNCA inventes datos clinicos.

Los 7 formatos: Analisis de Exigencias, Carta Medidas Preventivas, Carta Recomendaciones,
Cierre de Caso, Citacion a Empresas, Prueba de Trabajo, VOI (Valoracion Desempeno Ocupacional).

Al analizar documentos clinicos:
1. Extrae TODOS los datos y cruzalos entre formatos
2. Marca [OK] VERIFICADO / [!] PENDIENTE PORTAL / [X] DISCREPANCIA
3. Medifolios: verificar datos personales, siniestro (en OBSERVACIONES de Agenda Citas)
4. ARL Positiva: siniestro oficial, CIE-10, autorizaciones
5. Correcciones: ORIGINAL -> CORREGIDO + motivo en una linea

Cuando hay multiples formatos, incluye tabla de consistencia con campos clave.
Campos faltantes: [!!] FALTA: [campo] -- [donde obtenerlo]
CUANDO PIDAN FORMATOS: solo links de descarga. Nunca el contenido como texto."""

#  WORKSPACE

def _get_active_workspace() -> Path:
    """Devuelve el workspace activo (principal o fallback)."""
    if WORKSPACE_DIR.exists():
        return WORKSPACE_DIR
    if FALLBACK_WORKSPACE.exists():
        _log(f"WORKSPACE: usando fallback {FALLBACK_WORKSPACE}")
        return FALLBACK_WORKSPACE
    return WORKSPACE_DIR  # retornar el principal aunque no exista


def _fast_scandir(directory: Path, max_depth: int = 2, timeout: float = 3.0) -> list:
    """
    Escanea directorio rápido. MÁX 2 niveles, SIN rglob, con timeout.
    Retorna lista de (path, is_dir, size_kb, mtime_str).
    """
    results = []
    start = time.time()
    
    def _scan(dirpath: Path, depth: int):
        if depth > max_depth or time.time() - start > timeout:
            return
        try:
            for entry in os.scandir(dirpath):
                if time.time() - start > timeout:
                    return
                if entry.name.startswith('.'):
                    continue
                try:
                    st = entry.stat()
                    size_kb = st.st_size // 1024 if entry.is_file() else 0
                    mtime = datetime.fromtimestamp(st.st_mtime).strftime('%d/%m/%Y')
                except OSError:
                    size_kb, mtime = 0, "?"
                
                if entry.is_file():
                    ext = entry.name.lower()
                    if ext.endswith(('.docx', '.txt', '.json', '.pdf')):
                        results.append((Path(entry.path), False, size_kb, mtime))
                elif entry.is_dir() and depth < max_depth:
                    results.append((Path(entry.path), True, 0, mtime))
                    _scan(entry.path, depth + 1)
        except PermissionError:
            pass
        except Exception as e:
            _log(f"SCAN: error en {dirpath}: {e}")
    
    _scan(directory, 0)
    return results


def _listar_workspace() -> str:
    """Lista archivos del workspace para inyectar al LLM. Incluye storage/docs/."""
    global _cache_workspace
    now = time.time()
    if now - _cache_workspace["ts"] < 60:
        return _cache_workspace["files"]

    t0 = time.time()
    all_entries = []

    if FALLBACK_WORKSPACE.exists():
        fb_entries = _fast_scandir(FALLBACK_WORKSPACE, max_depth=1, timeout=3.0)
        all_entries.extend(fb_entries)

    ws = WORKSPACE_DIR if WORKSPACE_DIR.exists() else None
    if ws and ws != FALLBACK_WORKSPACE:
        ws_entries = _fast_scandir(ws, max_depth=2, timeout=5.0)
        all_entries.extend(ws_entries)

    if not all_entries:
        _cache_workspace = {"files": "", "ts": now}
        return ""

    lines = []
    for path, is_dir, size_kb, mtime in all_entries[:50]:
        base = FALLBACK_WORKSPACE if str(path).startswith(str(FALLBACK_WORKSPACE)) else (ws or FALLBACK_WORKSPACE)
        try: rel = str(path.relative_to(base))
        except ValueError: rel = path.name
        if is_dir: lines.append(f"  {rel}/")
        else: lines.append(f"  {rel} ({size_kb}KB, {mtime})")

    result = "\n".join(lines[:30]) if lines else "  (carpeta vacia)"
    _log(f"WORKSPACE: {len(all_entries)} entradas en {time.time()-t0:.1f}s")
    _cache_workspace = {"files": result, "ts": now}
    return result


def _buscar_archivo(mensaje: str) -> Optional[Path]:
    """Busca archivo mencionado en el mensaje. Soporta nombres parciales."""
    _log(f"BUSCAR: '{mensaje[:100]}...'")
    
    ws = _get_active_workspace()
    if not ws.exists():
        _log(f"BUSCAR: workspace no existe: {ws}")
        return None
    
    t0 = time.time()
    entries = _fast_scandir(ws, max_depth=3, timeout=4.0)
    archivos = [(p, s, m) for p, is_dir, s, m in entries if not is_dir]
    
    _log(f"BUSCAR: {len(archivos)} archivos en {time.time()-t0:.1f}s")
    if not archivos:
        return None
    
    ml = mensaje.lower()
    
    # Estrategia 0: NUEVA -- nombre PARCIAL del archivo en el mensaje
    # "VOI_JOHN DEIVER" -> buscar archivos que contengan "voi" y "john"
    palabras_msg = set(re.findall(r'[a-záéíóúñ0-9]{3,}', ml))
    for path, size_kb, mtime in archivos:
        stem_lower = path.stem.lower()
        # Contar cuántas palabras del mensaje aparecen en el nombre del archivo
        coincidencias = sum(1 for w in palabras_msg if w in stem_lower)
        if coincidencias >= 2:
            _log(f"BUSCAR: parcial ({coincidencias} palabras) -> {path.name}")
            return path
    
    # Estrategia 1: nombre exacto del archivo (sin extensión) en el mensaje
    for path, size_kb, mtime in archivos:
        if path.stem.lower() in ml:
            _log(f"BUSCAR: exacto -> {path.name}")
            return path
    
    # Estrategia 2: coincidencia de palabras "reales" (≥4 letras)
    palabras_largas = set(re.findall(r'[a-záéíóúñ]{4,}', ml))
    for path, size_kb, mtime in archivos:
        palabras_archivo = set(re.findall(r'[a-záéíóúñ]{4,}', path.stem.lower()))
        if len(palabras_largas & palabras_archivo) >= 2:
            _log(f"BUSCAR: palabras largas -> {path.name}")
            return path
    
    # Estrategia 3: tipo de formato + más reciente
    tipos = {
        "analisis": "analisis", "análisis": "analisis",
        "voi": "valoracion", "valoracion": "valoracion", "desempeño": "valoracion",
        "medidas": "medidas", "recomendaciones": "recomendaciones",
        "cierre": "cierre", "citacion": "citacion", "empresas": "citacion",
        "prueba": "prueba", "trabajo": "prueba",
    }
    for kw, tipo in tipos.items():
        if kw in ml:
            cand = [(p, s, m) for p, s, m in archivos if tipo in p.stem.lower()]
            if cand:
                # Priorizar por coincidencia de nombre
                for p, s, m in cand:
                    if palabras_largas & set(re.findall(r'[a-záéíóúñ]{4,}', p.stem.lower())):
                        _log(f"BUSCAR: {tipo} con nombre -> {p.name}")
                        return p
                # Fallback: más reciente
                best = max(cand, key=lambda x: x[1])  # por mtime string (dd/mm/yyyy)
                _log(f"BUSCAR: {tipo} más reciente -> {best[0].name}")
                return best[0]
    
    _log("BUSCAR: no encontrado")
    return None


def _buscar_archivos_paciente(cc: str = "", nombre: str = "", mensaje: str = "") -> list:
    """
    Busca TODOS los archivos de un paciente por CC, nombre.
    Busca en workspace principal Y en storage/docs/.
    """
    ws = _get_active_workspace()
    dirs = [FALLBACK_WORKSPACE] if FALLBACK_WORKSPACE.exists() else []
    if ws.exists() and ws not in dirs:
        dirs.append(ws)

    all_archivos = []
    for d in dirs:
        entries = _fast_scandir(d, max_depth=2, timeout=4.0)
        all_archivos.extend([(p, s, m) for p, is_dir, s, m in entries if not is_dir and p.suffix.lower() == '.docx'])

    if not all_archivos:
        return []
    
    entries = _fast_scandir(ws, max_depth=3, timeout=4.0)
    all_archivos = sorted(set(all_archivos), key=lambda x: x[0].stem)
    archivos = all_archivos
    if not archivos:
        return []
    
    ml = mensaje.lower()
    
    # Detectar si el usuario quiere TODOS los formatos
    quiere_todos = any(kw in ml for kw in [
        "todos los formatos", "todos los documentos", "los 7 formatos",
        "todos los archivos", "completos", "todos los del paciente",
        "revisa todos", "revisame todos", "organiza todos", "verifica todos"
    ])
    
    # Buscar por CC
    if cc:
        matches = [(p, s, m) for p, s, m in archivos if cc in p.stem]
        if matches:
            _log(f"MULTI: {len(matches)} archivos para CC {cc}")
            return sorted(matches, key=lambda x: x[0].stem)
    
    # Buscar por nombre de paciente
    if nombre:
        nombre_lower = nombre.lower()
        matches = [(p, s, m) for p, s, m in archivos if nombre_lower in p.stem.lower()]
        if matches:
            _log(f"MULTI: {len(matches)} archivos para '{nombre}'")
            return sorted(matches, key=lambda x: x[0].stem)
    
    # Extraer CC del mensaje (mantener espacios para word boundaries)
    m_cc = re.search(r'\b(\d{6,12})\b', mensaje.replace("'", "").replace(".", "").replace(",", ""))
    if m_cc:
        cc_found = m_cc.group(1)
        matches = [(p, s, m) for p, s, m in archivos if cc_found in p.stem]
        if matches and (quiere_todos or len(matches) >= 3):
            _log(f"MULTI: {len(matches)} archivos para CC {cc_found}")
            return sorted(matches, key=lambda x: x[0].stem)
    
    # Si quiere todos pero sin CC, buscar por palabras del nombre
    if quiere_todos:
        palabras = set(re.findall(r'[a-záéíóúñ]{4,}', ml))
        # Agrupar archivos que comparten CC o nombre
        for p, s, m in archivos:
            if palabras & set(re.findall(r'[a-záéíóúñ]{4,}', p.stem.lower())):
                # Encontrar el grupo: todos los archivos con mismo prefijo (CC)
                cc_match = re.search(r'(\d{6,12})', p.stem)
                if cc_match:
                    cc_group = cc_match.group(1)
                    matches = [(p2, s2, m2) for p2, s2, m2 in archivos if cc_group in p2.stem]
                    if len(matches) >= 3:
                        _log(f"MULTI: {len(matches)} archivos agrupados por CC {cc_group}")
                        return sorted(matches, key=lambda x: x[0].stem)
                break
    
    return []


def _leer_multiples_docx(rutas: list, max_chars_total: int = 15000) -> str:
    """
    Lee múltiples archivos .docx con límite de chars total.
    Cada archivo: versión resumida (~1500 chars c/u para que quepan 7+).
    """
    if not rutas:
        return ""
    
    _log(f"MULTI-LEER: {len(rutas)} archivos, límite {max_chars_total} chars")
    t0 = time.time()
    
    chunks = []
    chars_used = 0
    chars_per_file = max(500, max_chars_total // max(len(rutas), 1))
    
    for path, size_kb, mtime in rutas:
        if chars_used >= max_chars_total:
            _log(f"MULTI-LEER: límite alcanzado tras {len(chunks)} archivos")
            break
        
        content = _leer_docx(path)
        if content:
            # Truncar cada archivo a su cuota
            truncated = content[:chars_per_file]
            chunks.append(truncated)
            chars_used += len(truncated)
    
    result = f" {len(chunks)} documentos del paciente:\n\n" + "\n\n---\n\n".join(chunks)
    _log(f"MULTI-LEER: {len(chunks)} archivos, {chars_used} chars en {time.time()-t0:.1f}s")
    return result


def _leer_docx(ruta: Path) -> str:
    """Lee .docx rápido, primeras 60 líneas + 5 tablas."""
    _log(f"LEER: {ruta.name}...")
    t0 = time.time()
    
    if not ruta.exists():
        _log(f"LEER: [X] no existe")
        return ""
    
    if ruta.suffix.lower() == '.txt':
        try:
            content = ruta.read_text(encoding='utf-8', errors='ignore')[:2000]
            _log(f"LEER: txt {len(content)} chars en {time.time()-t0:.1f}s")
            return f" {ruta.name}:\n{content}"
        except Exception as e:
            _log(f"LEER: error txt: {e}")
            return ""
    
    if ruta.suffix.lower() != '.docx':
        return ""
    
    try:
        from docx import Document
        doc = Document(str(ruta))
        _log(f"LEER: {len(doc.paragraphs)}p, {len(doc.tables)}t en {time.time()-t0:.1f}s")
        
        lines = [f" {ruta.name} | {len(doc.paragraphs)} párrafos, {len(doc.tables)} tablas\n"]
        
        # Párrafos con estructura (más para docs pesados de varias páginas)
        for p in doc.paragraphs[:60]:
            t = p.text.strip()
            if not t: continue
            if re.match(r'^\d+[\.\)]\s', t):
                lines.append(f"\n {t}")
            elif t.isupper() and len(t) < 80:
                lines.append(f"\n {t}")
            else:
                lines.append(f"  {t[:150]}")
        
        # Tablas (primeras 8 filas de c/u, máx 5 tablas)
        for ti, tabla in enumerate(doc.tables[:5]):
            lines.append(f"\n Tabla {ti+1}:")
            for ri, row in enumerate(tabla.rows[:5]):
                cells = [cell.text.strip()[:50] for cell in row.cells[:6] if cell.text.strip()]
                if cells:
                    lines.append(f"  {' | '.join(cells)}")
        
        result = "\n".join(lines)[:6000]
        _log(f"LEER: {len(result)} chars total en {time.time()-t0:.1f}s")
        return result
    
    except ImportError:
        _log("LEER: [X] python-docx no instalado")
        return ""
    except Exception as e:
        _log(f"LEER: [X] {type(e).__name__}: {e}")
        return ""


def _llamar_llm(mensaje: str, cc: str = "", archivo_doc: str = "", workspace_files: str = "",
                historial=None, ctx_sistema: str = "", modo_simple: bool = False) -> str:
    """Llama al LLM. DeepSeek v4: modelo de razonamiento - necesita max_tokens alto."""
    if not API_KEY:
        _log("LLM: sin API key")
        return "No tengo conexion con mi cerebro (falta API key). Revisa el archivo .env"

    prompt = SYSTEM_PROMPT_SIMPLE if modo_simple else SYSTEM_PROMPT
    messages = [{"role": "system", "content": prompt}]

    # Historial (ultimos 6 turnos)
    if historial:
        for h in historial[-6:]:
            rol = h.get("rol", "")
            contenido = (h.get("contenido", "") or "")[:2000]
            if rol == "usuario":
                messages.append({"role": "user", "content": contenido})
            elif rol == "asistente":
                messages.append({"role": "assistant", "content": contenido})

    # Construir contexto
    partes = []
    if ctx_sistema:
        partes.append(ctx_sistema)
    partes.append(f"\n=== MENSAJE DE SANDRA ===\n{mensaje}")
    if archivo_doc:
        partes.append(f"\n\n[Documentos encontrados]:\n{archivo_doc}")
    if cc:
        partes.append(f"\n[CC: {cc}]")
    if workspace_files:
        partes.append(f"\n\nArchivos disponibles:\n{workspace_files}")
    else:
        partes.append("\n\nNo tengo acceso a tu carpeta de trabajo.")

    ctx = "\n".join(partes)
    messages.append({"role": "user", "content": ctx})

    total_chars = sum(len(m["content"]) for m in messages)
    _log(f"LLM: {total_chars} chars -> {MODEL}")
    t0 = time.time()

    # DeepSeek v4 es modelo de razonamiento -- PIENSA antes de responder.
    # Con archivos pesados (varias páginas, tablas, datos clínicos) el razonamiento
    # puede consumir 10K-15K tokens. Necesita margen holgado para pensar + responder.
    # Sin límites artificiales que fuercen finish=length y respuesta vacía.
    max_tokens_values = [16000, 24000]  # intento 1: 16000, intento 2: 24000 (reducido para lotes)
    
    for intento in range(2):
        try:
            max_tok = max_tokens_values[intento]
            resp = requests.post(
                f"{BASE_URL}/chat/completions",
                headers={
                    "Authorization": f"Bearer {API_KEY}",
                    "Content-Type": "application/json",
                    "User-Agent": "Mozilla/5.0",
                },
                json={
                     "model": os.getenv("LLM_MODEL_EXTRACCION", "deepseek-v4-flash"),
                    "messages": messages,
                    "temperature": 0.7,
                    "max_tokens": max_tok,
                },
                timeout=300
            )
            elapsed = time.time() - t0
            usage_info = ""
            
            if resp.status_code == 200:
                body = resp.json()
                choices = body.get("choices", [])
                if choices:
                    choice = choices[0]
                    msg = choice.get("message", {})
                    content = (msg.get("content", "") or "").strip()
                    reasoning = (msg.get("reasoning_content", "") or "").strip()
                    finish = choice.get("finish_reason", "?")
                    usage = body.get("usage", {})
                    total_tok = usage.get("total_tokens", "?")
                    
                    _log(f"LLM: content={len(content)}chars reasoning={len(reasoning)}chars "
                         f"finish={finish} tokens={total_tok} max_tokens={max_tok} "
                         f"(intento {intento+1})")
                    
                    if content:
                        return content
                    
                    # Si content vacío pero finish=length -> los tokens se agotaron en razonamiento
                    if finish == "length" and reasoning:
                        _log(f"LLM: finish=length, {len(reasoning)} chars de razonamiento. "
                             f"Reintentando con max_tokens={max_tokens_values[1] if intento==0 else 'N/A'}...")
                        if intento == 0:
                            continue  # reintentar con más tokens
                        else:
                            # Último intento: extraer algo del razonamiento
                            lines = [l.strip() for l in reasoning.split('\n') if l.strip()]
                            last_lines = lines[-3:] if len(lines) > 3 else lines
                            fallback = " Estoy procesando tu solicitud. " + " ".join(last_lines)[:300]
                            _log(f"LLM: usando fallback de razonamiento ({len(fallback)} chars)")
                            return fallback
                    
                    if finish == "stop" and not content:
                        _log("LLM: finish=stop pero content vacío -- modelo no generó respuesta")
                        if intento == 0:
                            # Reintentar con prompt más directo
                            messages.append({"role": "user", "content": "Responde directamente en español, por favor."})
                            continue
                        return " El modelo no generó respuesta. Intentá con un mensaje más específico."
                    
                    # Otro caso raro
                    _log(f"LLM: caso no manejado -- content={len(content)} finish={finish}")
                    if intento == 0:
                        continue
                else:
                    _log(f"LLM: sin choices: {json.dumps(body)[:200]}")
                    return " El servicio de IA no devolvió respuesta. Intentá de nuevo."
            else:
                _log(f"LLM: HTTP {resp.status_code}: {resp.text[:200]}")
                return f" Error del servicio IA (HTTP {resp.status_code})."
        
        except requests.exceptions.Timeout:
            _log(f"LLM: timeout tras {time.time()-t0:.1f}s")
            if intento == 0:
                continue
            return "⏰ El servicio está tardando mucho. ¿Podés intentar con un mensaje más corto?"
        except Exception as e:
            _log(f"LLM: {type(e).__name__}: {e}")
            if intento == 0:
                continue
            return f"[X] Error de conexión: {type(e).__name__}."

    _log("LLM: [X] agotados los reintentos")
    return " No pude generar respuesta después de varios intentos. ¿Probás con un mensaje más concreto?"


def _procesar_via_workers(archivos: list, cc: str) -> str:
    """
    Reemplazo de _procesar_en_lotes que aísla cada llamada LLM en subprocess.

    Cada subprocess procesa 1 archivo, imprime JSON con datos extraídos, muere.
    El OS libera la RAM del reasoning_content de DeepSeek al terminar el proceso.

    Para extracción usa LLM_MODEL_EXTRACCION (deepseek-v4-flash por defecto, sin razonamiento).
    Los resultados se concatenan y se devuelven como string para inyectar al contexto.
    """
    modelo_extraccion = os.getenv("LLM_MODEL_EXTRACCION", "deepseek-v4-flash")
    _log(f"WORKERS: {len(archivos)} archivos con modelo={modelo_extraccion}")
    t0 = time.time()

    todos_datos = []
    for i, (path, size_kb, mtime) in enumerate(archivos):
        try:
            proc = subprocess.run(
                [sys.executable, "-m", "backend.lote_worker",
                 "--tipo", "extraer",
                 "--archivos", str(path),
                 "--modelo", modelo_extraccion,
                 "--cc", cc],
                capture_output=True, text=True, timeout=180,
                cwd=str(Path(__file__).parent.parent),
            )
            if proc.returncode == 0 and proc.stdout:
                resultado = json.loads(proc.stdout)
                if resultado.get("ok") and resultado.get("datos_raw"):
                    todos_datos.append(
                        f" DOC {i+1}/{len(archivos)} ({path.name}) \n"
                        f"{resultado['datos_raw']}"
                    )
                    tokens = resultado.get("tokens", {})
                    _log(f"WORKER {i+1}: {tokens.get('total', '?')}t en {resultado.get('tiempo_s', '?')}s")
            else:
                _log(f"WORKER {i+1}: FAILED exit={proc.returncode} stderr={proc.stderr[:200]}")
                todos_datos.append(f" DOC {i+1} ({path.name}) \n[ERROR de extracción]")
        except subprocess.TimeoutExpired:
            _log(f"WORKER {i+1}: TIMEOUT")
            todos_datos.append(f" DOC {i+1} ({path.name}) \n[TIMEOUT]")
        except Exception as e:
            _log(f"WORKER {i+1}: {type(e).__name__}: {e}")
            todos_datos.append(f" DOC {i+1} ({path.name}) \n[{type(e).__name__}]")

    elapsed = time.time() - t0
    resultado_str = "\n\n".join(todos_datos)
    _log(f"WORKERS: {len(archivos)} procesados en {elapsed:.1f}s -> {len(resultado_str)} chars")

    return f"""[DATOS EXTRAÍDOS DE {len(archivos)} FORMATOS POR WORKERS AISLADOS]

{resultado_str}

[!] Tu tarea con estos datos:
1. CRUZAR campos entre documentos, detectar inconsistencias
2. VERIFICAR contra portales (marcar [OK]/[!]/[X])
3. ORGANIZAR información por secciones
4. CORREGIR discrepancias detectadas
5. COMPLETAR campos faltantes con prioridad"""


def _procesar_en_lotes(archivos: list, cc: str, mensaje_original: str) -> str:
    """
    Procesa muchos archivos en lotes de 3-4. Cada lote: extraer datos clave.
    Luego junta todo y lo devuelve como contexto comprimido para la llamada final.
    Evita OOM en el contenedor sin sacrificar calidad.
    """
    TAMANO_LOTE = 2
    lotes = [archivos[i:i+TAMANO_LOTE] for i in range(0, len(archivos), TAMANO_LOTE)]
    _log(f"LOTES: {len(archivos)} archivos -> {len(lotes)} lotes de ~{TAMANO_LOTE}")
    
    EXTRACT_PROMPT = """Extrae TODOS los datos clínicos y administrativos de estos documentos.
Devuelve SOLO datos, en este formato EXACTO:

PACIENTE: [nombre completo] | CC: [número]
SINIESTRO: [id] | FECHA: [dd/mm/aaaa] | TIPO: [AT/EL/EG]
DIAGNÓSTICO: [código CIE-10] - [descripción]
EMPRESA: [nombre] | NIT: [número]
FECHAS: consulta=[fecha], valoración=[fecha], cierre=[fecha]
SEGMENTO: [segmento corporal]
OBSERVACIONES: [datos relevantes, notas de Sandra, campos incompletos]
CAMPOS_VACIOS: [lista de campos que faltan en este documento]

Si un dato no aparece, escribe FALTA."""
    
    todos_los_datos = []
    t0 = time.time()
    
    for i, lote in enumerate(lotes):
        _log(f"LOTE {i+1}/{len(lotes)}: {len(lote)} archivos")
        
        # Leer archivos del lote
        contenido_lote = _leer_multiples_docx(lote, max_chars_total=3000)
        if not contenido_lote:
            continue
        
        # Llamar al LLM para extraer datos
        if not API_KEY:
            todos_los_datos.append(f"[LOTE {i+1}] Sin API key -- no se pudo procesar")
            continue
        
        try:
            resp = requests.post(
                f"{BASE_URL}/chat/completions",
                headers={"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json", "User-Agent": "Mozilla/5.0"},
                json={
                     "model": os.getenv("LLM_MODEL_EXTRACCION", "deepseek-v4-flash"),
                    "messages": [
                        {"role": "system", "content": EXTRACT_PROMPT},
                        {"role": "user", "content": contenido_lote}
                    ],
                    "temperature": 0.3,
                    "max_tokens": 4000,
                },
                timeout=120
            )
            if resp.status_code == 200:
                body = resp.json()
                choices = body.get("choices", [])
                if choices:
                    content = (choices[0].get("message", {}).get("content", "") or "").strip()
                    if content:
                        todos_los_datos.append(f" LOTE {i+1} \n{content}")
                        _log(f"LOTE {i+1}: extraídos {len(content)} chars")
                    else:
                        _log(f"LOTE {i+1}: respuesta vacía")
                else:
                    _log(f"LOTE {i+1}: sin choices")
            else:
                _log(f"LOTE {i+1}: HTTP {resp.status_code}")
        except Exception as e:
            _log(f"LOTE {i+1}: error {e}")
    
    elapsed = time.time() - t0
    resultado = "\n\n".join(todos_los_datos)
    _log(f"LOTES: {len(lotes)} procesados en {elapsed:.1f}s -> {len(resultado)} chars comprimidos")
    
    # Devolver datos comprimidos + instrucción para el cruce final
    return f"""[DATOS EXTRAÍDOS DE {len(archivos)} FORMATOS EN {len(lotes)} LOTES]

{resultado}

[!] Con estos datos extraídos, tu tarea es:
1. CRUZAR: comparar campos entre lotes, detectar inconsistencias
2. VERIFICAR: marcar qué confirmar en Medifolios y Positiva
3. ORGANIZAR: estructurar la información por secciones
4. CORREGIR: señalar discrepancias y sugerir correcciones
5. COMPLETAR: listar campos faltantes con prioridad

Mensaje original de Sandra: "{mensaje_original[:200]}" """


def _get_data_dir() -> Path:
    """Retorna el directorio de datos (configurable vía STORAGE_DIR)."""
    storage = Path(os.getenv("STORAGE_DIR", str(Path(__file__).parent.parent / "storage")))
    return storage / "data"


def _cargar_datos_verificados(cc: str) -> Optional[dict]:
    """Busca storage/data/*{cc}*completo.json y lo carga."""
    if not cc:
        return None
    data_dir = _get_data_dir()
    if not data_dir.exists():
        return None
    pattern = f"*{cc}*completo.json"
    matches = sorted(data_dir.glob(pattern), key=os.path.getmtime, reverse=True)
    if not matches:
        return None
    try:
        with open(matches[0], encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def _formatear_datos_verificados(datos: dict) -> str:
    """Genera texto con marcas [OK]/[!]/[X] a partir del JSON de orquestador."""
    if not datos:
        return ""
    partes = []
    medi = datos.get("medifolios", {})
    pos = datos.get("positiva", {})

    siniestro_medi = medi.get("siniestro_medi", "[SIN DATO]")
    siniestros_pos = pos.get("siniestros", [])
    siniestro_pos_id = siniestros_pos[0].get("id", "[SIN DATO]") if siniestros_pos else "[SIN DATO]"

    if siniestro_medi == siniestro_pos_id and siniestro_medi != "[SIN DATO]":
        partes.append(f"[OK] Siniestro: {siniestro_pos_id} (coincide Medifolios + Positiva)")
    elif siniestro_medi and siniestro_pos_id and siniestro_medi != "[SIN DATO]" and siniestro_pos_id != "[SIN DATO]":
        partes.append(f"[X] Siniestro DISCREPANCIA: Medifolios={siniestro_medi} vs Positiva={siniestro_pos_id}")
        partes.append("   -> Prevalece Positiva (fuente oficial)")
    elif siniestro_pos_id and siniestro_pos_id != "[SIN DATO]":
        partes.append(f"[OK] Siniestro: {siniestro_pos_id} (Positiva)")
        if siniestro_medi == "[SIN DATO]" or not siniestro_medi:
            partes.append("   [!] No disponible en Medifolios")
    else:
        partes.append("[!] Siniestro no disponible en ningún portal")

    nombre_medi = medi.get("nombre1") or medi.get("nombre", "")
    apellido_medi = medi.get("apellido1", "")
    nombre = f"{nombre_medi} {apellido_medi}".strip()
    if nombre:
        partes.append(f"[OK] Paciente: {nombre}")
    if medi.get("telefono"):
        partes.append(f"[OK] Teléfono: {medi['telefono']}")
    if medi.get("direccion"):
        partes.append(f"[OK] Dirección: {medi['direccion']}")

    discrepancias = datos.get("_meta", {}).get("discrepancias", [])
    for d in discrepancias:
        partes.append(f"[X] {d.get('campo', 'desconocido')}: Medifolios={d.get('medifolios','?')} vs Positiva={d.get('positiva','?')}")
        if d.get("resolucion"):
            partes.append(f"   -> {d['resolucion']}")

    if not partes:
        partes.append("ℹ No hay datos verificados disponibles")

    fecha = datos.get("_meta", {}).get("extraido_en", "")
    header = f" DATOS VERIFICADOS (extraído: {fecha})" if fecha else " DATOS VERIFICADOS"
    return f"\n\n\n{header}\n" + "\n".join(partes) + "\n\n"


def _construir_contexto_sistema() -> str:
    """Captura estado actual del sistema para inyectar al LLM."""
    lineas = ["=== ESTADO ACTUAL DEL SISTEMA ==="]
    
    try:
        db_path = os.getenv("WORKFLOW_DB_PATH", "./storage/workflow.db")
        from backend.task_db import TaskDB
        db = TaskDB(db_path)
        activas = db.listar_activos()
        if activas:
            lineas.append(f"\nPROCESANDO AHORA ({len(activas)} tarea(s)):")
            PASOS = {1:"Transcribiendo audio",2:"Buscando datos",3:"Leyendo notas",
                     4:"Leyendo formatos",5:"Sintetizando (IA)",6:"Verificando portales",
                     7:"Generando formatos",8:"QA",9:"PDF",10:"Notificando"}
            for t in activas[:3]:
                label = PASOS.get(t.get("paso_actual",0), t.get("estado","?"))
                lineas.append(f"  CC {t['paciente_cc']}: {label} (paso {t.get('paso_actual','?')}/10)")
        else:
            lineas.append("\nSin tareas activas - sistema libre")
    except Exception:
        lineas.append("\n[Tareas: no disponible]")

    try:
        from backend.notificador import api_get_agenda
        agenda = api_get_agenda()
        if agenda and agenda.get("citas"):
            citas = agenda["citas"]
            lineas.append(f"\nAGENDA HOY ({len(citas)} cita(s)):")
            for c in citas[:6]:
                lineas.append(f"  {c.get('hora','?')} - {c.get('paciente','?')}")
        else:
            lineas.append("\nAGENDA: No hay agenda cargada")
    except Exception:
        pass

    try:
        import glob as _glob
        storage = Path(os.getenv("STORAGE_DIR", "./storage"))
        data_dir = storage / "data"
        jsons = _glob.glob(str(data_dir / "*-completo.json"))
        lineas.append(f"\nPACIENTES: {len(jsons)} en sistema")
    except Exception:
        pass

    try:
        docs_dir = storage / "docs"
        docs = sorted(docs_dir.glob("*.docx"), key=lambda x: x.stat().st_mtime, reverse=True)
        if docs:
            lineas.append(f"\nFORMATOS RECIENTES:")
            for d in docs[:5]:
                lineas.append(f"  {d.name}")
    except Exception:
        pass

    tomy_on = os.getenv("TOMY_COMPLETO_ENABLED", "false").lower() == "true"
    lineas.append(f"\nPipeline: {'ACTIVO' if tomy_on else 'INACTIVO'}")

    return "\n".join(lineas)


# --- Intent detection ---

PORTAL_KEYWORDS = [
    "verifica", "verificar", "compara", "comparar", "coincide",
    "chequea", "confronta", "revisa en el portal", "que dice positiva",
    "que dice medifolios", "cruza con", "corrobora", "esta bien el",
    "busca en positiva", "extrae de",
]

CORRECCION_KEYWORDS = [
    "corrige", "cambia", "cambiar", "actualiza", "actualizar",
    "modifica", "el nombre es", "el dato correcto es", "esta mal el",
    "reemplaza", "pon", "quita", "agrega",
]


def _detectar_intencion_portal(mensaje: str, cc: str = ""):
    msg_lower = mensaje.lower()
    if any(k in msg_lower for k in PORTAL_KEYWORDS):
        return {"tipo": "verificar", "cc": cc or _extraer_cc_de_texto(mensaje),
                "forzar": "actualiza" in msg_lower or "de nuevo" in msg_lower}
    return None


def _es_intencion_correccion(mensaje: str) -> bool:
    return any(k in mensaje.lower() for k in CORRECCION_KEYWORDS)


def _extraer_cc_de_texto(texto: str) -> str:
    m = re.search(r"\b(\d{6,12})\b", texto.replace("'","").replace(".","").replace(",",""))
    return m.group(1) if m else ""


def _formatear_verificacion_para_llm(resultado: dict) -> str:
    cc = resultado.get("cc", "?")
    lineas = [f"\n=== VERIFICACION PORTALES - CC {cc} ==="]
    if resultado.get("error"):
        lineas.append(f"ERROR: {resultado['error']}")
        return "\n".join(lineas)
    for campo in resultado.get("campos", []):
        icono = "OK" if campo["coincide"] else "X"
        lineas.append(f"{icono} {campo['campo']}: {campo['medifolios']} = {campo['positiva']}")
    return "\n".join(lineas)


def procesar_mensaje(mensaje: str, paciente_cc: str = "", historial: list = None) -> dict:
    """Procesa mensaje del chat. Soporta hasta 7 formatos por paciente."""
    _log(f" '{mensaje[:100]}' ")
    t_total = time.time()

    if not mensaje.strip():
        return {"contenido": "¿En qué te ayudo, Sandra? ", "accion": None}

    # 1. Workspace (siempre, rápido)
    workspace_files = _listar_workspace()

    # 2. Detectar CC (sin limpiar espacios -- mantener word boundaries)
    cc = paciente_cc
    if not cc:
        m = re.search(r'\b(\d{6,12})\b', mensaje.replace("'", "").replace(".", "").replace(",", ""))
        if m: cc = m.group(1)

    # 3. NUEVO: Cargar datos verificados si Fase A activa
    datos_verificados = None
    sugerir_extraccion = False
    if os.getenv("FASE_A_ENABLED", "false").lower() == "true" and cc:
        datos_verificados = _cargar_datos_verificados(cc)
        if datos_verificados:
            _log(f"VERIFIED: datos cargados para CC {cc}")
        else:
            sugerir_extraccion = True
            _log(f"VERIFIED: sin datos para CC {cc}")

    # 4. Buscar archivos -- PRIMERO multi (todos los del paciente), luego individual
    archivos_paciente = _buscar_archivos_paciente(cc=cc, mensaje=mensaje)
    
    if archivos_paciente and len(archivos_paciente) >= 3:
        # MODO WORKERS (Fase A) o LOTES (legacy)
        if os.getenv("FASE_A_ENABLED", "false").lower() == "true":
            doc_content = _procesar_via_workers(archivos_paciente, cc)
            archivo_nombre = f"{len(archivos_paciente)} archivos (workers)"
        else:
            doc_content = _procesar_en_lotes(archivos_paciente, cc, mensaje)
            archivo_nombre = f"{len(archivos_paciente)} archivos (lotes-legacy)"
    elif archivos_paciente:
        # Pocos archivos: leer directo
        doc_content = _leer_multiples_docx(archivos_paciente, max_chars_total=8000)
        archivo_nombre = f"{len(archivos_paciente)} archivos"
    else:
        # MODO INDIVIDUAL: buscar un solo archivo
        archivo = _buscar_archivo(mensaje)
        doc_content = _leer_docx(archivo) if archivo else ""
        archivo_nombre = archivo.name if archivo else None

    # 4.5. NUEVO: detectar comando "organiza este formato"
    if any(kw in mensaje.lower() for kw in ["organiza este formato", "organízame este", "organizame el formato"]):
        archivo = _buscar_archivo(mensaje)
        if archivo and cc:
            try:
                from backend.organizador_formato import organizar_formato
                resultado_org = organizar_formato(str(archivo), cc)
                if resultado_org.get("ok"):
                    doc_content = (doc_content or "") + f"\n\n[FORMATO ORGANIZADO]:\n{json.dumps(resultado_org, ensure_ascii=False, indent=2)}"
            except Exception as e:
                _log(f"ORGANIZAR: error {e}")

    # 5. Inyectar datos verificados al contexto del LLM
    if datos_verificados:
        doc_content = (doc_content or "") + _formatear_datos_verificados(datos_verificados)

    # 5.5 Dame formatos - links reales, no inventar
    if cc and any(kw in mensaje.lower() for kw in [
        "dame los formatos", "pasame los formatos", "descargar formatos",
        "manda los formatos", "los formatos", "documentos generados",
    ]):
        DOCS_DIR = Path(os.getenv("STORAGE_DIR", "./storage")) / "docs"
        formatos_files = sorted(DOCS_DIR.glob(f"*{cc}*.docx"), key=lambda x: x.stat().st_mtime, reverse=True)
        if formatos_files:
            links = "\n".join(f"- [{f.name}](/api/download/{f.name})" for f in formatos_files[:10])
            return {"contenido": f"Formatos para CC {cc}:\n\n{links}\n\nClick para descargar.", "accion": "formatos", "archivo": None}
        else:
            return {"contenido": f"No hay formatos generados para CC {cc}. Queres que procese un audio?", "accion": None}

    # 5.6 Verificacion portales
    contexto_portal = ""
    intencion = _detectar_intencion_portal(mensaje, cc)
    if intencion and intencion.get("cc"):
        _log(f"PORTAL: verificar CC {intencion['cc']}")
        try:
            from backend.portal_verificador import PortalVerificador
            v = PortalVerificador()
            resultado = v.verificar(intencion["cc"], forzar_extraccion=intencion.get("forzar", False))
            contexto_portal = _formatear_verificacion_para_llm(resultado)
        except Exception as e:
            contexto_portal = f"\nNo pude verificar portales: {e}\n"

    # 5.7 Correccion de datos
    if cc and _es_intencion_correccion(mensaje):
        try:
            from backend.correction_resolver import interpretar_correccion
            from backend.aplicador_correccion import aplicar_correccion
            correccion = interpretar_correccion(mensaje, cc)
            if correccion and correccion.get("confianza", 0) >= 0.7:
                resultado_corr = aplicar_correccion(cc, correccion)
                if resultado_corr.get("ok"):
                    formatos = resultado_corr.get("formatos_regenerados", [])
                    links = "\n".join(f"- [{f}](/api/download/{f})" for f in formatos)
                    return {
                        "contenido": f"OK. Corregi '{correccion['campo']}' a '{correccion['valor']}'.\n\n{links}",
                        "accion": "correccion", "archivo": None
                    }
        except Exception as e:
            _log(f"CORRECCION: error {e}")

    # 6. Armar mensaje final
    if contexto_portal:
        mensaje_con_sugerencia = mensaje + "\n\n" + contexto_portal
    elif sugerir_extraccion:
        mensaje_con_sugerencia = mensaje + "\n\nNo tengo datos verificados de este paciente. Queres que verifique en Medifolios y Positiva?"
    else:
        mensaje_con_sugerencia = mensaje

    # 7. Llamar LLM
    ctx_sistema = _construir_contexto_sistema()
    modo_simple = not doc_content and not cc
    respuesta = _llamar_llm(mensaje_con_sugerencia, cc, doc_content, workspace_files,
                            historial=historial, ctx_sistema=ctx_sistema, modo_simple=modo_simple)
    
    total_time = time.time() - t_total
    _log(f" Listo en {total_time:.1f}s ")

    return {
        "contenido": respuesta,
        "accion": "documento" if doc_content else None,
        "archivo": archivo_nombre,
        "workspace_accesible": bool(workspace_files),
        "tiempo": f"{total_time:.1f}s",
    }


def diagnosticar_workspace() -> dict:
    """Devuelve info de diagnóstico del workspace (para /api/workspace)."""
    ws = _get_active_workspace()
    entries = _fast_scandir(ws, max_depth=2, timeout=5.0)
    
    archivos = []
    for path, is_dir, size_kb, mtime in entries:
        archivos.append({
            "nombre": path.name,
            "tipo": "carpeta" if is_dir else "archivo",
            "ruta": str(path.relative_to(ws)),
            "tamano_kb": size_kb,
            "modificado": mtime,
        })
    
    return {
        "workspace_principal": str(WORKSPACE_DIR),
        "workspace_principal_existe": WORKSPACE_DIR.exists(),
        "workspace_activo": str(ws),
        "total_archivos": len([a for a in archivos if a["tipo"] == "archivo"]),
        "total_carpetas": len([a for a in archivos if a["tipo"] == "carpeta"]),
        "archivos": archivos[:50],
        "api_key_configurada": bool(API_KEY),
        "modelo": MODEL,
    }
