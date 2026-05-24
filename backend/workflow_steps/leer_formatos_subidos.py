"""
Paso 4: Leer formatos DOCX que Sandra subió como referencia para esta task.

Sandra a veces sube un formato antiguo desorganizado para que Tomy lo
use como base de información (que cruzará con verificados).

Archivos se guardan en: storage/formatos_subidos/{task_id}/*.docx
"""
import os
import sys
from pathlib import Path
from datetime import datetime
from typing import List, Dict


def _log(msg: str):
    print(f"[FORMATOS_SUB {datetime.now().strftime('%H:%M:%S')}] {msg}", file=sys.stderr, flush=True)


def leer_formatos_subidos(task_id: str, max_chars: int = 5000) -> List[Dict]:
    storage = Path(os.getenv("STORAGE_DIR", "./storage"))
    formato_dir = storage / "formatos_subidos" / task_id
    if not formato_dir.exists():
        return []

    formatos = []
    try:
        from docx import Document
        for path in formato_dir.glob("*.docx"):
            try:
                doc = Document(str(path))
                texto = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                texto = texto[:max_chars]
                for tabla in doc.tables[:10]:
                    for row in tabla.rows[:20]:
                        celdas = [c.text.strip()[:60] for c in row.cells if c.text.strip()]
                        if celdas:
                            texto += "\n" + " | ".join(celdas)
                formatos.append({"nombre": path.name, "ruta": str(path), "contenido": texto[:max_chars]})
            except Exception as e:
                _log(f"Error leyendo {path.name}: {e}")
    except ImportError:
        _log("python-docx no disponible - no se leen formatos subidos")

    return formatos


def ejecutar(task_id: str) -> dict:
    formatos = leer_formatos_subidos(task_id)
    return {"formatos_subidos": formatos, "total": len(formatos)}
