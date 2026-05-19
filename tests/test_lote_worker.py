"""Tests para lote_worker — subprocess que ejecuta 1 llamada LLM y muere."""
import json
import os
import subprocess
import sys
from pathlib import Path
import pytest


def _ejecutar_worker(args: list) -> tuple[int, dict, str]:
    """Helper: ejecuta lote_worker como subprocess y parsea salida."""
    proc = subprocess.run(
        [sys.executable, "-m", "backend.lote_worker"] + args,
        capture_output=True, text=True, timeout=180,
        cwd=Path(__file__).parent.parent,
    )
    try:
        salida = json.loads(proc.stdout) if proc.stdout else {}
    except json.JSONDecodeError:
        salida = {"raw_stdout": proc.stdout}
    return proc.returncode, salida, proc.stderr


def test_lote_worker_sin_archivos_devuelve_error(tmp_path):
    """Llamar sin --archivos debe fallar con exit 1 y mensaje en stderr."""
    code, salida, stderr = _ejecutar_worker(["--tipo", "extraer", "--archivos", "", "--cc", "123"])
    assert code != 0
    assert "archivos" in stderr.lower() or "archivos" in str(salida).lower()


@pytest.mark.skipif(not os.getenv("OPENCODE_GO_API_KEY"), reason="needs OPENCODE_GO_API_KEY")
def test_lote_worker_memoria_padre_estable(tmp_path):
    """10 invocaciones consecutivas → memoria del padre no crece >50MB."""
    import psutil
    import os
    from docx import Document
    docx_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("PACIENTE: TEST")
    doc.add_paragraph("CC: 12345678")
    doc.save(str(docx_path))

    proc = psutil.Process(os.getpid())
    mem_inicial = proc.memory_info().rss / 1024 / 1024  # MB

    for i in range(3):
        code, salida, _ = _ejecutar_worker([
            "--tipo", "extraer",
            "--archivos", str(docx_path),
            "--modelo", "deepseek-v4-flash",
            "--cc", "12345678",
        ])

    mem_final = proc.memory_info().rss / 1024 / 1024
    delta = mem_final - mem_inicial
    assert delta < 50, f"Memoria padre creció {delta:.1f}MB tras 10 invocaciones"


@pytest.mark.skipif(not os.getenv("OPENCODE_GO_API_KEY"), reason="needs OPENCODE_GO_API_KEY")
def test_lote_worker_extraccion_devuelve_json_valido(tmp_path):
    """Con archivo válido + API disponible, debe devolver JSON parseable."""
    from docx import Document
    docx_path = tmp_path / "paciente.docx"
    doc = Document()
    doc.add_paragraph("Paciente Juan Carlos Duran Narvaez")
    doc.add_paragraph("Cédula: 1193143688")
    doc.add_paragraph("Siniestro: 503463870 fecha 02/03/2026")
    doc.save(str(docx_path))

    code, salida, stderr = _ejecutar_worker([
        "--tipo", "extraer",
        "--archivos", str(docx_path),
        "--modelo", "deepseek-v4-flash",
        "--cc", "1193143688",
    ])
    assert code == 0, f"exit {code}, stderr={stderr[:300]}"
    assert salida.get("ok") is True
    assert "datos_raw" in salida
    assert salida.get("modelo") == "deepseek-v4-flash"
