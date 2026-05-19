"""Tests para leer_formatos_subidos.py."""
import sys
from pathlib import Path
from docx import Document
import pytest

sys.path.insert(0, str(Path(__file__).parent.parent.parent))


def test_lee_formatos_subidos_para_task(tmp_path, monkeypatch):
    storage = tmp_path / "storage"
    formato_dir = storage / "formatos_subidos" / "abc12345"
    formato_dir.mkdir(parents=True)

    doc = Document()
    doc.add_paragraph("Análisis de exigencias antiguo")
    doc.add_paragraph("Datos incompletos del paciente")
    doc.save(str(formato_dir / "analisis_viejo.docx"))

    monkeypatch.setenv("STORAGE_DIR", str(storage))
    from backend.workflow_steps.leer_formatos_subidos import leer_formatos_subidos

    formatos = leer_formatos_subidos("abc12345")
    assert len(formatos) == 1
    assert "Análisis de exigencias antiguo" in formatos[0]["contenido"]


def test_devuelve_vacio_si_no_hay_formatos(tmp_path, monkeypatch):
    monkeypatch.setenv("STORAGE_DIR", str(tmp_path))
    from backend.workflow_steps.leer_formatos_subidos import leer_formatos_subidos
    formatos = leer_formatos_subidos("xyz999")
    assert formatos == []
