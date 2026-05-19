"""Tests para organizador_formato.py."""
import sys
from pathlib import Path
from unittest.mock import patch, MagicMock
import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))


def test_organiza_formato_crudo(tmp_path, monkeypatch):
    monkeypatch.setenv("STORAGE_DIR", str(tmp_path))

    from docx import Document
    crudo = tmp_path / "crudo.docx"
    doc = Document()
    doc.add_paragraph("Paciente Juan, dolor lumbar, dice que le duele al agacharse")
    doc.add_paragraph("Empresa: ACME")
    doc.save(str(crudo))

    with patch("backend.organizador_formato._llamar_llm_organizador") as mock_llm:
        mock_llm.return_value = {
            "ok": True,
            "datos_organizados": {
                "paciente": {"nombre": "JUAN", "documento": "1193143688"},
                "diagnostico": "Dolor lumbar",
            },
            "campos_faltantes": ["empresa.nit", "siniestro.id_siniestro"],
        }
        from backend.organizador_formato import organizar_formato
        resultado = organizar_formato(str(crudo), paciente_cc="1193143688")

    assert resultado["ok"]
    assert resultado["datos_organizados"]["paciente"]["nombre"] == "JUAN"
    assert "empresa.nit" in resultado["campos_faltantes"]
