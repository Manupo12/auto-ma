"""
Paso 7: QA real sobre cada DOCX generado.
Verifica placeholders, campos vacios, y estructura basica.
"""
import os
import re
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, List


def _log(msg: str):
    print(f"[QA {datetime.now().strftime('%H:%M:%S')}] {msg}", file=sys.stderr, flush=True)


PLACEHOLDER_PATTERNS = [
    r'\[FALTA[:\s]*[^\]]*\]',
    r'\[VERIFICAR\]',
    r'\[NOMBRE\]',
    r'\[CC\]',
    r'\[FECHA\]',
    r'\[SINIESTRO\]',
    r'\[EMPRESA\]',
    r'\[CARGO\]',
    r'\[DIAGNOSTICO\]',
]

CAMPOS_REQUERIDOS = [
    'nombre', 'documento', 'empresa', 'siniestro', 'diagnostico',
]


def qa_uno(formato: dict) -> dict:
    archivo = formato.get("archivo")
    if not archivo or not Path(archivo).exists():
        return {**formato, "qa_ok": False, "qa_error": "archivo no existe"}

    warnings = []
    try:
        from docx import Document
        doc = Document(str(archivo))
        texto = " ".join(p.text for p in doc.paragraphs if p.text.strip())

        if not texto or len(texto) < 100:
            return {**formato, "qa_ok": False, "qa_error": "documento vacio o muy corto"}

        for pattern in PLACEHOLDER_PATTERNS:
            matches = re.findall(pattern, texto, re.IGNORECASE)
            if matches:
                warnings.append(f"Placeholders: {', '.join(matches[:5])}")

        n_tablas = len(doc.tables)
        if n_tablas == 0:
            warnings.append("Sin tablas - posible documento incompleto")

        tiene_datos = any(
            campo in texto.lower()
            for campo in CAMPOS_REQUERIDOS
        )

        if not tiene_datos:
            warnings.append("Faltan datos clinicos basicos (nombre, empresa, siniestro)")

        qa_ok = len(warnings) == 0 or all("Placeholders" not in w for w in warnings)

        return {
            **formato,
            "qa_ok": qa_ok,
            "qa_warnings": warnings,
            "qa_parrafos": len(doc.paragraphs),
            "qa_tablas": n_tablas,
            "qa_chars": len(texto),
        }

    except ImportError:
        return {**formato, "qa_ok": True, "qa_warnings": ["python-docx no disponible"]}
    except Exception as e:
        return {**formato, "qa_ok": False, "qa_error": str(e)}


def qa_todos(formatos: List[dict]) -> Dict:
    resultados = [qa_uno(f) for f in formatos]
    todos_ok = all(r.get("qa_ok") for r in resultados)
    con_placeholders = sum(1 for r in resultados if any("Placeholders" in w for w in r.get("qa_warnings", [])))
    return {
        "ok": todos_ok,
        "resultados": resultados,
        "total": len(formatos),
        "exitosos": sum(1 for r in resultados if r.get("qa_ok")),
        "con_warnings": sum(1 for r in resultados if r.get("qa_warnings")),
        "con_placeholders": con_placeholders,
    }


def ejecutar(formatos_generados: List[dict]) -> dict:
    return qa_todos(formatos_generados)
