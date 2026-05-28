"""
Generador de documentos clínicos — Sistema por etiquetas (v2)
Sandra Patricia Polania Osorio — REHABILITACION INTEGRAL LABORAL Y OCUPACIONAL SAS

Estrategia: Las plantillas son los documentos de ejemplo. El generador
encuentra celdas/párrafos por su ETIQUETA (texto fijo del formulario oficial)
y reemplaza los valores con los datos del paciente.

v2 — Correcciones:
  - _poner_texto: arreglado bug del elif muerto, ahora maneja runs y
    párrafos correctamente.
  - _buscar_y_reemplazar: ahora detecta la frontera entre celdas de
    etiqueta y celdas de valor en tablas con celdas combinadas.
  - _reemplazar_parrafo: reemplazo seguro de párrafos enteros (busca
    etiqueta en el texto del párrafo y reescribe todo el contenido).
  - _poner_fecha_celdas: escribe fechas en celdas individuales (D1 D2
    / M1 M2 / A1 A2 A3 A4).
  - _marcar_opcion: marca X en la opción correcta entre varias.
"""
import shutil, subprocess, os, re
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _sanitize_filename(name: str) -> str:
    """Reemplaza caracteres invalidos en nombres de archivo (/, backslash, :, *, ?, \", <, >, |)."""
    return re.sub(r'[/\\:*?"<>|]', '-', name)

BASE = Path(os.getenv("RILO_ROOT", str(Path(__file__).parent.parent)))
TEMPLATES = Path(os.getenv("TEMPLATES_DIR", str(BASE / "templates/formatos")))
ASSETS = BASE / "templates/assets"
DOCS = Path(os.getenv("STORAGE_DIR", str(BASE / "storage"))) / "docs"
PDF = BASE / "storage/pdfs"
FIRMA = ASSETS / "firma_sandra.png"
NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

ATRIBUCION = (
    "SANDRA PATRICIA POLANIA OSORIO\n"
    "Fisioterapeuta Esp. SO\n"
    "REHABILITACION INTEGRAL LABORAL Y OCUPACIONAL SAS"
)

def _log(msg: str):
    print(f"[DOCGEN] {msg}", flush=True)

def _normalizar_valor(v, opcional=False):
    """Centraliza la decision de usar [VERIFICAR] para campos requeridos o '' para opcionales."""
    if v is None:
        return "" if opcional else "[VERIFICAR]"
    return str(v)

# ============================================================
# UTILIDADES (corregidas)
# ============================================================

def _poner_texto(celda, texto, bold=None, font_size=None, max_chars=None):
    """Limpia una celda y pone texto nuevo. Maneja runs correctamente.
    Si bold no es None, fuerza bold=True/False en el run.
    Si font_size no es None, configura el tamaño de fuente (en Pt).
    Si max_chars no es None, trunca el texto en el ultimo espacio dentro del limite.
    Elimina párrafos extra del template para evitar whitespace excesivo."""
    from docx.shared import Pt
    texto = str(texto) if texto is not None else ""
    if max_chars and len(texto) > max_chars:
        trunc = texto[:max_chars]
        ultimo_espacio = trunc.rfind(" ")
        if ultimo_espacio > 0:
            trunc = trunc[:ultimo_espacio]
        texto = trunc
    # Limpiar SOLO el sombreado amarillo (FFFF00) de la celda — NO otros colores decorativos
    tcPr = celda._tc.find(f'{{{NS_W}}}tcPr')
    if tcPr is not None:
        shd = tcPr.find(f'{{{NS_W}}}shd')
        if shd is not None:
            fill = shd.get(f'{{{NS_W}}}fill', '')
            val = shd.get(f'{{{NS_W}}}val', '')
            # Solo eliminar si es amarillo (FFFF00) — respetar colores como FCE9D9 (durazno)
            if fill == 'FFFF00' or val == 'FFFF00' or fill == 'yellow':
                tcPr.remove(shd)
    # Limpiar todos los runs Y hyperlinks de todos los párrafos
    for p in celda.paragraphs:
        for run in p.runs:
            run.text = ""
        _eliminar_hyperlinks(p)
    # Eliminar párrafos extra (dejando solo el primero) para evitar celdas con whitespace excesivo.
    # Solo se eliminan <w:p> que son hijos directos de <w:tc>, no tablas anidadas.
    parrafos = celda.paragraphs
    for p_extra in parrafos[1:]:
        parent = p_extra._element.getparent()
        if parent is not None:
            parent.remove(p_extra._element)
    # Poner el nuevo texto en el primer run del primer párrafo
    if celda.paragraphs:
        p = celda.paragraphs[0]
        if p.runs:
            p.runs[0].text = texto
            if bold is not None:
                p.runs[0].bold = bold
            if font_size is not None:
                p.runs[0].font.size = Pt(font_size)
        else:
            run = p.add_run(texto)
            if bold is not None:
                run.bold = bold
            if font_size is not None:
                run.font.size = Pt(font_size)
    else:
        run = celda.add_paragraph().add_run(texto)
        if bold is not None:
            run.bold = bold
        if font_size is not None:
            run.font.size = Pt(font_size)


def _poner_parrafo(parrafo, texto):
    """Limpia un párrafo y pone texto nuevo. Seguro con runs partidos."""
    texto = str(texto) if texto is not None else ""
    for run in parrafo.runs:
        run.text = ""
    if parrafo.runs:
        parrafo.runs[0].text = texto
    else:
        parrafo.add_run(texto)


def _eliminar_hyperlinks(parrafo):
    """Elimina todos los elementos <w:hyperlink> de un párrafo.
    Necesario porque paragraph.runs no incluye los runs dentro de hyperlinks,
    asi que _poner_parrafo() no los limpia, causando texto duplicado."""
    for hl in parrafo._element.findall(f"{{{NS_W}}}hyperlink"):
        parrafo._element.remove(hl)


def _buscar_y_reemplazar(doc, etiqueta, valor, table_indices=None):
    """
    Busca etiqueta en TODAS las tablas (o las especificadas en table_indices)
    y reemplaza el valor en la celda adyacente. Usa _tc para detectar el fin de la región fusionada de la
    etiqueta — evita el bug donde el texto de valor contiene la etiqueta.
    """
    etiqueta_lower = etiqueta.lower().strip()
    valor_str = _normalizar_valor(valor)
    celdas_escritas = set()

    tablas = doc.tables
    if table_indices is not None:
        tablas = [doc.tables[idx] for idx in table_indices if idx < len(doc.tables)]

    for tabla in tablas:
        for fila in tabla.rows:
            celdas = fila.cells
            # 1. Encontrar primera celda que contiene la etiqueta
            label_start = None
            for i, celda in enumerate(celdas):
                if etiqueta_lower in celda.text.lower() and len(celda.text) < 200:
                    label_start = i
                    break
            if label_start is None:
                continue

            # 2. Encontrar el fin de la región fusionada de la etiqueta usando _tc.
            # Esto evita el bug donde el texto del valor contiene la etiqueta word.
            label_tc = celdas[label_start]._tc
            value_start = None
            for i in range(label_start, len(celdas)):
                if celdas[i]._tc != label_tc:
                    value_start = i
                    break
            if value_start is None:
                tc = celdas[label_start]._tc
                if tc not in celdas_escritas:
                    celdas_escritas.add(tc)
                    texto_original = celdas[label_start].text
                    if ":" in texto_original:
                        partes = texto_original.split(":", 1)
                        nuevo_texto = f"{partes[0]}: {valor_str}"
                    else:
                        nuevo_texto = f"{etiqueta}: {valor_str}"
                    _poner_texto(celdas[label_start], nuevo_texto)
                return True

            # 3. Poner el valor en la primera celda de valor (solo si no se ha escrito ya)
            celda_valor = celdas[value_start]
            tc = celda_valor._tc
            if tc not in celdas_escritas:
                celdas_escritas.add(tc)
                _poner_texto(celda_valor, valor_str)
            return True
    return False


def _marcar_opcion(doc, etiqueta, opcion_a_marcar, table_indices=None):
    """
    Marca con X la opción correcta en una fila de checkboxes.
    Identifica los bloques de opciones y sus casillas adyacentes
    (las celdas entre el final de una opción y el inicio de la siguiente).
    Primero limpia TODAS las X en la fila, luego marca solo la correcta.
    """
    etiqueta_lower = etiqueta.lower().strip()
    opcion_lower = opcion_a_marcar.lower().strip()
    marcado = False

    tablas = doc.tables
    if table_indices is not None:
        tablas = [doc.tables[idx] for idx in table_indices if idx < len(doc.tables)]

    for tabla in tablas:
        for fila in tabla.rows:
            celdas = fila.cells
            if not any(etiqueta_lower in c.text.lower() for c in celdas):
                continue

            # Encontrar dónde termina la etiqueta (ej: "Dominancia")
            label_end = 0
            for i, c in enumerate(celdas):
                if etiqueta_lower in c.text.lower():
                    label_end = i + 1

            # Limpiar TODAS las X en la fila primero (no solo las del bloque)
            tcs_todas = set()
            for j in range(label_end, len(celdas)):
                tc = celdas[j]._tc
                if tc not in tcs_todas:
                    tcs_todas.add(tc)
                    if celdas[j].text.strip().upper() == "X":
                        _poner_texto(celdas[j], "")

            # Mapear los bloques de opciones (texto != "" y != "X")
            bloques = []  # [(start, end, text)]
            i = label_end
            while i < len(celdas):
                txt = celdas[i].text.strip()
                if not txt or txt.upper() == "X":
                    i += 1
                    continue
                start = i
                while i < len(celdas) and celdas[i].text.strip() == txt:
                    i += 1
                bloques.append((start, i, txt))

            # Para cada opción, sus casillas van desde su end hasta
            # el start del siguiente bloque. Para el ÚLTIMO bloque,
            # limitar al _tc del primer checkbox para no invadir celdas
            # fusionadas masivamente (bug: vinculación laboral, etc.)
            for bi, (b_start, b_end, b_text) in enumerate(bloques):
                cb_start = b_end
                if bi + 1 < len(bloques):
                    cb_end = bloques[bi + 1][0]
                else:
                    # Último bloque: detectar frontera real del checkbox.
                    # El checkbox son las celdas vacías/X entre el final de la opción
                    # y el inicio de una celda fusionada masiva (diferente _tc).
                    # Tomamos solo celdas que compartan el _tc de cb_start,
                    # o máximo 4 celdas, lo que llegue primero.
                    cb_tc = celdas[b_end]._tc if b_end < len(celdas) else None
                    limit = b_end
                    while limit < len(celdas):
                        if celdas[limit]._tc != cb_tc:
                            # Entramos a nueva región fusionada → parar aquí
                            break
                        limit += 1
                        if limit - b_end >= 4:
                            break
                    cb_end = limit

                # Limpiar X existentes (dedup por _tc)
                tcs_limpias = set()
                for j in range(cb_start, cb_end):
                    tc = celdas[j]._tc
                    if tc not in tcs_limpias:
                        tcs_limpias.add(tc)
                        if celdas[j].text.strip().upper() == "X":
                            _poner_texto(celdas[j], "")

                # Si esta es la opción a marcar, poner X (dedup por _tc).
                # Para el ÚLTIMO bloque, limitar también por cantidad de _tc únicos
                # (máx 4) para no invadir celdas fusionadas masivamente (bug: vinculación).
                # Normalize slash spacing for comparison (template may have "A/B" vs data "A / B")
                import re as _re
                def _norm_slash(s): return _re.sub(r'\s*/\s*', '/', s)
                if (_norm_slash(opcion_lower) in _norm_slash(b_text.lower()) or
                        opcion_lower in b_text.lower()) and not marcado:
                    tcs_marcadas = set()
                    max_tcs = 4 if bi + 1 == len(bloques) else 999
                    for j in range(cb_start, cb_end):
                        tc = celdas[j]._tc
                        if tc not in tcs_marcadas:
                            if len(tcs_marcadas) >= max_tcs:
                                break
                            tcs_marcadas.add(tc)
                            _poner_texto(celdas[j], "X")
                    marcado = True
                    break  # stop processing blocks in this row; continue to next row

            # Continue to next matching row to clear its pre-filled X values
    return marcado


def _poner_fecha_celdas(doc, etiqueta_fecha, fecha_str):
    """
    Escribe una fecha en celdas individuales (D1 D2 / M1 M2 / A1 A2 A3 A4).

    Detecta si la fila encontrada es una FILA DE ENCABEZADO (las celdas
    tienen texto "Día", "mes", "año") o una FILA DE DATOS (celdas vacías o
    con dígitos). Si es encabezado, escribe en la SIGUIENTE FILA a las mismas
    posiciones de columna, saltando celdas fusionadas con dedup por _tc.

    fecha_str: 'YYYY-MM-DD' o 'DD/MM/YYYY' o 'DD-MM-YYYY'
    """
    import re
    fecha_str = str(fecha_str).strip()
    parts = re.split(r'[-/]', fecha_str)
    if len(parts) != 3:
        return False

    if len(parts[0]) == 4:
        y, m, d = parts[0], parts[1], parts[2]
    else:
        d, m, y = parts[0], parts[1], parts[2]

    d = d.zfill(2)
    m = m.zfill(2)
    y = y.zfill(4)

    etiqueta_lower = etiqueta_fecha.lower().strip()

    def _escribir_digitos_dedup(celdas_fila, col_inicio, digitos):
        """Escribe digitos en celdas únicas (_tc) empezando en col_inicio."""
        seen = set()
        di = 0
        for col in range(col_inicio, len(celdas_fila)):
            if di >= len(digitos):
                break
            tc = celdas_fila[col]._tc
            if tc not in seen:
                seen.add(tc)
                _poner_texto(celdas_fila[col], digitos[di])
                di += 1

    for ti, tabla in enumerate(doc.tables):
        for ri, fila in enumerate(tabla.rows):
            celdas = fila.cells
            if not any(etiqueta_lower in c.text.lower() for c in celdas):
                continue

            dia_col = mes_col = ano_col = None
            for i, celda in enumerate(celdas):
                txt = celda.text.lower().strip()
                if txt in ('día', 'dia') and dia_col is None:
                    dia_col = i
                elif txt == 'mes' and mes_col is None:
                    mes_col = i
                elif txt in ('año', 'ano') and ano_col is None:
                    ano_col = i

            if dia_col is None and mes_col is None and ano_col is None:
                continue

            # ¿Es fila encabezado? (los textos SON las palabras de label)
            es_encabezado = (
                dia_col is not None and
                celdas[dia_col].text.lower().strip() in ('día', 'dia')
            )

            if es_encabezado:
                # Buscar la fila de datos: la siguiente fila del mismo bloque
                # que también contenga la etiqueta (celda fusionada verticalmente)
                fila_datos = None
                for ri2 in range(ri + 1, len(tabla.rows)):
                    celdas2 = tabla.rows[ri2].cells
                    if any(etiqueta_lower in c.text.lower() for c in celdas2):
                        fila_datos = celdas2
                        break
                if fila_datos is None:
                    continue
                target = fila_datos
            else:
                target = celdas

            if dia_col is not None:
                _escribir_digitos_dedup(target, dia_col, d)
            if mes_col is not None:
                _escribir_digitos_dedup(target, mes_col, m)
            if ano_col is not None:
                _escribir_digitos_dedup(target, ano_col, y)
            return True
    return False


def _reemplazar_parrafo(doc, texto_busqueda, texto_nuevo):
    """
    Busca un párrafo que CONTENGA texto_busqueda y reemplaza TODO
    su contenido con texto_nuevo. Seguro: limpia todos los runs primero.
    También elimina hyperlinks heredados de la plantilla para evitar
    duplicación de texto (los hyperlinks no son accesibles vía paragraph.runs).
    """
    texto_nuevo = str(texto_nuevo) if texto_nuevo is not None else ""
    for p in doc.paragraphs:
        if texto_busqueda in p.text:
            _poner_parrafo(p, texto_nuevo)
            _eliminar_hyperlinks(p)
            return True
    return False


def _reemplazar_seccion_parrafos(doc, inicio, fin, texto_nuevo):
    """
    Reemplaza todos los párrafos entre 'inicio' (exclusive) y 'fin'
    con el contenido de texto_nuevo. Útil para secciones de recomendaciones
    en cartas donde el contenido varía completamente.

    Cada línea de texto_nuevo se convierte en un párrafo nuevo.
    Usa paragraph.clear() para eliminar formato previo (evita guión doble
    • - texto) y NO incluye guiones manuales si la plantilla ya usa bullets.
    Elimina párrafos sobrantes para no dejar viñetas vacías.
    """
    lineas_raw = texto_nuevo.strip().split('\n')
    # Limpiar guiones manuales: el bullet lo pone el estilo del párrafo
    lineas = []
    for linea in lineas_raw:
        linea = re.sub(r'^[\s]*[-•*]\s*', '', linea)
        lineas.append(linea)

    parrafos = doc.paragraphs

    # Encontrar índices de inicio y fin
    idx_inicio = None
    idx_fin = None
    for i, p in enumerate(parrafos):
        if idx_inicio is None and inicio in p.text:
            idx_inicio = i
        elif idx_inicio is not None and fin in p.text:
            idx_fin = i
            break

    if idx_inicio is None:
        return False

    # Si no hay fin, reemplazar hasta el final
    if idx_fin is None:
        idx_fin = len(parrafos)

    # Limpiar párrafos entre inicio+1 y fin-1
    for i in range(idx_inicio + 1, idx_fin):
        _poner_parrafo(parrafos[i], "")

    # Insertar nuevo contenido
    rango = idx_fin - idx_inicio - 1
    # Detectar numPr del primer párrafo con bullet en la sección para heredarlo
    numPr_ref = None
    for i in range(idx_inicio + 1, idx_fin):
        pPr = parrafos[i]._element.find(f"{{{NS_W}}}pPr")
        if pPr is not None:
            numPr = pPr.find(f"{{{NS_W}}}numPr")
            if numPr is not None:
                numPr_ref = numPr
                break
    for li, linea in enumerate(lineas):
        if li < rango:
            p = parrafos[idx_inicio + 1 + li]
            # Verificar si este párrafo ya tiene bullet formatting
            pPr = p._element.find(f"{{{NS_W}}}pPr")
            tiene_numPr = pPr is not None and pPr.find(f"{{{NS_W}}}numPr") is not None
            p.clear()
            if tiene_numPr:
                # El bullet lo pone numPr, no incluir guion manual
                p.add_run(linea)
            elif numPr_ref is not None:
                # Heredar bullet del párrafo hermano
                if pPr is None:
                    from lxml import etree
                    pPr = etree.SubElement(p._element, f"{{{NS_W}}}pPr")
                    p._element.insert(0, pPr)
                pPr.append(deepcopy(numPr_ref))
                p.add_run(linea)
            else:
                # Sin bullet disponible, incluir guion manual
                p.add_run(f"- {linea}")

    # Eliminar párrafos sobrantes (viñetas vacías de la plantilla)
    # Iterar en reversa para que los índices no se desplacen al borrar
    sobrantes = rango - len(lineas)
    if sobrantes > 0:
        body = parrafos[idx_inicio]._element.getparent()
        for i in range(idx_fin - 1, idx_inicio + len(lineas), -1):
            body.remove(parrafos[i]._element)

    return True


def _reemplazar_en_parrafos(doc, viejo, nuevo):
    """
    Reemplaza 'viejo' por 'nuevo' dentro de los párrafos.
    Versión segura: reconstruye cada párrafo afectado.
    No hace nada si viejo está vacío.
    También elimina hyperlinks heredados para evitar duplicación.
    """
    if not viejo:
        return
    nuevo = _normalizar_valor(nuevo, opcional=True)
    for p in doc.paragraphs:
        if viejo in p.text:
            new_text = p.text.replace(viejo, nuevo)
            _poner_parrafo(p, new_text)
            _eliminar_hyperlinks(p)


def _reemplazar_en_tablas(doc, viejo, nuevo, table_indices=None):
    """Reemplaza 'viejo' por 'nuevo' dentro de las celdas de tablas (o las especificadas).
    Usa cell._tc para deduplicar celdas fusionadas (merged cells)."""
    if not viejo:
        return
    nuevo = str(nuevo) if nuevo is not None else ""
    procesadas = set()
    tablas = doc.tables
    if table_indices is not None:
        tablas = [doc.tables[idx] for idx in table_indices if idx < len(doc.tables)]
    for tabla in tablas:
        for fila in tabla.rows:
            for celda in fila.cells:
                tc = celda._tc
                if tc in procesadas:
                    continue
                procesadas.add(tc)
                if viejo in celda.text:
                    new_text = celda.text.replace(viejo, nuevo)
                    _poner_texto(celda, new_text)


# Nombres femeninos comunes en Colombia
_NOMBRES_FEMENINOS = {
    "maria", "rosa", "maribel", "laura", "claudia", "patricia", "sandra",
    "diana", "carolina", "andrea", "alejandra", "luz", "gloria", "martha",
    "marta", "pilar", "natalia", "paola", "viviana", "marcela", "adriana",
    "liliana", "yolanda", "beatriz", "cecilia", "helena", "irene", "julia",
    "blanca", "esperanza", "graciela", "lucía", "lucia", "silvia", "teresa",
    "susana", "mónica", "monica", "margarita", "dolores", "consuelo",
}

def _detectar_genero(nombre: str) -> str:
    """
    Retorna 'F' si el primer nombre es femenino, 'M' si masculino, '' si no se sabe.
    Detecta por lista de nombres comunes colombianos.
    """
    if not nombre:
        return ""
    primer_nombre = nombre.strip().split()[0].lower()
    if primer_nombre in _NOMBRES_FEMENINOS:
        return "F"
    # Heurística: nombres que terminan en 'a' son frecuentemente femeninos
    if primer_nombre.endswith("a") and len(primer_nombre) > 3:
        return "F"
    return "M"


def _reemplazar_celda_entera(doc, texto_busqueda, texto_nuevo, bold=None, font_size=None, table_indices=None):
    """
    Busca una celda que CONTENGA texto_busqueda (robusto y case-insensitive)
    y reemplaza TODO su contenido con texto_nuevo. Usa cell._tc para deduplicar.
    Ideal para celdas con texto largo (motivo consulta, concepto, etc.)
    donde _reemplazar_en_tablas causaría duplicación.
    Si bold no es None, fuerza el formato bold del texto.
    Si font_size no es None, configura el tamaño de fuente.
    """
    if not texto_busqueda:
        return False
    
    def _norm(t):
        import re
        import unicodedata
        t = re.sub(r'[\s\-]+', ' ', t.lower().strip())
        t = unicodedata.normalize("NFD", t)
        t = "".join(c for c in t if unicodedata.category(c) != "Mn")
        return t

    busqueda_norm = _norm(texto_busqueda)
    procesadas = set()
    
    tablas = doc.tables
    if table_indices is not None:
        tablas = [doc.tables[idx] for idx in table_indices if idx < len(doc.tables)]

    for tabla in tablas:
        for fila in tabla.rows:
            for celda in fila.cells:
                tc = celda._tc
                if tc in procesadas:
                    continue
                procesadas.add(tc)
                if busqueda_norm in _norm(celda.text):
                    _poner_texto(celda, texto_nuevo, bold=bold, font_size=font_size)
                    return True
    return False


def _insertar_firma(doc, etiqueta="Insertar Firma"):
    """Busca 'Insertar Firma' en tablas y reemplaza con la imagen.
    Usa cell._tc para deduplicar celdas fusionadas."""
    if not FIRMA.exists():
        return False
    procesadas = set()
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                tc = celda._tc
                if tc in procesadas:
                    continue
                procesadas.add(tc)
                if etiqueta in celda.text:
                    for p in celda.paragraphs:
                        p.clear()
                    run = celda.paragraphs[0].add_run()
                    run.add_picture(str(FIRMA), width=Inches(1.3))
                    return True
    return False


def _convertir_a_pdf(docx_path):
    """Convierte DOCX a PDF con LibreOffice headless.
    Lanza RuntimeError si falla la conversion o LibreOffice no esta disponible."""
    soffice = None
    for candidate in ["/usr/lib/libreoffice/program/soffice", "libreoffice", "soffice"]:
        try:
            subprocess.run([candidate, "--version"], capture_output=True, timeout=5)
            soffice = candidate
            break
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    if not soffice:
        raise RuntimeError(f"LibreOffice no encontrado. No se pudo convertir a PDF: {docx_path}")

    try:
        result = subprocess.run([
            soffice, "--headless", "--convert-to", "pdf",
            docx_path, "--outdir", str(PDF)
        ], check=True, timeout=120, capture_output=True, text=True)
        pdf_name = Path(docx_path).stem + ".pdf"
        pdf_path = PDF / pdf_name
        if pdf_path.exists():
            return str(pdf_path)
        else:
            raise RuntimeError(f"PDF esperado en {pdf_path} pero no se encontro. stderr: {result.stderr[:200]}")
    except subprocess.TimeoutExpired:
        raise RuntimeError(f"Timeout convirtiendo {docx_path} (120s)")
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"Error LibreOffice convirtiendo {docx_path}: {e.stderr[:300] if e.stderr else str(e)}")
    except RuntimeError:
        raise
    except Exception as e:
        raise RuntimeError(f"Error inesperado convirtiendo {docx_path}: {e}")


def _preparar_documento(src_name, output_name, datos):
    """Copia la plantilla y devuelve (doc, dst_path)."""
    src = TEMPLATES / src_name
    if not src.exists():
        raise FileNotFoundError(f"Plantilla no encontrada: {src}")
    output_name = _sanitize_filename(output_name)
    dst = DOCS / f"{output_name}.docx"
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(str(src), str(dst))
    doc = Document(str(dst))
    return doc, dst


def _safe_get(d, *keys, default="[VERIFICAR]"):
    """Acceso seguro a diccionarios anidados."""
    for k in keys:
        if isinstance(d, dict):
            d = d.get(k, default)
        else:
            return default
    return d if d is not None else default


# ============================================================
# GENERADORES POR FORMATO
# ============================================================


def _marcar_opcion_en_tabla(tabla, opcion):
    """Marca X en la celda checkbox de una tabla de conclusiones.
    Busca la fila que contiene 'Conclusión:' y marca la opción correcta."""
    opcion_lower = opcion.lower().strip()
    for fila in tabla.rows:
        for ci, celda in enumerate(fila.cells):
            if "conclusión" in celda.text.lower() or "conclusion" in celda.text.lower():
                # La fila de conclusión tiene checkbox en las celdas siguientes
                for cj in range(ci + 1, len(fila.cells)):
                    txt = fila.cells[cj].text.strip()
                    if opcion_lower in txt.lower():
                        # Limpiar X previas en celdas checkbox cercanas
                        for ck in range(ci + 1, min(cj + 3, len(fila.cells))):
                            if fila.cells[ck].text.strip().upper() == "X":
                                _poner_texto(fila.cells[ck], "")
                        # Marcar X en la celda ANTERIOR a la opción (es el checkbox)
                        if cj > ci:
                            _poner_texto(fila.cells[cj - 1], "X")
                        return True
    return False


def _llenar_perfil_exigencias(tabla, perfil):
    """Llena el Anexo Perfil de Exigencias (Tabla 6, 96 filas).
    Marca X en la columna de calificación (0-4) según el score de cada ítem.
    Los scores vienen como int (0-4) o tuple (score, observacion)."""
    # Mapa de secciones a filas en la tabla
    secciones = [
        ("vision", "VISION"),
        ("audicion", "AUDICIÓN"),
        ("sensibilidad", "SENSIBILIDAD SUPERFICIAL Y PROFUNDA"),
        ("olfato_gusto", "OLFATO / GUSTO"),
        ("motricidad_gruesa", "MOTRICIDAD GRUESA"),
        ("motricidad_fina", "MOTRICIDAD FINA"),
        ("armonia", "ARMONÍA"),
        ("cognitivos", "COGNITIVOS"),
        ("sociales_laborales", "Adaptación al Grupo de Trabajo"),  # sub-sección de COGNITIVOS
        ("laborales", "LABORALES"),
    ]
    
    perfil_dict = perfil if isinstance(perfil, dict) else {}
    # Flatten: iterate through all rows of tabla 6, find items by name
    for ri, fila in enumerate(tabla.rows):
        item_name = fila.cells[0].text.strip()
        if not item_name:
            continue

        # Deduplicar celdas fusionadas antes de indexar por columna
        celdas_unicas = []
        seen_tcs = set()
        for c in fila.cells:
            tc = c._tc
            if tc not in seen_tcs:
                seen_tcs.add(tc)
                celdas_unicas.append(c)

        # Always clear pre-filled x in score columns (1-6), regardless of data
        for sc in range(1, 7):
            if sc < len(celdas_unicas) and celdas_unicas[sc].text.strip().upper() == "X":
                _poner_texto(celdas_unicas[sc], "")

        # Search for this item in the perfil data
        score = None
        obs = ""
        for section_name, section_items in perfil_dict.items():
            if isinstance(section_items, dict) and item_name in section_items:
                val = section_items[item_name]
                if isinstance(val, tuple):
                    score, obs = val
                else:
                    score = val
                break

        if score is not None:
            # Dedup layout: [0]=name, [1]=score-0, [2]=score-1, [3]=score-2,
            #               [4]=score-3, [5]=score-4, [6]=Observaciones
            score_col = 1 + int(score)
            if score_col < len(celdas_unicas):
                _poner_texto(celdas_unicas[score_col], "X")

            # Observación en columna [6]
            if obs and 6 < len(celdas_unicas):
                _poner_texto(celdas_unicas[6], obs)


def _poner_texto_multiparrafo(celda, texto, bold=None, font_size=None):
    """Escribe texto en una celda respetando multiples parrafos. Separa por saltos de linea (\\n o \\n\\n)."""
    from docx.shared import Pt
    partes = [p for p in texto.replace("\n\n", "\n").split("\n") if p.strip()]
    if not partes:
        partes = [texto]
    
    # Limpiar shading amarillo
    tcPr = celda._tc.find(f'{{{NS_W}}}tcPr')
    if tcPr is not None:
        shd = tcPr.find(f'{{{NS_W}}}shd')
        if shd is not None:
            fill = shd.get(f'{{{NS_W}}}fill', '')
            if fill == 'FFFF00':
                tcPr.remove(shd)
    
    # Limpiar todos los parrafos existentes
    for p in celda.paragraphs:
        for run in p.runs:
            run.text = ""
        _eliminar_hyperlinks(p)

    # Eliminar párrafos extra que excedan el número de partes
    parrafos_actuales = celda.paragraphs
    if len(parrafos_actuales) > len(partes):
        for p_extra in parrafos_actuales[len(partes):]:
            parent = p_extra._element.getparent()
            if parent is not None:
                parent.remove(p_extra._element)

    # Escribir cada parte en un parrafo
    for idx, parte in enumerate(partes):
        if idx < len(celda.paragraphs):
            p = celda.paragraphs[idx]
            if p.runs:
                p.runs[0].text = parte
                if bold is not None:
                    p.runs[0].bold = bold
                if font_size is not None:
                    p.runs[0].font.size = Pt(font_size)
            else:
                run = p.add_run(parte)
                if bold is not None:
                    run.bold = bold
                if font_size is not None:
                    run.font.size = Pt(font_size)
        else:
            run = celda.add_paragraph().add_run(parte)
            if bold is not None:
                run.bold = bold
            if font_size is not None:
                run.font.size = Pt(font_size)

def generar_analisis_exigencia(datos, output_name=None):
    """FORMATO 1 — Análisis de Exigencias / Homologación.
    
    Estructura del template:
      Tabla 0 (32f×22c): Encabezado + Sección 1 (Identificación)
      Tabla 1 (33f×5c): Contacto empresa + Secciones 2-5 (Metodología, Antecedentes, 
                         Condiciones de trabajo, Tarea 1 datos)
      Tabla 2 (10f×8c): Tarea 1 — Apreciaciones y Conclusión
      Tabla 3 (4f×4c):  Tarea 2 — Datos
      Tabla 4 (10f×8c): Tarea 2 — Apreciaciones y Conclusión
      Tabla 5 (25f×6c): Secciones 6-9 (Materiales, Peligros, Concepto, Recomendaciones)
      Tabla 6 (96f×10c): Anexo Perfil de Exigencias + Sección 10 (Registro)
    """
    p, e, l, s, c = _extraer_secciones(datos)
    output_name = output_name or f"analisis-{p.get('documento','sinid')}-{c.get('fecha','sinfecha')}"
    doc, dst = _preparar_documento("ejemplo analisis de exigencia.docx", output_name, datos)

    # ═══ TABLA 0 — Encabezado + Sección 1: Identificación ═══
    # Fecha en celdas individuales (F00-F01)
    if c.get("fecha"):
        _poner_fecha_celdas(doc, "FECHA", c["fecha"])

    # Campos label→value (substrings del template)
    _buscar_y_reemplazar(doc, "Nombre del trabajador", p.get("nombre"))
    _buscar_y_reemplazar(doc, "Número de documento", p.get("documento"))
    _buscar_y_reemplazar(doc, "Identificación del siniestro", s.get("id_siniestro"))

    # Fecha de nacimiento: _poner_fecha_celdas detecta fila encabezado y escribe
    # en la fila de datos (T0 F9) automáticamente.
    if p.get("fecha_nacimiento"):
        _poner_fecha_celdas(doc, "Fecha de nacimiento/edad", p["fecha_nacimiento"])
    # Edad: T0 F9 C17 (celda de valor, no el encabezado "Edad" que está en F8 C12-C21)
    edad_val = p.get("edad", "")
    if edad_val:
        try:
            _poner_texto(doc.tables[0].rows[9].cells[17], str(edad_val))
        except (IndexError, AttributeError) as e:
            _log(f"Celda no encontrada ({e}) - campo 'edad' omitido en analisis_exigencia")

    # Dominancia — checkbox
    dom = p.get("dominancia", "")
    if dom:
        _marcar_opcion(doc, "Dominancia", dom)

    _buscar_y_reemplazar(doc, "Estado civil", p.get("estado_civil"))

    # Nivel educativo — checkbox
    nivel = p.get("nivel_educativo", "")
    if nivel:
        _marcar_opcion(doc, "Nivel educativo", nivel)
    # Formación "Otros"
    otros = p.get("formacion_otros", "")
    if otros:
        _buscar_y_reemplazar(doc, "Otros", otros)

    _buscar_y_reemplazar(doc, "Teléfonos trabajador", p.get("telefono"))
    _buscar_y_reemplazar(doc, "Dirección residencia", p.get("direccion"))
    diag_texto = s.get("diagnosticos", "")
    if diag_texto:
        diag_texto = diag_texto.replace(" / ", chr(10))
        # Tabla 0 F18: C0 es label (gs=2, ocupa C0-C1), C2 es valor
        _reemplazar_celda_entera(doc, "ESGUINCE GRADO I", diag_texto, bold=False, font_size=7.5)
    _buscar_y_reemplazar(doc, "Fecha(s) del evento(s) ATEL", s.get("fecha_evento"))
    _buscar_y_reemplazar(doc, "EPS - IPS", p.get("eps_ips"))
    _buscar_y_reemplazar(doc, "AFP", p.get("afp"))
    _buscar_y_reemplazar(doc, "Tiempo total de incapacidad", s.get("tiempo_incapacidad"))
    _buscar_y_reemplazar(doc, "Empresa donde labora", e.get("nombre"))
    _buscar_y_reemplazar(doc, "NIT de la Empresa", e.get("nit"))
    _buscar_y_reemplazar(doc, "Cargo actual", l.get("cargo"))
    _buscar_y_reemplazar(doc, "Área/sección/proceso", l.get("area"))

    # Fecha ingreso cargo: _poner_fecha_celdas detecta encabezado y escribe en fila datos
    if l.get("fecha_ingreso_cargo"):
        _poner_fecha_celdas(doc, "Fecha ingreso cargo", l["fecha_ingreso_cargo"])
    # Antigüedad en el cargo: T0 F28 C12
    ant_cargo = l.get("antiguedad_cargo", "")
    if ant_cargo:
        try:
            _poner_texto(doc.tables[0].rows[28].cells[12], ant_cargo)
        except (IndexError, AttributeError) as e:
            _log(f"Celda no encontrada ({e}) - campo 'antiguedad_cargo' omitido en analisis_exigencia")

    # Fecha ingreso empresa: mismo patrón
    if l.get("fecha_ingreso_empresa"):
        _poner_fecha_celdas(doc, "Fecha ingreso a la empresa", l["fecha_ingreso_empresa"])
    # Antigüedad en la empresa: T0 F30 C12
    ant_emp = l.get("antiguedad_empresa", "")
    if ant_emp:
        try:
            _poner_texto(doc.tables[0].rows[30].cells[12], ant_emp)
        except (IndexError, AttributeError) as e:
            _log(f"Celda no encontrada ({e}) - campo 'antiguedad_empresa' omitido en analisis_exigencia")

    _buscar_y_reemplazar(doc, "Forma de vinculación laboral", l.get("vinculacion"))

    # ═══ TABLA 1 — Contacto empresa + Secciones 2-5 ═══
    tabla1 = doc.tables[1]
    
    # F00-F03: Contacto, Correo, Teléfono, Dirección empresa
    # Contacto: poner los dos contactos en líneas separadas como en el ejemplo
    contacto_val = e.get("contacto", "")
    if contacto_val:
        # Separar por " / " entre el nombre del primer contacto y el segundo
        contacto_val = contacto_val.replace(" / Franky", chr(10) + "Franky")  # chr(10) = salto de linea
        _buscar_y_reemplazar(doc, "Contacto en empresa", contacto_val)
    _buscar_y_reemplazar(doc, "Correo(s) electrónico(s)", e.get("correo"))
    _buscar_y_reemplazar(doc, "Teléfonos de contacto empresa", e.get("telefono"))
    _buscar_y_reemplazar(doc, "Dirección de empresa", e.get("direccion"))

    # Sección 2 — Metodología (F06) — preservar estructura de párrafos del template
    metodo = c.get("metodologia", "")
    if metodo:
        # Buscar la celda de metodología en Tabla 1 fila 6
        for fila in doc.tables[1].rows:
            if "La metodología utilizada" in fila.cells[0].text:
                _poner_texto_multiparrafo(fila.cells[0], metodo, bold=False, font_size=7.5)
                break

    # Sección 3 — Antecedentes del caso (F09-F10)
    resultados = c.get("resultados_valoracion", "")
    if resultados:
        _reemplazar_celda_entera(doc, "Servidora de", resultados, bold=False, font_size=7.5)
    
    rhi = c.get("programa_rhi", "")
    if rhi:
        # RHI: reemplazar " / " por saltos de línea como en el ejemplo
        rhi = rhi.replace(" / ", chr(10))  # chr(10) = salto de linea
        # Escribir directo a Tabla 1 Fila 10 Celda 2 (celda de valor de RHI)
        _poner_texto(doc.tables[1].rows[10].cells[2], rhi, bold=False, font_size=7.5)

    # Sección 4 — Condiciones de trabajo
    # 4.1 Descripción del proceso productivo (F14)
    pp = c.get("proceso_productivo", "")
    if pp:
        _reemplazar_celda_entera(doc, "La Servidora manifiesta las siguientes funciones", pp, bold=False, font_size=7.5)

    # 4.2 Apreciación del trabajador (F17)
    at = c.get("apreciacion_trabajador", "")
    if at:
        _reemplazar_celda_entera(doc, "La servidora menciona que ejecuta las tareas", at, bold=False, font_size=7.5)

    # 4.3 Estándares de productividad (F20)
    ep = c.get("estandares_productividad", "")
    if ep:
        # The template row F20 has "El ritmo de trabajo se desarrolla a diario..."
        _reemplazar_celda_entera(doc, "El ritmo de trabajo se desarrolla", ep, bold=False, font_size=7.5)

    # 4.4 Requerimientos del desempeño organizacional (F23-F26)
    _buscar_y_reemplazar(doc, "Jornada", l.get("jornada"))
    _buscar_y_reemplazar(doc, "Ritmo", l.get("ritmo"))
    _buscar_y_reemplazar(doc, "Descansos programados", l.get("descansos"))
    _buscar_y_reemplazar(doc, "Turnos", l.get("turnos"))
    _buscar_y_reemplazar(doc, "Tiempos efectivos", l.get("tiempos_efectivos"))
    _buscar_y_reemplazar(doc, "Rotaciones", l.get("rotaciones"))
    _buscar_y_reemplazar(doc, "Horas extras", l.get("horas_extras"))
    _buscar_y_reemplazar(doc, "Distribución semanal", l.get("distribucion_semanal"))

    # ═══ SECCIÓN 5 — Tareas críticas ═══
    _tareas_raw = datos.get("tareas_criticas") or datos.get("laboral", {}).get("tareas_criticas", [])
    tareas = [t for t in (_tareas_raw if isinstance(_tareas_raw, list) else []) if isinstance(t, dict)]

    # Pre-clean conclusion checkboxes in T2 and T4 (template pre-fills "x" for first option)
    for ti_conc in [2, 4]:
        if ti_conc < len(doc.tables):
            for fila_c in doc.tables[ti_conc].rows:
                celdas_c = fila_c.cells
                row_txt = " ".join(cel.text.strip().lower() for cel in celdas_c)
                if "puede desempeñarla" in row_txt or "restricción parcial" in row_txt:
                    seen_c = set()
                    for cel in celdas_c:
                        if cel._tc not in seen_c:
                            seen_c.add(cel._tc)
                            if cel.text.strip().upper() == "X":
                                _poner_texto(cel, "")

    # Tarea 1: Datos en Tabla 1 F28-F32, Apreciaciones en Tabla 2
    if len(tareas) >= 1:
        t1 = tareas[0]
        # Tabla 1: Actividad, Ciclo, Subactividad, Estándar, Descripción
        _buscar_y_reemplazar(doc, "Actividad", t1.get("actividad"))  # first occurrence = Tarea 1
        # The template has Actividad at F29[0], value at F29[2]
        _buscar_y_reemplazar(doc, "Ciclo", t1.get("ciclo"))
        _buscar_y_reemplazar(doc, "Subactividad", t1.get("subactividad"))
        _buscar_y_reemplazar(doc, "Estándar de Productividad", t1.get("estandar"))
        desc1 = t1.get("descripcion", "")
        if desc1:
            _reemplazar_celda_entera(doc, "La Servidora realiza las siguientes tareas con uso de computador", desc1, bold=False, font_size=7.5)

        # Tabla 2: Apreciaciones y Conclusión Tarea 1
        t2 = doc.tables[2]
        ap_trab1 = t1.get("apreciacion_trabajador", "")
        if ap_trab1:
            _reemplazar_celda_entera(doc, "La servidora indica que realiza las tareas administrativas", ap_trab1, bold=False, font_size=7.5)
        ap_prof1 = t1.get("apreciacion_profesional", "")
        if ap_prof1:
            _reemplazar_celda_entera(doc, "Servidora sin limitación en las funciones", ap_prof1, bold=False, font_size=7.5)
        # Conclusión checkbox + descripción en Tabla 2
        conc1_tipo = t1.get("conclusion_tipo", "")
        if conc1_tipo:
            _marcar_opcion_en_tabla(t2, conc1_tipo)
        conc1_desc = t1.get("conclusion_descripcion", "")
        if conc1_desc:
            _reemplazar_celda_entera(doc, "La actividad puede ser realizada por la trabajadora", conc1_desc, bold=False, font_size=7.5)

    # Tarea 2: Datos en Tabla 3, Apreciaciones en Tabla 4
    if len(tareas) >= 2:
        t2_data = tareas[1]
        # Tabla 3: Actividad, Ciclo, Subactividad, Estándar, Descripción
        t3 = doc.tables[3]
        if t3.rows[0].cells[1].text.strip():
            _poner_texto(t3.rows[0].cells[1], t2_data.get("actividad", ""))
        _poner_texto(t3.rows[0].cells[3], t2_data.get("ciclo", ""))
        _poner_texto(t3.rows[1].cells[1], t2_data.get("subactividad", ""))
        _poner_texto(t3.rows[1].cells[3], t2_data.get("estandar", ""))
        desc2 = t2_data.get("descripcion", "")
        if desc2:
            _poner_texto(t3.rows[3].cells[1], desc2)

        # Tabla 4: Apreciaciones y Conclusión Tarea 2
        t4 = doc.tables[4]
        ap_trab2 = t2_data.get("apreciacion_trabajador", "")
        if ap_trab2:
            _reemplazar_celda_entera(doc, "La trabajadora manifiesta limitación moderada", ap_trab2, bold=False, font_size=7.5)
        ap_prof2 = t2_data.get("apreciacion_profesional", "")
        if ap_prof2:
            _reemplazar_celda_entera(doc, "La servidora puede realizar la actividad teniendo", ap_prof2, bold=False, font_size=7.5)
        conc2_tipo = t2_data.get("conclusion_tipo", "")
        if conc2_tipo:
            _marcar_opcion_en_tabla(t4, conc2_tipo)
        conc2_desc = t2_data.get("conclusion_descripcion", "")
        if conc2_desc:
            _reemplazar_celda_entera(doc, "La actividad puede ser realizada por la trabajadora", conc2_desc, bold=False, font_size=7.5)

    # ═══ TABLA 5 — Secciones 6-9 ═══
    t5 = doc.tables[5]

    # Sección 6 — Materiales (F02-F06)
    _mat_raw = datos.get("materiales") or datos.get("laboral", {}).get("materiales", [])
    materiales = [m for m in (_mat_raw if isinstance(_mat_raw, list) else []) if isinstance(m, dict)]
    for idx, mat in enumerate(materiales):
        if idx >= 5:
            break
        fila_mat = t5.rows[2 + idx]  # F02-F06
        _poner_texto(fila_mat.cells[0], mat.get("nombre", ""))
        _poner_texto(fila_mat.cells[2], mat.get("estado", ""))
        _poner_texto(fila_mat.cells[3], mat.get("requerimientos", ""))

    # Sección 7 — Peligros (F10-F15)
    _pel_raw = datos.get("peligros") or datos.get("laboral", {}).get("peligros", [])
    peligros = [p for p in (_pel_raw if isinstance(_pel_raw, list) else []) if isinstance(p, dict)]
    for idx, pel in enumerate(peligros):
        if idx >= 6:
            break
        fila_pel = t5.rows[10 + idx]  # F10-F15
        _poner_texto(fila_pel.cells[0], pel.get("nombre", ""))
        _poner_texto(fila_pel.cells[2], pel.get("descripcion", ""))
        _poner_texto(fila_pel.cells[4], pel.get("control", ""))
        # Recomendación at cell[5] if exists
        if len(fila_pel.cells) > 5:
            _poner_texto(fila_pel.cells[5], pel.get("recomendacion", ""))

    # Sección 8 — Concepto capacidad de desempeño (F18)
    concepto = c.get("concepto_desempeno", "")
    if concepto:
        _reemplazar_celda_entera(doc, "La servidora se desempeña en el cargo", concepto, bold=False, font_size=7.5)

    # Sección 9 — Recomendaciones (F22, F24)
    rec = datos.get("recomendaciones", {})
    rec_trab = rec.get("trabajador", "")
    if rec_trab:
        _reemplazar_celda_entera(doc, "Las recomendaciones relacionadas con el presente análisis", rec_trab, bold=False, font_size=7.5)
    rec_emp = rec.get("empresa", "")
    if rec_emp and rec_emp != rec_trab:
        # Same text appears in two cells; the second will match too
        pass  # First match already handled

    # ═══ TABLA 6 — Anexo Perfil de Exigencias + Sección 10 Registro ═══
    t6 = doc.tables[6]
    perfil = datos.get("perfil_exigencias", {})
    _llenar_perfil_exigencias(t6, perfil)

    # Sección 10 — Registro (F91-F95)
    registro = datos.get("registro", {})
    if registro.get("elaboro"):
        _reemplazar_celda_entera(doc, "BIBIANA HORTA PERDOMO", registro["elaboro"], bold=False, font_size=7.5)
    if registro.get("reviso"):
        _reemplazar_celda_entera(doc, "JENNY MARITZA RIVERA POLANIA", registro["reviso"], bold=False, font_size=7.5)

    # Firma
    _insertar_firma(doc)

    doc.save(str(dst))
    _convertir_a_pdf(str(dst))
    return str(dst)


def generar_carta_medidas(datos, output_name=None):
    """FORMATO 2 — Carta de Medidas Preventivas"""
    p, e, l, s, c, v, r, rec, *_ = _extraer_secciones(datos, extra=["visita", "rehabilitacion", "recomendaciones"])
    output_name = output_name or f"medidas-{p.get('documento','sinid')}-{c.get('fecha','sinfecha')}"
    doc, dst = _preparar_documento("ejemplo carta de medidas.docx", output_name, datos)

    # --- Encabezado ---
    _reemplazar_en_parrafos(doc, "FISCALÍA GENERAL DE LA NACIÓN HUILA", e.get("nombre") or "[VERIFICAR]")
    _reemplazar_en_parrafos(doc, "Dra. Diana Cristina Sánchez Pama", e.get("contacto") or "[VERIFICAR]")
    _reemplazar_en_parrafos(doc, "Profesional de seguridad y salud en el trabajo", e.get("cargo_contacto") or "[VERIFICAR]")
    _reemplazar_en_parrafos(doc, "Carrera 21ª # 26-65 Sur", e.get("direccion") or "[VERIFICAR]")
    _reemplazar_en_parrafos(doc, "3223089360", e.get("telefono") or "[VERIFICAR]")
    _reemplazar_en_parrafos(doc, "dianac.sanchez@fiscalia.gov.co", e.get("correo") or "[VERIFICAR]")

    # --- Asunto ---
    _reemplazar_en_parrafos(doc, "OLGA LUCIA PAREDES ORTIZ", p.get("nombre", "[NOMBRE]"))
    _reemplazar_en_parrafos(doc, "40.776.172", p.get("documento", "[DOCUMENTO]"))
    _reemplazar_en_parrafos(doc, "423022266", s.get("id_siniestro", "[SINIESTRO]"))
    _reemplazar_en_parrafos(doc, "12/08/20", s.get("fecha_evento", "[FECHA EVENTO]"))

    # --- Párrafo inicial (cuerpo) ---
    cuerpo = c.get("cuerpo_carta", "")
    if cuerpo:
        _reemplazar_parrafo(doc, "Una vez realizada la visita al puesto de trabajo", cuerpo)

    # --- PARA LA SERVIDORA ---
    texto_trabajador = rec.get("trabajador_texto", "")
    if texto_trabajador:
        _reemplazar_seccion_parrafos(doc, "PARA LA SERVIDORA:", "RECOMENDACIONES PARA EL DESARROLLO", texto_trabajador)

    # --- RECOMENDACIONES PARA LAS TAREAS ---
    texto_tareas = rec.get("tareas_texto", "")
    if texto_tareas:
        _reemplazar_seccion_parrafos(doc, "RECOMENDACIONES PARA EL DESARROLLO", "PARA LA ENTIDAD", texto_tareas)

    # --- PARA LA ENTIDAD ---
    texto_empresa = rec.get("empresa_texto", "")
    if texto_empresa:
        _reemplazar_seccion_parrafos(doc, "PARA LA ENTIDAD", "En caso de cualquier inquietud", texto_empresa)

    # --- Firmas ---
    _reemplazar_en_parrafos(doc, "YAMILE M SUAREZ VARGAS", r.get("nombre_revisor", "[NOMBRE REVISOR]"))

    doc.save(str(dst))
    _convertir_a_pdf(str(dst))
    return str(dst)


def generar_carta_recomendaciones(datos, output_name=None):
    """FORMATO 3 — Carta de Recomendaciones / Reincorporación Laboral"""
    p, e, l, s, c, v, r, rec = _extraer_secciones(datos, extra=["visita", "rehabilitacion", "recomendaciones"])
    output_name = output_name or f"recomendaciones-{p.get('documento','sinid')}-{c.get('fecha','sinfecha')}"
    doc, dst = _preparar_documento("ejemplo carta de recomendaciones.docx", output_name, datos)

    # --- Tabla 0 — Datos del afiliado (label-value) ---
    _buscar_y_reemplazar(doc, "NOMBRE AFILIADO", p.get("nombre"))
    _buscar_y_reemplazar(doc, "DOCUMENTO CC", p.get("documento"))
    _buscar_y_reemplazar(doc, "TIPO DE SINIESTRO", s.get("tipo"))
    _buscar_y_reemplazar(doc, "FECHA DEL EVENTO", s.get("fecha_evento"))
    _buscar_y_reemplazar(doc, "SEGMENTO LESIONADO", s.get("segmento_lesionado"))
    _buscar_y_reemplazar(doc, "CARGO:", l.get("cargo"))
    _buscar_y_reemplazar(doc, "TIEMPO DE VIGENCIA", r.get("tiempo_vigencia"))
    _buscar_y_reemplazar(doc, "INCAPACITADO", "Sí" if s.get("incapacitado") else "No")
    _buscar_y_reemplazar(doc, "FECHA ULTIMA INCAPACIDAD", s.get("fecha_ultima_incapacidad"))
    _buscar_y_reemplazar(doc, "FORMA DE INTEGRACIÓN", r.get("forma_integracion"))

    # --- Concepto (texto largo en Tabla 0, reemplazar celda entera) ---
    concepto = c.get("concepto", "")
    if concepto:
        _reemplazar_celda_entera(doc, "Una vez realizada la Valoración de desempeño", concepto)

    # --- Reemplazar en párrafos del encabezado ---
    _reemplazar_en_parrafos(doc, "BOLIVARIANA DE MINERALES Y CIA LTDA", e.get("nombre") or "[VERIFICAR]")
    _reemplazar_en_parrafos(doc, "Ingrid Johan Vargas Reyes", e.get("contacto") or "[VERIFICAR]")
    _reemplazar_en_parrafos(doc, "jefe Talento Humano", e.get("cargo_contacto") or "[VERIFICAR]")
    _reemplazar_en_parrafos(doc, "KM 2 VIA NEIVA - PALERMO", e.get("direccion") or "[VERIFICAR]")
    _reemplazar_en_parrafos(doc, "3143009418", e.get("telefono") or "[VERIFICAR]")

    # También reemplazar en tablas (por si BOLIVARIANA aparece en celdas, no solo párrafos)
    nombre_empresa = e.get("nombre") or "[VERIFICAR]"
    _reemplazar_en_tablas(doc, "BOLIVARIANA DE MINERALES Y CIA LTDA", nombre_empresa)
    _reemplazar_en_tablas(doc, "Obrero de tratamiento roca", l.get("cargo") or e.get("cargo") or "[VERIFICAR]")
    _reemplazar_en_tablas(doc, "1193143688", p.get("documento") or "[VERIFICAR]")

    # --- Fecha de la carta ---
    if c.get("fecha"):
        _reemplazar_en_parrafos(doc, "09 de abril de 2026", c.get("fecha", ""))

    # --- Tabla 1: Tareas del cargo + Recomendaciones ---
    # La tabla tiene 1 fila con 2 celdas (izq=tareas, der=recomendaciones)
    tareas = rec.get("tareas", [])
    recs_trabajador = rec.get("trabajador", [])
    if len(doc.tables) >= 2:
        tabla1 = doc.tables[1]
        if tabla1.rows:
            fila = tabla1.rows[0]
            celdas_unicas = []
            vistos = set()
            for celda in fila.cells:
                tc = celda._tc
                if tc not in vistos:
                    vistos.add(tc)
                    celdas_unicas.append(celda)
            if len(celdas_unicas) >= 2:
                # Celda izquierda: tareas del cargo
                if tareas:
                    _poner_texto(celdas_unicas[0], "\n\n".join(tareas))
                # Celda derecha: recomendaciones para el trabajador
                if recs_trabajador:
                    _poner_texto(celdas_unicas[1], "\n\n".join(recs_trabajador))

    # --- RECOMENDACIONES PARA LA EMPRESA ---
    texto_empresa = rec.get("empresa_texto", "")
    if texto_empresa:
        _reemplazar_seccion_parrafos(
            doc,
            "RECOMENDACIONES PARA LA EMPRESA:",
            "Una vez se cumpla el tiempo de vigencia",
            texto_empresa
        )

    # --- Firma ---
    _insertar_firma(doc)

    doc.save(str(dst))
    _convertir_a_pdf(str(dst))
    return str(dst)


def generar_cierre_caso(datos, output_name=None):
    """FORMATO 4 — Certificado de Rehabilitación Integral / Cierre de Caso"""
    p, e, l, s, c, v, r, *_ = _extraer_secciones(datos, extra=["visita", "rehabilitacion", "recomendaciones"])
    output_name = output_name or f"cierre-{p.get('documento','sinid')}-{c.get('fecha','sinfecha')}"
    doc, dst = _preparar_documento("ejemplo cierre de caso.docx", output_name, datos)

    # Fecha y último día de incapacidad
    if c.get("fecha"):
        _poner_fecha_celdas(doc, "FECHA", c["fecha"])
    if s.get("ultimo_dia_incapacidad"):
        _poner_fecha_celdas(doc, "ÚLTIMO DIA DE INCAPACIDAD", s["ultimo_dia_incapacidad"])

    # Proveedor RHI
    _reemplazar_en_tablas(doc, "RILO SAS", "REHABILITACION INTEGRAL LABORAL Y OCUPACIONAL SAS")

    # Sección 1 — Identificación
    _buscar_y_reemplazar(doc, "Nombre del trabajador", p.get("nombre"))
    _buscar_y_reemplazar(doc, "Número de documento", p.get("documento"))
    _buscar_y_reemplazar(doc, "Identificación del siniestro", s.get("id_siniestro"))
    _buscar_y_reemplazar(doc, "EDAD (años cumplidos)", p.get("edad"))

    # Dominancia (checkbox)
    dom = p.get("dominancia", "")
    if dom:
        _marcar_opcion(doc, "Dominancia", dom.title())

    _buscar_y_reemplazar(doc, "Estado civil", p.get("estado_civil"))
    _buscar_y_reemplazar(doc, "Teléfonos del trabajador", p.get("telefono"))
    _buscar_y_reemplazar(doc, "Dirección residencia/ciudad", p.get("direccion"))
    _buscar_y_reemplazar(doc, "Diagnóstico(s) clínico(s)", s.get("diagnosticos"))

    # Nivel educativo — marcar opción
    nivel = p.get("nivel_educativo", "")
    if nivel:
        _marcar_opcion(doc, "Nivel educativo", nivel)

    # Datos laborales (están en filas siguientes, buscar por etiquetas cercanas)
    _reemplazar_en_tablas(doc, "JUAN CARLOS DURAN NARVAEZ", p.get("nombre") or "[VERIFICAR]")

    # Empresa
    _buscar_y_reemplazar(doc, "Nombre de la empresa", e.get("nombre"))
    _buscar_y_reemplazar(doc, "NIT de la empresa", e.get("nit"))
    _buscar_y_reemplazar(doc, "Cargo que desempeña", l.get("cargo") or e.get("cargo"))
    # Limpiar datos hardcodeados de Juan Carlos que puedan quedar en el template
    _reemplazar_en_tablas(doc, "BOLIVARIANA DE MINERALES Y CIA LTDA", e.get("nombre") or "[VERIFICAR]")
    _reemplazar_en_tablas(doc, "Obrero de tratamiento roca", l.get("cargo") or e.get("cargo") or "[VERIFICAR]")

    # Sección 2 — Decisión de reintegro (marcar según corresponda)
    forma = r.get("forma_integracion", "")
    if forma:
        _marcar_opcion(doc, "Reintegro sin modificaciones", forma)
        _marcar_opcion(doc, "Reintegro con modificaciones", forma)
        _marcar_opcion(doc, "Reubicación", forma)

    # Sección 3 — Compromiso / Deserción
    if r.get("compromiso"):
        _marcar_opcion(doc, "Sin compromiso", r.get("compromiso"))
    if r.get("desercion"):
        _marcar_opcion(doc, "Deserción", r.get("desercion"))

    # Sección 4 — Actividades (reemplazar la celda entera, no hacer append)
    actividades = (
        f"Proveedor Funcional:\n{r.get('actividades_funcional', '[VERIFICAR]')}\n"
        f"Proveedor Ocupacional:\n{r.get('actividades_ocupacional', '[VERIFICAR]')}"
    )
    _reemplazar_celda_entera(doc, "Proveedor Funcional:", actividades)

    # Sección 5 — Logros (reemplazar celda entera, sin bold en el cuerpo)
    motivo = c.get("motivo_consulta", "")
    enfermedad = c.get("enfermedad_actual", "")
    furat = c.get("descripcion_furat", "")
    examen = c.get("examen_fisico", "")
    analisis = c.get("analisis_recomendaciones", "")
    plan = c.get("plan_manejo", "")
    seccion5 = (
        f"MOTIVO DE CONSULTA:\n{motivo}\n\n"
        f"ENFERMEDAD ACTUAL:\n{enfermedad}"
    )
    if furat:
        seccion5 += f"\n\n{furat}"
    seccion5 += f"\n\nEXAMEN FÍSICO:\n{examen}"
    if analisis:
        seccion5 += f"\n\nANÁLISIS / RECOMENDACIONES:\n{analisis}"
    if plan:
        seccion5 += f"\n\nPLAN DE MANEJO:\n{plan}"
    _reemplazar_celda_entera(doc, "MOTIVO DE CONSULTA:", seccion5, bold=False)

    # Sección 6 — Obstáculos
    obstaculos = r.get("obstaculos") or "Ninguno frente al caso."
    _reemplazar_en_tablas(doc, "Ninguno frente al caso.", obstaculos)

    # Sección 7 — Concepto integral (reemplazar celda entera)
    concepto = r.get("concepto_integral", "[VERIFICAR]")
    _reemplazar_celda_entera(doc, "Afiliado de 35 años", concepto)

    # Sección 8 — Observación
    obs = r.get("observaciones") or "Ninguna frente al caso"
    _reemplazar_en_tablas(doc, "Ninguna frente al caso", obs)

    # Sección 9 — Firmas
    _insertar_firma(doc)

    doc.save(str(dst))
    _convertir_a_pdf(str(dst))
    return str(dst)


def generar_citacion_empresas(datos, output_name=None):
    """FORMATO 5 — Citacion de Empresas"""
    p, e, l, s, c, v = _extraer_secciones(datos, extra=["visita"])
    prov = datos.get("proveedor", {})
    output_name = output_name or f"citacion-{p.get('documento','sinid')}-{c.get('fecha','sinfecha')}"
    doc, dst = _preparar_documento("ejemplo formato de citacion de empresas.docx", output_name, datos)

    # --- Encabezado: fecha, destinatario ---
    if c.get("fecha"):
        _reemplazar_en_parrafos(doc, "Neiva, 23 febrero 2026", f"Neiva, {c['fecha']}")
    _reemplazar_en_parrafos(doc, "Seccional Rama Judicial Neiva", e.get("nombre") or "[VERIFICAR]")
    _reemplazar_en_parrafos(doc, "HEBERTH ARMANDO RUIZ PAVA", e.get("contacto") or "[VERIFICAR]")
    _reemplazar_en_parrafos(doc, "Coordinador de SST", e.get("cargo_contacto") or "[VERIFICAR]")
    _reemplazar_en_parrafos(doc, "Neiva-Huila", e.get("direccion") or "[VERIFICAR]")

    # --- Asunto (párrafo largo con nombre afiliado, CC, empresa) ---
    asunto_template = (
        f"Asunto: Programación de la visita a la empresa {e.get('nombre', '[EMPRESA]')} "
        f"con el fin de realizar: {v.get('tipo_estudio', '[ESTUDIO]')} del afiliado "
        f"{p.get('nombre', '[NOMBRE]')} identificado con CC {p.get('documento', '[CC]')} "
        f"por parte de la ARL positiva como parte del Programa de Rehabilitación "
        f"y Reincorporación Laboral Segura."
    )
    _reemplazar_parrafo(doc, "Programación de la visita a la empresa", asunto_template)

    # --- Tabla de datos de visita ---
    _buscar_y_reemplazar(doc, "CARGO", l.get("cargo"))
    _buscar_y_reemplazar(doc, "TIPO DE ESTUDIO", v.get("tipo_estudio"))
    _buscar_y_reemplazar(doc, "TIEMPO DE EJECUCIÓN", v.get("tiempo_ejecucion"))
    _buscar_y_reemplazar(doc, "OBJETIVO DEL ESTUDIO", v.get("objetivo"))
    _buscar_y_reemplazar(doc, "FECHA Y HORA", v.get("fecha_hora") or v.get("fecha"))
    _buscar_y_reemplazar(doc, "NOMBRE DEL PROFESIONAL", v.get("nombre_profesional"))
    _buscar_y_reemplazar(doc, "AUDITOR", v.get("nombre_auditor"))
    _buscar_y_reemplazar(doc, "DESCRIPCI", v.get("descripcion_estado"))
    _buscar_y_reemplazar(doc, "REQUISITOS", v.get("requisitos"))

    # --- Cierre: teléfono y correo del proveedor ---
    if prov.get("telefono"):
        _reemplazar_en_parrafos(doc, "3182675427", prov["telefono"])
    if prov.get("correo"):
        _reemplazar_en_parrafos(doc, "tocupacionalrilo@gmail.com", prov["correo"])

    # --- Firma ---
    _insertar_firma(doc)

    doc.save(str(dst))
    _convertir_a_pdf(str(dst))
    return str(dst)


# ============================================================
# AUXILIARES PARA PRUEBA DE TRABAJO
# ============================================================

def _reemplazar_tabla_tarea_critica(doc, idx, tarea):
    """Llena los datos de una tarea crítica en sus tablas correspondientes.
    Cada tarea usa 2 tablas: una de actividad/subactividad y otra de apreciaciones/conclusiones.
    idx=0 usa tablas 1-2, idx=1 usa tablas 3-4, idx=2 usa tablas 5-6.

    Estructura de las tablas de actividad en el template:
      T1 (idx=0, 5r): R0=encabezado sección "4 REQUERIMIENTOS...", R1=Actividad, R2=Subactividad,
                      R3=headers(Registro/Descripción), R4=contenido descripción
      T3,T5 (idx=1,2, 4r): R0=Actividad, R1=Subactividad,
                            R2=headers, R3=contenido descripción
    En todas: cols únicas = [label, valor, label, valor]

    Estructura de las tablas de apreciaciones (T2,T4,T6, 10r):
      R0=label "Apreciación del trabajador", R1=contenido
      R2=label "Apreciación del profesional...", R3=contenido
      R5=label "Conclusión...", R6=checkboxes
      R8=label "Descripción de la conclusión...", R9=contenido
    """
    if not doc.tables:
        return

    tabla_act = 1 + idx * 2
    tabla_conc = 2 + idx * 2

    if tabla_act >= len(doc.tables):
        return

    def _unique(row):
        seen = set()
        result = []
        for c in row.cells:
            if c._tc not in seen:
                seen.add(c._tc)
                result.append(c)
        return result

    # --- Tabla de Actividad/Subactividad ---
    t_act = doc.tables[tabla_act]
    # T1 (idx=0) has a section header in R0; actual data starts at R1
    # T3, T5 (idx>0) data starts at R0
    act_off = 1 if idx == 0 else 0

    # Actividad row: unique cells [label, value, Ciclo_label, ciclo_value]
    if len(t_act.rows) > act_off:
        u = _unique(t_act.rows[act_off])
        if len(u) > 1:
            _poner_texto(u[1], tarea.get("actividad", ""))
        if len(u) > 3:
            _poner_texto(u[3], tarea.get("ciclo", ""))

    # Subactividad row: unique cells [label, value, Autoadministrado_label, estandar_value]
    if len(t_act.rows) > act_off + 1:
        u = _unique(t_act.rows[act_off + 1])
        if len(u) > 1:
            _poner_texto(u[1], tarea.get("subactividad", ""))
        if len(u) > 3:
            _poner_texto(u[3], tarea.get("estandar", ""))

    # Descripción row: R(act_off+3) unique cells ['', description]
    desc_row = act_off + 3
    if len(t_act.rows) > desc_row:
        u = _unique(t_act.rows[desc_row])
        if len(u) > 1:
            _poner_texto(u[1], tarea.get("descripcion", ""))

    # --- Tabla de Apreciaciones/Conclusiones ---
    if tabla_conc >= len(doc.tables):
        return
    t_conc = doc.tables[tabla_conc]

    # R0 is label "Apreciación del trabajador" — preserve; write content to R1
    if len(t_conc.rows) > 1:
        u = _unique(t_conc.rows[1])
        if u:
            _poner_texto(u[0], tarea.get("apreciacion_trabajador", ""))

    # R2 is label "Apreciación del profesional..." — preserve; write content to R3
    if len(t_conc.rows) > 3:
        u = _unique(t_conc.rows[3])
        if u:
            _poner_texto(u[0], tarea.get("apreciacion_profesional", ""))

    # R5 is label "Conclusión..." — preserve; R6 has checkboxes
    conclusion = tarea.get("conclusion_tipo", "")
    if conclusion and len(t_conc.rows) > 5:
        _marcar_opcion_en_tabla(t_conc, conclusion)

    # R8 is label "Descripción de la conclusión..." — preserve; write content to R9
    if len(t_conc.rows) > 9:
        u = _unique(t_conc.rows[9])
        if u:
            _poner_texto(u[0], tarea.get("conclusion_descripcion", ""))


def _reemplazar_en_tabla_por_idx(tabla, fila_idx, col_start, texto):
    """Reemplaza el texto en celdas a partir de col_start en una fila.
    Útil para tablas con celdas fusionadas donde el contenido está en un bloque mergeado."""
    if fila_idx >= len(tabla.rows):
        return
    fila = tabla.rows[fila_idx]
    celdas_unicas = []
    vistos = set()
    for c in fila.cells:
        tc = c._tc
        if tc not in vistos:
            vistos.add(tc)
            celdas_unicas.append(c)
    if col_start < len(celdas_unicas):
        _poner_texto(celdas_unicas[col_start], texto)


def _marcar_opcion_en_tabla(tabla, opcion):
    """Marca X en la celda correspondiente a la opción dentro de una tabla de conclusiones."""
    for fila in tabla.rows:
        celdas_unicas = []
        vistos = set()
        for c in fila.cells:
            tc = c._tc
            if tc not in vistos:
                vistos.add(tc)
                celdas_unicas.append(c)
        for i, celda in enumerate(celdas_unicas):
            if opcion.lower() in celda.text.lower():
                # La X va en la celda anterior (columna del checkbox)
                if i > 0:
                    _poner_texto(celdas_unicas[i - 1], "X")
                return


def _reemplazar_tabla_materiales(doc, materiales):
    """Llena la tabla de Materiales, Equipos y Herramientas (Tabla 7, filas 2-8)."""
    if len(doc.tables) < 8:
        return
    t7 = doc.tables[7]
    # Las filas de materiales empiezan en fila 2
    for mi, mat in enumerate(materiales):
        ri = 2 + mi
        if ri >= len(t7.rows):
            break
        fila = t7.rows[ri]
        celdas_unicas = []
        vistos = set()
        for c in fila.cells:
            tc = c._tc
            if tc not in vistos:
                vistos.add(tc)
                celdas_unicas.append(c)
        # Col 0-2: Nombre, Col 3-5: Estado, Col 6-11: Requerimientos, Col 12-14: Observaciones
        # En celdas_unicas: [Nombre, Estado, Requerimientos, Observaciones, ...]
        if len(celdas_unicas) >= 1:
            _poner_texto(celdas_unicas[0], mat.get("nombre", ""))
        if len(celdas_unicas) >= 2:
            _poner_texto(celdas_unicas[1], mat.get("estado", ""))
        if len(celdas_unicas) >= 3:
            _poner_texto(celdas_unicas[2], mat.get("requerimientos", ""))
        if len(celdas_unicas) >= 4:
            _poner_texto(celdas_unicas[3], mat.get("observaciones", ""))


def _reemplazar_tabla_peligros(doc, peligros):
    """Llena la tabla de Peligros (Tabla 7, filas 12+)."""
    if len(doc.tables) < 8:
        return
    t7 = doc.tables[7]
    # Las filas de peligros empiezan en fila 12
    for pi, pel in enumerate(peligros):
        ri = 12 + pi
        if ri >= len(t7.rows):
            break
        fila = t7.rows[ri]
        celdas_unicas = []
        vistos = set()
        for c in fila.cells:
            tc = c._tc
            if tc not in vistos:
                vistos.add(tc)
                celdas_unicas.append(c)
        # Col 0-2: Nombre, Col 3-9: Descripción, Col 10-12: Control, Col 13-14: Recomendaciones
        if len(celdas_unicas) >= 1:
            _poner_texto(celdas_unicas[0], pel.get("nombre", ""))
        if len(celdas_unicas) >= 2:
            _poner_texto(celdas_unicas[1], pel.get("descripcion", ""))
        if len(celdas_unicas) >= 3:
            _poner_texto(celdas_unicas[2], pel.get("control", ""))
        if len(celdas_unicas) >= 4:
            _poner_texto(celdas_unicas[3], pel.get("recomendaciones", ""))


def generar_prueba_trabajo(datos, output_name=None):
    """FORMATO 6 — Prueba de Trabajo"""
    p, e, l, s, c, v = _extraer_secciones(datos, extra=["visita"])
    r = datos.get("rehabilitacion", {})
    _tareas_raw = datos.get("tareas_criticas") or datos.get("laboral", {}).get("tareas_criticas", [])
    tareas = [t for t in (_tareas_raw if isinstance(_tareas_raw, list) else []) if isinstance(t, dict)]
    _mat_raw = datos.get("materiales") or datos.get("laboral", {}).get("materiales", [])
    materiales = [m for m in (_mat_raw if isinstance(_mat_raw, list) else []) if isinstance(m, dict)]
    _pel_raw = datos.get("peligros") or datos.get("laboral", {}).get("peligros", [])
    peligros = [p for p in (_pel_raw if isinstance(_pel_raw, list) else []) if isinstance(p, dict)]
    output_name = output_name or f"prueba-{p.get('documento','sinid')}-{c.get('fecha','sinfecha')}"
    doc, dst = _preparar_documento("ejemplo prueba de trabajo.docx", output_name, datos)

    # === FECHAS ===
    if c.get("fecha"):
        _poner_fecha_celdas(doc, "FECHA DE VALORACIÓN", c["fecha"])
    if s.get("ultimo_dia_incapacidad"):
        _poner_fecha_celdas(doc, "ÚLTIMO DIA DE INCAPACIDAD", s["ultimo_dia_incapacidad"])

    # === SECCIÓN 1 — IDENTIFICACIÓN ===
    _buscar_y_reemplazar(doc, "Nombre del trabajador", p.get("nombre"))
    _buscar_y_reemplazar(doc, "Número de documento", p.get("documento"))
    _buscar_y_reemplazar(doc, "Identificación del siniestro", s.get("id_siniestro"))
    _buscar_y_reemplazar(doc, "Teléfonos trabajadores", p.get("telefono"))
    _buscar_y_reemplazar(doc, "Dirección residencia/ciudad", p.get("direccion"))
    _reemplazar_celda_entera(doc, "HERIDA DEL CUARTO DEDO", s.get("diagnosticos"), bold=False, font_size=7.5)
    _buscar_y_reemplazar(doc, "Fecha(s) del evento(s) ATEL", s.get("fecha_evento"))
    _buscar_y_reemplazar(doc, "EPS - IPS", p.get("eps_ips"))
    _buscar_y_reemplazar(doc, "AFP", p.get("afp"))
    # Tiempo total de incapacidad: solo reemplazar si NO es "Sin datos" (template tiene "0 | Sin datos")
    t_incap = s.get("tiempo_incapacidad", "")
    if t_incap and t_incap.strip().lower() != "sin datos":
        _buscar_y_reemplazar(doc, "Tiempo total de incapacidad", t_incap)
    _buscar_y_reemplazar(doc, "Empresa donde labora", e.get("nombre"))
    _buscar_y_reemplazar(doc, "NIT de la Empresa", e.get("nit"))
    _buscar_y_reemplazar(doc, "Cargo actual", l.get("cargo"))
    _buscar_y_reemplazar(doc, "Área/sección/proceso", l.get("area"))
    _buscar_y_reemplazar(doc, "Forma de vinculación laboral", l.get("vinculacion"))
    _buscar_y_reemplazar(doc, "Contacto en empresa/cargo", f"{e.get('contacto','')} / {e.get('cargo_contacto','')}")
    _buscar_y_reemplazar(doc, "Correo(s) electrónico(s)", e.get("correo"))
    _buscar_y_reemplazar(doc, "Teléfonos de contacto empresa", e.get("telefono"))
    _buscar_y_reemplazar(doc, "Dirección de empresa/ciudad", e.get("direccion"))

    # Edad y fecha nacimiento
    if p.get("fecha_nacimiento"):
        _poner_fecha_celdas(doc, "Fecha de nacimiento", p["fecha_nacimiento"])
    _buscar_y_reemplazar(doc, "63 años", p.get("edad", ""))

    # Fecha ingreso cargo
    if l.get("fecha_ingreso_cargo"):
        _poner_fecha_celdas(doc, "Fecha ingreso cargo", l["fecha_ingreso_cargo"])
    # Antigüedad cargo: acceso directo por índice (T0 F29 C14) para evitar
    # colisión de búsqueda con el mismo texto en antigüedad de empresa (mismo template text)
    ant_cargo = l.get("antiguedad_cargo", "")
    if ant_cargo:
        try:
            _poner_texto(doc.tables[0].rows[29].cells[14], ant_cargo)
        except (IndexError, AttributeError) as e:
            _log(f"Celda no encontrada ({e}) - campo 'antiguedad_cargo' omitido en prueba_trabajo")

    # Fecha ingreso empresa
    if l.get("fecha_ingreso_empresa"):
        _poner_fecha_celdas(doc, "Fecha ingreso a la empresa", l["fecha_ingreso_empresa"])
    # Antigüedad empresa: acceso directo por índice (T0 F31 C14)
    ant_emp = l.get("antiguedad_empresa", "")
    if ant_emp:
        try:
            _poner_texto(doc.tables[0].rows[31].cells[14], ant_emp)
        except (IndexError, AttributeError) as e:
            _log(f"Celda no encontrada ({e}) - campo 'antiguedad_empresa' omitido en prueba_trabajo")

    # Dominancia (checkbox)
    dom = p.get("dominancia", "")
    if dom:
        _marcar_opcion(doc, "Dominancia", dom.title())

    # Nivel educativo (checkbox)
    nivel = p.get("nivel_educativo", "")
    if nivel:
        _marcar_opcion(doc, "Nivel educativo", nivel)

    # Estado civil
    _buscar_y_reemplazar(doc, "Estado civil", p.get("estado_civil"))

    # === SECCIÓN 2 — METODOLOGÍA ===
    metodologia = c.get("metodologia", "")
    if metodologia:
        _reemplazar_celda_entera(doc, "La metodología utilizada para la Prueba de Trabajo", metodologia)

    # === SECCIÓN 3.1 — PROCESO PRODUCTIVO ===
    proceso = c.get("proceso_productivo", "")
    if proceso:
        _reemplazar_celda_entera(doc, "La Servidora manifiesta las siguientes funciones", proceso)

    # === SECCIÓN 3.2 — APRECIACIÓN DEL TRABAJADOR ===
    apreciacion = c.get("apreciacion_trabajador", "")
    if apreciacion:
        _reemplazar_celda_entera(doc, "La servidora manifiesta que permanece con dolor", apreciacion)

    # === SECCIÓN 3.3 — ESTÁNDARES DE PRODUCTIVIDAD ===
    estandares = c.get("estandares_productividad", "")
    if estandares:
        _reemplazar_celda_entera(doc, "La servidora menciona que el trabajo depende", estandares)

    # === SECCIÓN 3.4 — REQUERIMIENTOS DEL DESEMPEÑO ===
    _buscar_y_reemplazar(doc, "JORNADA", l.get("jornada", ""))
    _buscar_y_reemplazar(doc, "RITMO", l.get("ritmo", ""))
    _buscar_y_reemplazar(doc, "DESCANSOS", l.get("descansos", ""))
    _buscar_y_reemplazar(doc, "TURNOS", l.get("turnos", ""))
    _buscar_y_reemplazar(doc, "HORAS EXTRAS", l.get("horas_extras", ""))
    _buscar_y_reemplazar(doc, "ROTACIONES", l.get("rotaciones", ""))
    _buscar_y_reemplazar(doc, "TIEMPOS EFECTIVOS", l.get("tiempos_efectivos", ""))
    _buscar_y_reemplazar(doc, "DISTRIBUCIÓN SEMANAL", l.get("distribucion_semanal", ""))

    # Pre-clean conclusion checkboxes in T2, T4, T6 (template pre-fills "x" for first option)
    for ti_conc in [2, 4, 6]:
        if ti_conc < len(doc.tables):
            for fila_c in doc.tables[ti_conc].rows:
                celdas_c = fila_c.cells
                row_txt = " ".join(cel.text.strip().lower() for cel in celdas_c)
                if "puede desempeñarla" in row_txt or "restricción parcial" in row_txt:
                    seen_c = set()
                    for cel in celdas_c:
                        if cel._tc not in seen_c:
                            seen_c.add(cel._tc)
                            if cel.text.strip().upper() == "X":
                                _poner_texto(cel, "")

    # === SECCIÓN 4 — TAREAS CRÍTICAS ===
    # Cada tarea tiene su propia tabla (1-6). Buscar la tabla por el nombre de actividad.
    for idx, tarea in enumerate(tareas):
        _reemplazar_tabla_tarea_critica(doc, idx, tarea)

    # === SECCIÓN 5 — MATERIALES (Tabla 7, filas 2-8) ===
    _reemplazar_tabla_materiales(doc, materiales)

    # === SECCIÓN 6 — PELIGROS (Tabla 7, filas 11+) ===
    _reemplazar_tabla_peligros(doc, peligros)

    # === SECCIÓN 7 — CONCEPTO CAPACIDAD DE DESEMPEÑO ===
    concepto = c.get("concepto_desempeno", "")
    if concepto:
        _reemplazar_celda_entera(doc, "Se realizó la Prueba de Trabajo (PT)", concepto)

    # === SECCIÓN 9 — FIRMAS ===
    _insertar_firma(doc)

    doc.save(str(dst))
    _convertir_a_pdf(str(dst))
    return str(dst)


def generar_valoracion_desempeno(datos, output_name=None):
    """FORMATO 7 — Valoración del Desempeño Ocupacional Final.
    
    Estructura del template:
      Tabla 0 (38 filas x 35 cols): Sección 1-2 (Identificación completa)
      Tabla 1 (24 filas x 14 cols): Sección 3-6 (Historia, Descripción, Rol, Tratamiento)
      Tabla 2 (39 filas x 12 cols): Adaptaciones + Sección 7 (Composición familiar) + Sección 8 parcial
      Tabla 3 (40 filas x 7 cols): Sección 8 cont. (Movilidad, Aprendizaje, Vida doméstica) + Sección 9
      Tabla 4 (8 filas x 3 cols): Concepto cont. + Sección 10 (Orientación) + Sección 11 (Registro)
    """
    p, e, l, s, c = _extraer_secciones(datos)
    # Detectar género del paciente para adaptar el lenguaje
    genero = _detectar_genero(p.get("nombre", ""))
    output_name = output_name or f"valoracion-{p.get('documento','sinid')}-{c.get('fecha','sinfecha')}"
    doc, dst = _preparar_documento("ejemplo valoracion de desempeño ocupacional.docx", output_name, datos)

    # ─── SECCIÓN 1 — Objetivo (Tabla 0, F02-F03, texto fijo de plantilla, NO se modifica) ───

    # ─── SECCIÓN 2 — Identificación (Tabla 0, F06-F36) ───

    # Fecha de valoración (celdas individuales día/mes/año en F00-F01)
    if c.get("fecha"):
        _poner_fecha_celdas(doc, "FECHA DE VALORACIÓN", c["fecha"])
    else:
        # Limpiar la fecha del template (viene con 06/04/2026 de Juan Carlos)
        for tabla_fv in doc.tables:
            for ri_fv, fila_fv in enumerate(tabla_fv.rows):
                if any("fecha de valoraci" in cel.text.lower() for cel in fila_fv.cells):
                    for ri2 in range(ri_fv, min(ri_fv + 3, len(tabla_fv.rows))):
                        fila2 = tabla_fv.rows[ri2]
                        for cel2 in fila2.cells:
                            import re as _re
                            if _re.match(r'^\d{1,4}$', cel2.text.strip()):
                                _poner_texto(cel2, "")
                    break

    # Campos label→value (buscar por substrings exactos del template)
    _buscar_y_reemplazar(doc, "Nombre del Trabajador", p.get("nombre"), table_indices=[0])
    _buscar_y_reemplazar(doc, "Número de documento", p.get("documento"), table_indices=[0])
    _buscar_y_reemplazar(doc, "Identificación del siniestro", s.get("id_siniestro"), table_indices=[0])

    # Fecha de nacimiento en celdas individuales (F09-F10)
    if p.get("fecha_nacimiento"):
        _poner_fecha_celdas(doc, "Fecha de nacimiento/edad", p["fecha_nacimiento"])
    # Edad: acceso directo por índice (T0 F10 C17) para evitar reemplazo
    # incorrecto con búsqueda genérica de "años" que matchea múltiples celdas
    edad_val = p.get("edad", "")
    if edad_val:
        try:
            _poner_texto(doc.tables[0].rows[10].cells[17], str(edad_val))
        except (IndexError, AttributeError) as e:
            _log(f"Celda no encontrada ({e}) - campo 'edad' en valoracion_desempeno, usando fallback")
            _reemplazar_en_tablas(doc, "35", str(edad_val))

    # Dominancia con checkbox
    dom = p.get("dominancia", "")
    if dom:
        _marcar_opcion(doc, "Dominancia", dom, table_indices=[0])

    _buscar_y_reemplazar(doc, "Estado civil", p.get("estado_civil"), table_indices=[0])

    # Nivel educativo — checkbox directo por fila (Tabla 0, F13-F15)
    # Estructura: cada fila F13-F15 tiene opciones agrupadas y el checkbox X en C31-C34
    # F13: Formacion empirica | Basica primaria | Bachillerato vocacional 9°
    # F14: Bachillerato modalidad | Tecnico/Tecnologico | Profesional
    # F15: Especializacion/postgrado/maestria | Formacion informal oficios | Analfabeta
    nivel = p.get("nivel_educativo", "")
    if nivel:
        nivel_lower = nivel.lower().strip()
        _nivel_fila_map = {
            "formacion empirica": 13, "formación empírica": 13,
            "basica primaria": 13, "básica primaria": 13,
            "bachillerato vocacional": 13,
            "bachillerato modalidad": 14,
            "tecnico": 14, "tecnológico": 14,
            "profesional": 14,
            "especializacion": 15, "postgrado": 15, "maestria": 15, "maestría": 15,
            "formacion informal": 15, "formación informal": 15,
            "analfabeta": 15,
        }
        _target_fila = None
        for k, v in _nivel_fila_map.items():
            if k in nivel_lower:
                _target_fila = v
                break
        if _target_fila is not None:
            for _f in [13, 14, 15]:
                try:
                    _row = doc.tables[0].rows[_f]
                    if len(_row.cells) > 31 and _row.cells[31].text.strip().upper() == "X":
                        _poner_texto(_row.cells[31], "")
                except (IndexError, AttributeError):
                    pass
            try:
                _poner_texto(doc.tables[0].rows[_target_fila].cells[31], "X")
            except (IndexError, AttributeError) as e:
                _log(f"[WARN] No se pudo marcar nivel educativo: {e}")
        else:
            _log(f"[WARN] Nivel educativo no reconocido: '{nivel}'")

    # Especificar formación y oficios
    formacion = p.get("formacion_oficios") or datos.get("formacion_oficios", "")
    if formacion:
        _buscar_y_reemplazar(doc, "Especificar formación y oficios que conoce", formacion, table_indices=[0])

    _buscar_y_reemplazar(doc, "Teléfonos trabajadores", p.get("telefono"), table_indices=[0])

    # Dirección residencia y ciudad — con checkbox Urbano/Rural
    dir_val = p.get("direccion", "")
    if dir_val:
        _buscar_y_reemplazar(doc, "Dirección residencia y ciudad", dir_val, table_indices=[0])
    # Urbano/Rural — checkbox directo (Tabla 0 F18, C28=Urbano, C34=Rural)
    zona = p.get("zona", "")
    if zona:
        try:
            _fila_ur = doc.tables[0].rows[18]
            _zona_clean = zona.strip().lower()
            if "urbano" in _zona_clean:
                _poner_texto(_fila_ur.cells[28], "X")
                if len(_fila_ur.cells) > 34:
                    _poner_texto(_fila_ur.cells[34], "")
            elif "rural" in _zona_clean:
                _poner_texto(_fila_ur.cells[28], "")
                if len(_fila_ur.cells) > 34:
                    _poner_texto(_fila_ur.cells[34], "X")
        except (IndexError, AttributeError) as e:
            _log(f"[WARN] Error marcando Urbano/Rural: {e}")

    diag_val = s.get("diagnostico_cie10") or s.get("diagnosticos") or ""
    if diag_val:
        # Intentar reemplazar la celda con el diagnóstico del template de Juan Carlos
        reemplazado = _reemplazar_celda_entera(doc, "HERIDA DEL CUARTO DEDO", diag_val, bold=False, font_size=7.5, table_indices=[0])
        if not reemplazado:
            # Fallback: buscar la celda de Diagnóstico por su label
            _reemplazar_celda_entera(doc, "Diagnóstico(s) clínico(s)", diag_val, bold=False, font_size=7.5, table_indices=[0])
    else:
        # Limpiar datos de Juan Carlos aunque no tengamos diagnóstico propio
        _reemplazar_celda_entera(doc, "HERIDA DEL CUARTO DEDO", "", bold=False, font_size=7.5, table_indices=[0])
    _buscar_y_reemplazar(doc, "Fecha(s) del evento(s) ATEL", s.get("fecha_evento"), table_indices=[0])

    # Eventos No laborales (Si/No checkbox + campos fecha/diagnóstico)
    ev_nolab = s.get("eventos_no_laborales", "")
    if ev_nolab:
        _marcar_opcion(doc, "Eventos No laborales", ev_nolab, table_indices=[0])
        # Workaround: limpiar celdas checkbox de Diagnostico en la fila
        # de Eventos No laborales. La celda tc[9] abarca columnas 22-34
        # (gridSpan=13) y _marcar_opcion puede dejar X residuales del template.
        for tabla in doc.tables:
            for fila in tabla.rows:
                if any("eventos no laborales" in c.text.lower() for c in fila.cells):
                    celdas = fila.cells
                    for j in range(22, len(celdas)):
                        if celdas[j].text.strip().upper() == "X":
                            _poner_texto(celdas[j], "")

    _buscar_y_reemplazar(doc, "EPS - IPS", p.get("eps_ips"), table_indices=[0])
    _buscar_y_reemplazar(doc, "AFP", p.get("afp"), table_indices=[0])
    # Tiempo total de incapacidad: solo reemplazar si NO es "Sin datos" (template tiene "0 | Sin datos")
    t_incap = s.get("tiempo_incapacidad", "")
    if t_incap and t_incap.strip().lower() != "sin datos":
        _buscar_y_reemplazar(doc, "Tiempo total de incapacidad", t_incap, table_indices=[0])
    _buscar_y_reemplazar(doc, "Empresa donde labora", e.get("nombre"), table_indices=[0])

    # Vinculación laboral (NO/SI checkbox) — acceso directo (Tabla 0 F26, C1=NO, C8=SI)
    vinc = l.get("vinculacion", "")
    if vinc:
        try:
            _fila_vinc = doc.tables[0].rows[26]
            _vinc_clean = vinc.strip().upper()
            _poner_texto(_fila_vinc.cells[1], "")  # limpiar NO
            _poner_texto(_fila_vinc.cells[8], "")  # limpiar SI
            if "SI" in _vinc_clean:
                _poner_texto(_fila_vinc.cells[8], "X")
            elif "NO" in _vinc_clean:
                _poner_texto(_fila_vinc.cells[1], "X")
        except (IndexError, AttributeError) as e:
            _log(f"[WARN] Error marcando Vinculacion laboral: {e}")

    _buscar_y_reemplazar(doc, "Forma de vinculación laboral", l.get("forma_vinculacion"), table_indices=[0])

    # Modalidad — checkbox directo (Tabla 0 F28, C9=Presencial, C21=teletrabajo, C32=trabajo en casa)
    mod = l.get("modalidad", "")
    if mod:
        try:
            _fila_mod = doc.tables[0].rows[28]
            _mod_clean = mod.strip().lower()
            _poner_texto(_fila_mod.cells[9], "")   # limpiar Presencial X
            _poner_texto(_fila_mod.cells[21], "")  # limpiar teletrabajo X
            _poner_texto(_fila_mod.cells[32], "")  # limpiar trabajo en casa X
            if "presencial" in _mod_clean:
                _poner_texto(_fila_mod.cells[9], "X")
            elif "teletrabajo" in _mod_clean or "virtual" in _mod_clean:
                _poner_texto(_fila_mod.cells[21], "X")
            elif "trabajo en casa" in _mod_clean:
                _poner_texto(_fila_mod.cells[32], "X")
        except (IndexError, AttributeError) as e:
            _log(f"[WARN] Error marcando Modalidad: {e}")
    tiempo_modal = l.get("tiempo_modalidad") or l.get("antiguedad_empresa") or ""
    _buscar_y_reemplazar(doc, "Tiempo de la modalidad", tiempo_modal, table_indices=[0])

    _buscar_y_reemplazar(doc, "NIT de la Empresa", e.get("nit"), table_indices=[0])

    # Fecha ingreso empresa en celdas individuales (F30)
    if l.get("fecha_ingreso_empresa"):
        _poner_fecha_celdas(doc, "Fecha ingreso a la empresa", l["fecha_ingreso_empresa"])
    else:
        # Limpiar la fecha del template (viene con 16/12/2024 de Juan Carlos)
        # Buscar la fila y limpiar celdas numéricas de fecha
        for tabla_fi in doc.tables:
            for ri_fi, fila_fi in enumerate(tabla_fi.rows):
                if any("fecha ingreso a la empresa" in cel.text.lower() for cel in fila_fi.cells):
                    # Limpiar la siguiente fila que tenga dígitos de fecha
                    for ri2 in range(ri_fi, min(ri_fi + 3, len(tabla_fi.rows))):
                        fila2 = tabla_fi.rows[ri2]
                        for cel2 in fila2.cells:
                            import re as _re
                            if _re.match(r'^\d{1,4}$', cel2.text.strip()):
                                _poner_texto(cel2, "")
                    break
    # Antigüedad en la empresa: acceso directo (T0 F32 C19)
    ant_emp = l.get("antiguedad_empresa", "")
    if ant_emp:
        try:
            _poner_texto(doc.tables[0].rows[32].cells[19], ant_emp)
        except (IndexError, AttributeError) as e:
            _log(f"Celda no encontrada ({e}) - campo 'antiguedad_empresa' omitido en valoracion_desempeno")
    else:
        # Limpiar la antigüedad del template (viene con "1 año 4 meses aprox." de Juan Carlos)
        try:
            cel_ant = doc.tables[0].rows[32].cells[19]
            if cel_ant.text.strip() not in ("", "0"):
                _poner_texto(cel_ant, "")
        except (IndexError, AttributeError):
            pass

    contacto_str = e.get("contacto") or ""
    cargo_str = e.get("cargo_contacto") or ""
    if contacto_str and cargo_str:
        contacto_str = f"{contacto_str}/ {cargo_str}"
    elif cargo_str:
        contacto_str = cargo_str
    _buscar_y_reemplazar(doc, "Contacto en empresa", contacto_str, table_indices=[0])
    _buscar_y_reemplazar(doc, "Correo(s) electrónico(s)", e.get("correo"), table_indices=[0])
    _buscar_y_reemplazar(doc, "Teléfonos de contacto empresa", e.get("telefono"), table_indices=[0])
    _buscar_y_reemplazar(doc, "Dirección empresa", e.get("direccion"), table_indices=[0])

    # ─── SECCIÓN 3 — Historia ocupacional (Tabla 1, F02-F06) ───
    tabla1 = doc.tables[1]
    _hist_raw = datos.get("historia_ocupacional")
    historia = [h for h in (_hist_raw if isinstance(_hist_raw, list) else []) if isinstance(h, dict)]
    # La plantilla tiene 2 filas de ejemplo (F03, F04) con 4 columnas:
    # Empresa(cols 0-1) | Cargo(cols 2-3) | Tiempo(cols 4-7) | Motivo(cols 8-13)
    # Escribir cada entrada de historia en la fila correspondiente
    fila_base = 3  # F03 es la primera fila de datos en Tabla 1
    # Limpiar SOLO las 2 filas de historia del template (F03-F04 tienen datos de Juan Carlos)
    # range(2): solo F03 and F04. range(4) rompería "Otros Oficios" (F05) y "Oficios de interés" (F06)
    for idx_clear in range(2):
        if fila_base + idx_clear < len(tabla1.rows):
            fila_clear = tabla1.rows[fila_base + idx_clear]
            celdas_clear = fila_clear.cells
            _poner_texto(celdas_clear[0], "")
            _poner_texto(celdas_clear[2], "")
            _poner_texto(celdas_clear[4], "")
            _poner_texto(celdas_clear[8], "")
    for idx_hist, entrada in enumerate(historia[:2]):
        fila = tabla1.rows[fila_base + idx_hist]
        celdas = fila.cells
        _poner_texto(celdas[0], entrada.get("empresa", ""))
        _poner_texto(celdas[2], entrada.get("cargo", ""))
        _poner_texto(celdas[4], entrada.get("duracion", ""))
        _poner_texto(celdas[8], entrada.get("motivo_retiro", ""))

    # Para las entradas 3+: insertar nuevas filas ANTES de "Otros Oficios" en lugar de texto plano
    historia_extra = historia[2:]
    if historia_extra:
        fila_modelo = tabla1.rows[fila_base + 1]  # copia F04 (segunda fila de datos del template)
        offset = 0
        for entrada in historia_extra:
            new_tr = deepcopy(fila_modelo._tr)
            fila_otros_idx = fila_base + 2 + offset  # "Otros Oficios" se desplaza 1 por cada inserción
            if fila_otros_idx < len(tabla1.rows):
                fila_otros = tabla1.rows[fila_otros_idx]
                fila_otros._tr.addprevious(new_tr)
                nueva_fila = tabla1.rows[fila_base + 2 + offset]
                celdas_nueva = nueva_fila.cells
                _poner_texto(celdas_nueva[0], entrada.get("empresa", ""))
                _poner_texto(celdas_nueva[2], entrada.get("cargo", ""))
                _poner_texto(celdas_nueva[4], entrada.get("duracion", ""))
                _poner_texto(celdas_nueva[8], entrada.get("motivo_retiro", ""))
                offset += 1

    # Otros oficios y oficios de interés (F05, F06)
    otros = datos.get("otros_oficios", "")
    if otros:
        _buscar_y_reemplazar(doc, "Otros Oficios desempeñados", otros, table_indices=[1])
    interes = datos.get("oficios_interes", "")
    if interes:
        _buscar_y_reemplazar(doc, "Oficios de interés", interes, table_indices=[1])

    # ─── SECCIÓN 4 — Descripción actividad laboral actual (Tabla 1, F08-F15) ───
    # Sección 4: filas con 1 celda merged (formato "Label: valor") — _buscar_y_reemplazar falla
    # Usar _reemplazar_celda_entera que reemplaza TODO el contenido de la celda encontrada
    # Fallback a empresa.cargo cuando laboral.cargo está vacío (el LLM puede poner el cargo en empresa)
    cargo_val = l.get("cargo") or e.get("cargo") or ""
    tareas_val = l.get("tareas") or ""
    herr_val = l.get("herramientas") or ""
    horario_val = l.get("horario") or ""
    epp_val = l.get("epp") or ""
    ant_cargo_val = l.get("antiguedad_cargo") or ""
    req_val = l.get("requerimientos_motrices") or ""
    _reemplazar_celda_entera(doc, "Nombre del cargo:",
        f"Nombre del cargo: {cargo_val}", font_size=7.5, table_indices=[1])
    _reemplazar_celda_entera(doc, "Tareas (nombre y descripción):",
        f"Tareas (nombre y descripción): {tareas_val}", font_size=7.5, table_indices=[1])
    _reemplazar_celda_entera(doc, "Herramientas de trabajo:",
        f"Herramientas de trabajo: {herr_val}", font_size=7.5, table_indices=[1])
    _reemplazar_celda_entera(doc, "Horario de trabajo:",
        f"Horario de trabajo: {horario_val}", font_size=7.5, table_indices=[1])
    _reemplazar_celda_entera(doc, "Elementos de Protección:",
        f"Elementos de Protección: {epp_val}", font_size=7.5, table_indices=[1])
    _reemplazar_celda_entera(doc, "Antigüedad en el cargo:",
        f"Antigüedad en el cargo: {ant_cargo_val}", font_size=7.5, table_indices=[1])
    _reemplazar_celda_entera(doc, "Requerimientos motrices de la actividad:",
        f"Requerimientos motrices de la actividad: {req_val}", font_size=7.5, table_indices=[1])

    # Ocurrencia del ATEL (Tabla 1, buscar por label "Ocurrencia del ATEL")
    ocurrencia = l.get("ocurrencia_atel", "")
    lugar = l.get("lugar_atel", "")
    if ocurrencia or lugar:
        try:
            _fila_oc = None
            for _r in doc.tables[1].rows:
                if len(_r.cells) > 0 and "ocurrencia del atel" in _r.cells[0].text.lower():
                    _fila_oc = _r
                    break
            if _fila_oc is not None and len(_fila_oc.cells) > 10:
                # Limpiar todos los checkboxes previos
                for _ci in [1, 4, 5, 9, 10]:
                    if _ci < len(_fila_oc.cells):
                        _poner_texto(_fila_oc.cells[_ci], "")
                # Marcar SI/NO ocurrencia
                if ocurrencia:
                    if ocurrencia.strip().upper() == "SI":
                        _poner_texto(_fila_oc.cells[1], "X")
                    elif ocurrencia.strip().upper() == "NO":
                        pass  # NO checkbox no tiene celda separada
                # Marcar lugar
                if lugar:
                    lugar_upper = lugar.strip().upper()
                    if "PUESTO DE TRABAJO" in lugar_upper:
                        _poner_texto(_fila_oc.cells[4], "X")  # SI Puesto
                    elif "AREA" in lugar_upper:
                        _poner_texto(_fila_oc.cells[9], "X")  # SI Area
        except (IndexError, AttributeError) as e:
            _log(f"[WARN] Error marcando Ocurrencia ATEL: {e}")

    # ─── SECCIÓN 5 — Rol laboral (Tabla 1, buscar por label) ───
    rol = datos.get("rol_laboral", {}) or {}
    try:
        _t1 = doc.tables[1]
        # Buscar filas por label en lugar de indice fijo (la historia ocupacional inserta filas)
        _rol_filas = {}
        for _ri, _r in enumerate(_t1.rows):
            _txt = _r.cells[0].text.lower() if len(_r.cells) > 0 else ""
            if "tareas y operaciones" in _txt:
                _rol_filas["tareas"] = _r
            elif "componentes del desempe" in _txt:
                _rol_filas["componentes"] = _r
            elif "tiempo de ejecuci" in _txt:
                _rol_filas["tiempo"] = _r
            elif "forma de integraci" in _txt:
                _rol_filas["forma"] = _r
        
        _tareas_op = rol.get("tareas_operaciones") or ""
        if _tareas_op and "tareas" in _rol_filas and len(_rol_filas["tareas"].cells) > 2:
            _poner_texto(_rol_filas["tareas"].cells[2], f"Rol Laboral: {_tareas_op}", font_size=7.5)
        
        _sens = rol.get("sensorio_motor") or ""
        _cog = rol.get("cognitivo") or ""
        _psic = rol.get("psicologicos") or ""
        _soc = rol.get("social") or ""
        _componentes = ""
        if _sens: _componentes += f"Sensorio - motor: {_sens}\n"
        if _cog: _componentes += f"Cognitivo: {_cog}\n"
        if _psic: _componentes += f"Psicológicos: {_psic}\n"
        if _soc: _componentes += f"Social: {_soc}"
        if _componentes.strip() and "componentes" in _rol_filas and len(_rol_filas["componentes"].cells) > 2:
            _poner_texto(_rol_filas["componentes"].cells[2], _componentes.strip(), font_size=7.5)
        
        _tiempo_ej = rol.get("tiempo_ejecucion") or ""
        if _tiempo_ej and "tiempo" in _rol_filas and len(_rol_filas["tiempo"].cells) > 2:
            _poner_texto(_rol_filas["tiempo"].cells[2], f"Tiempo de Ejecución: {_tiempo_ej}", font_size=7.5)
        
        _forma_int = rol.get("forma_integracion") or ""
        if _forma_int and "forma" in _rol_filas and len(_rol_filas["forma"].cells) > 2:
            _poner_texto(_rol_filas["forma"].cells[2], f"Forma de integración laboral: {_forma_int}", font_size=7.5)
    except (IndexError, AttributeError) as e:
        _log(f"[WARN] Error en Rol Laboral: {e}")

    # ─── SECCIÓN 6 — Tratamiento ATEL (Tabla 1, F23) ───
    # tratamiento_atel puede ser: string directo, lista de dicts, o None
    trat_raw = datos.get("tratamiento_atel") or ""
    if isinstance(trat_raw, list):
        # Lista de dicts → formatear
        partes_trat = []
        for t in trat_raw:
            if isinstance(t, dict):
                partes_trat.append(f"{t.get('fecha','')} — {t.get('especialidad','')}. {t.get('ips','')}\n{t.get('contenido','')}")
            elif isinstance(t, str):
                partes_trat.append(t)
        texto_trat = "\n\n".join(partes_trat)
    else:
        texto_trat = str(trat_raw)
    
    texto_trat_clean = texto_trat.strip()
    
    # Sección 6 — Tratamiento ATEL (buscar por label "Tratamiento recibido")
    tratamiento_reemplazado = False
    try:
        for _r_trat in doc.tables[1].rows:
            if len(_r_trat.cells) > 0 and "tratamiento recibido" in _r_trat.cells[0].text.lower():
                if len(_r_trat.cells) > 2:
                    _poner_texto(_r_trat.cells[2], texto_trat_clean, bold=False, font_size=7.5)
                    tratamiento_reemplazado = True
                elif len(_r_trat.cells) > 1:
                    _poner_texto(_r_trat.cells[1], texto_trat_clean, bold=False, font_size=7.5)
                    tratamiento_reemplazado = True
                break
    except (IndexError, AttributeError) as e:
        _log(f"[WARN] Error en busqueda tratamiento: {e}")
                
    if not tratamiento_reemplazado:
        if texto_trat_clean:
            reemplazado = _reemplazar_celda_entera(doc, "ENFERMEDAD ACTUAL",
                                                   texto_trat_clean, bold=False, font_size=7.5, table_indices=[1])
            if not reemplazado:
                _reemplazar_celda_entera(doc, "Tratamiento recibido por Rehabilitación",
                                         texto_trat_clean, bold=False, font_size=7.5, table_indices=[1])
        else:
            _reemplazar_celda_entera(doc, "ENFERMEDAD ACTUAL", "", bold=False, font_size=7.5, table_indices=[1])

    # ─── SECCIÓN 7 — Composición familiar (Tabla 2, F11-F17) ───
    fam = datos.get("composicion_familiar", {})
    _buscar_y_reemplazar(doc, "Composición del núcleo familiar", fam.get("nucleo"), table_indices=[2])
    _buscar_y_reemplazar(doc, "Fecha de nacimiento de cada integrante", fam.get("integrantes"), table_indices=[2])
    _buscar_y_reemplazar(doc, "Persona(s) que sostiene económicamente el hogar", fam.get("sostenedor"), table_indices=[2])
    _buscar_y_reemplazar(doc, "Ingreso Promedio en el hogar", fam.get("ingreso_promedio"), table_indices=[2])
    _buscar_y_reemplazar(doc, "Responsabilidad económica en el hogar", fam.get("responsabilidad"), table_indices=[2])
    _buscar_y_reemplazar(doc, "Convivencia actual", fam.get("convivencia"), table_indices=[2])

    # ─── Adaptaciones (Tabla 2, F00-F09) ───
    adaptaciones = datos.get("adaptaciones", {})
    if isinstance(adaptaciones, dict):
        for nombre, estado in adaptaciones.items():
            if estado and estado.lower() != "no":
                _marcar_opcion(doc, nombre, "SI")

    # Calificación de PCL (F09)
    if s.get("pcl_aplica"):
        _marcar_opcion(doc, "Calificación de PCL", "SI", table_indices=[2])
    else:
        _marcar_opcion(doc, "Calificación de PCL", "NO", table_indices=[2])

    # ─── SECCIÓN 8 — Evaluación otras áreas ocupacionales ───
    areas = datos.get("areas_ocupacionales", {})

    # Función helper para marcar nivel de dificultad en tabla 2 o 3 (v2 corregida)
    def _marcar_nivel_dificultad(doc, item_label, nivel_str, observacion=""):
        """Busca la fila que contiene item_label y marca X en la columna de nivel.
        Tabla 2 (12 cols): label(0-1) | NO DIF(2-3) | LEVE(4-5) | MOD(6-7) | SEV(8-9) | COMP(10) | OBS(11)
        Tabla 3 (7 cols):  label(0) | NO DIF(1) | LEVE(2) | MOD(3) | SEV(4) | COMP(5) | OBS(6)
        - Busca SOLO en las primeras 2 celdas (label area) para evitar falsos matches
        - Añade logging detallado para items no encontrados
        """
        import unicodedata
        import re
        import sys
        
        def _clean(s):
            s = unicodedata.normalize("NFD", s.lower().strip())
            s = "".join(c for c in s if unicodedata.category(c) != "Mn")
            return re.sub(r'[\s\-/,.:;()]+', '', s)

        target = _clean(item_label)
        if not target:
            return False

        nivel_lower = nivel_str.lower().strip() if nivel_str else ""
        if not nivel_lower or nivel_lower in ("n/a", "na", ""):
            return False

        # Mapeos semánticos específicos entre el JSON del LLM y el template del Word
        _mapeos_items = {
            "higiene personal": "higiene personal",
            "hacer los quehaceres de la casa": "realizar quehaceres",
            "quehaceres de la casa": "realizar quehaceres",
            "realizar quehaceres de la casa": "realizar quehaceres",
            "limpiar / hacer oficio": "limpieza de la vivienda",
            "limpiar hacer oficio": "limpieza de la vivienda",
            "limpieza de la vivienda": "limpieza de la vivienda",
            "produccion de mensajes no verbales": "produccion de mensajes no",
        }
        for k, v in _mapeos_items.items():
            if _clean(k) in target:
                target = _clean(v)
                break

        for tabla in doc.tables[2:4]:
            num_cols = len(tabla.columns)
            if num_cols >= 12:
                _nivel_map = {
                    "no dificultad": 2, "dificultad leve": 4,
                    "dificultad moderada": 6, "dificultad severa": 8,
                    "dificultad completa": 10,
                }
                _obs_col = 11
            else:
                _nivel_map = {
                    "no dificultad": 1, "dificultad leve": 2,
                    "dificultad moderada": 3, "dificultad severa": 4,
                    "dificultad completa": 5,
                }
                _obs_col = 6

            col_target = None
            for k, col in _nivel_map.items():
                if k in nivel_lower:
                    col_target = col
                    break
            if col_target is None:
                continue

            for fila in tabla.rows:
                celdas = fila.cells
                # Buscar SOLO en las primeras 2 celdas (label area)
                found = False
                for c in celdas[:2]:
                    cell_text = _clean(c.text)
                    if target in cell_text:
                        found = True
                        break
                if not found:
                    continue

                # Limpiar SOLO celdas con X previa
                for col_clear in _nivel_map.values():
                    if col_clear < len(celdas):
                        if celdas[col_clear].text.strip().upper() == "X":
                            _poner_texto(celdas[col_clear], "")

                # Marcar X
                if col_target < len(celdas):
                    _poner_texto(celdas[col_target], "X")
                # Observación
                if observacion and _obs_col < len(celdas):
                    _poner_texto(celdas[_obs_col], observacion)
                return True

        _log(f"[WARN] Ítem de evaluación no encontrado en template: '{item_label}'")
        return False

    # Cuidado personal (Tabla 2, F22-F31)
    cp = areas.get("cuidado_personal", {})
    for item, val in cp.items():
        if isinstance(val, dict):
            _marcar_nivel_dificultad(doc, item, val.get("nivel", ""), val.get("observacion", ""))
        else:
            _marcar_nivel_dificultad(doc, item, str(val))

    # Comunicación (Tabla 2 F32-F38, Tabla 3 F00-F04)
    com = areas.get("comunicacion", {})
    for item, val in com.items():
        if isinstance(val, dict):
            _marcar_nivel_dificultad(doc, item, val.get("nivel", ""), val.get("observacion", ""))
        else:
            _marcar_nivel_dificultad(doc, item, str(val))

    # Movilidad (Tabla 3, F06-F15)
    mov = areas.get("movilidad", {})
    for item, val in mov.items():
        if isinstance(val, dict):
            _marcar_nivel_dificultad(doc, item, val.get("nivel", ""), val.get("observacion", ""))
        elif isinstance(val, tuple):
            nivel, obs = val
            _marcar_nivel_dificultad(doc, item, nivel, obs)
        else:
            _marcar_nivel_dificultad(doc, item, str(val))

    # Aprendizaje (Tabla 3, F17-F25)
    apr = areas.get("aprendizaje", {})
    for item, val in apr.items():
        if isinstance(val, dict):
            _marcar_nivel_dificultad(doc, item, val.get("nivel", ""), val.get("observacion", ""))
        else:
            _marcar_nivel_dificultad(doc, item, str(val))

    # Vida doméstica (Tabla 3, F27-F36)
    vd = areas.get("vida_domestica", {})
    for item, val in vd.items():
        if isinstance(val, dict):
            _marcar_nivel_dificultad(doc, item, val.get("nivel", ""), val.get("observacion", ""))
        else:
            _marcar_nivel_dificultad(doc, item, str(val))

    # ─── SECCIÓN 9 — Concepto ocupacional (Tabla 3 F39 + Tabla 4 F00) ───
    concepto = datos.get("concepto_ocupacional") or c.get("concepto") or c.get("concepto_desempeno") or datos.get("concepto_desempeno") or ""
    if not concepto or len(concepto) < 300:
        nombre = p.get("nombre", "el paciente")
        edad = p.get("edad", "")
        edad_str = f" de {edad} años" if edad else ""
        tipo_ev = s.get("tipo_evento") or "evento"
        fecha_ev = s.get("fecha_evento", "")
        fecha_str = f" ocurrido el {fecha_ev}" if fecha_ev else ""
        diag = s.get("diagnosticos") or s.get("diagnostico_cie10") or "diagnóstico en estudio"
        cargo = l.get("cargo") or e.get("cargo") or "su cargo"
        tareas = l.get("tareas", "")
        horario = l.get("horario", "")
        fam = datos.get("composicion_familiar", {})
        nucleo = fam.get("nucleo", "")
        ingreso = fam.get("ingreso_promedio", "")
        responsabilidad = fam.get("responsabilidad", "")
        rol_sens = (datos.get("rol_laboral") or {}).get("sensorio_motor", "")
        rol_psic = (datos.get("rol_laboral") or {}).get("psicologicos", "")
        rol_soc = (datos.get("rol_laboral") or {}).get("social", "")
        if genero == "F":
            afix = "Afiliada"
        else:
            afix = "Afiliado"
        concepto = f"{afix} {nombre}{edad_str}, quien ingresa al consultorio caminando por sus propios medios y quien a la valoración del desempeño ocupacional se encontró sin dificultad en los procesos cognitivos de pensamiento lógico y coherente, orientado en tiempo, lugar y espacio, quien ingresa por {tipo_ev}{fecha_str} con diagnóstico: {diag}."
        if nucleo:
            concepto += f" {afix} vive con {nucleo}"
            if ingreso or responsabilidad:
                extras_fam = []
                if ingreso: extras_fam.append(f"sostiene el hogar con un ingreso promedio de {ingreso}")
                if responsabilidad: extras_fam.append(f"con responsabilidades económicas de {responsabilidad}")
                concepto += ", " + ", ".join(extras_fam) + "."
            else:
                concepto += "."
        if cargo:
            concepto += f" Labora en el cargo de {cargo}"
            if tareas:
                concepto += f" con funciones de: {tareas}"
            concepto += "."
        if horario:
            concepto += f" Horario de trabajo: {horario}."
        componentes_list = []
        if rol_sens: componentes_list.append(f"Sensorio - motor: {rol_sens}")
        if rol_psic: componentes_list.append(f"Psicológicos: {rol_psic}")
        if rol_soc: componentes_list.append(f"Social: {rol_soc}")
        if componentes_list:
            concepto += "\n\nComponentes de desempeño:\n" + "\n".join(componentes_list)
        concepto += f"\n\n{afix} refiere que se encuentra laborando."
    # (concepto se escribe al final, despues de toda la limpieza)

    # ─── SECCIÓN 10 — Orientación ocupacional (Tabla 4, F02) ───
    orientacion = datos.get("orientacion_ocupacional") or datos.get("orientacion") or ""
    if orientacion:
        if genero == "F":
            _reemplazar_celda_entera(doc, "El afiliado puede realizar",
                orientacion.replace("El afiliado", "La afiliada"), bold=False, table_indices=[4])
        else:
            _reemplazar_celda_entera(doc, "El afiliado puede realizar", orientacion, bold=False, table_indices=[4])

    # ─── SECCIÓN 11 — Registro / Firmas (Tabla 4, F04-F07) ───
    prov = datos.get("proveedor", {})
    nombre_prov = prov.get("nombre", "SANDRA PATRICIA POLANIA OSORIO")
    ips_prov = prov.get("ips", "REHABILITACION INTEGRAL LABORAL Y OCUPACIONAL SAS")
    # Only replace in Tabla 4 (table_indices=[4]) to avoid corrupting concepto in Tabla 3
    _reemplazar_en_tablas(doc, "SANDRA PATRICIA POLANIA OSORIO", nombre_prov, table_indices=[4])
    _reemplazar_en_tablas(doc, "REHABILITACION INTEGRAL LABORAL Y OCUPACIONAL SAS", ips_prov, table_indices=[4])

    # Firma
    _insertar_firma(doc)

    # Limpieza agresiva de datos residuales de Juan Carlos Duran en el template
    _valores_a_limpiar = [
        "JUAN CARLOS DURAN NARVAEZ",
        "BOLIVARIANA DE MINERALES Y CIA LTDA",
        "Obrero de tratamiento roca",
        "Soldado Profesional",
        "EJERCITO NACIONAL",
        "Ingrid Johan Vargas Reyes",
        "bolivarianamineralesgh@gmail.com",
        "3143009418",
        "KM 2 VIA NEIVA - PALERMO",
        "860517272",
        "1 anio 4 meses aprox.",
        "1 año 4 meses aprox.",
        "HERIDA DEL CUARTO DEDO",
        "S611 HERIDA DEL CUARTO DEDO",
        "S601 CONTUSION DEL CUARTO DEDO",
        "02/03/2026",
        "06/04/2026",
        "Sin datos",
    ]
    for valor in _valores_a_limpiar:
        _reemplazar_en_parrafos(doc, valor, "[VERIFICAR]")
        _reemplazar_en_tablas(doc, valor, "[VERIFICAR]")

    # Escribir concepto al final (despues de limpieza para evitar corruptelas)
    try:
        if len(doc.tables) > 3 and len(doc.tables[3].rows) > 39:
            _poner_texto(doc.tables[3].rows[39].cells[0], concepto, bold=False, font_size=7.5)
        # T4F0: solo limpiar si tiene contenido residual de Juan Carlos (no dejar fila en blanco)
        if len(doc.tables) > 4 and len(doc.tables[4].rows) > 0:
            t4f0_text = doc.tables[4].rows[0].cells[0].text.strip()
            if t4f0_text and ("afiliado" in t4f0_text.lower() or "núcleo" in t4f0_text.lower() or "nucleo" in t4f0_text.lower()):
                _poner_texto(doc.tables[4].rows[0].cells[0], "", bold=False, font_size=7.5)
    except (IndexError, AttributeError) as e:
        _log(f"[WARN] Error escribiendo concepto: {e}")
        _reemplazar_celda_entera(doc, "Afiliado de", concepto, bold=False, font_size=7.5, table_indices=[3, 4])
        _reemplazar_celda_entera(doc, "Núcleo familiar:", "", bold=False, table_indices=[3, 4])

    doc.save(str(dst))
    _convertir_a_pdf(str(dst))
    return str(dst)
def _extraer_secciones(datos, extra=None):
    """Extrae secciones comunes del diccionario de datos.
    Retorna (p, e, l, s, c, *extras) donde extras son los valores
    de las claves adicionales solicitadas (ej: 'visita', 'rehabilitacion').
    """
    p = datos.get("paciente", {}) or {}
    e = datos.get("empresa", {}) or {}
    l = dict(datos.get("laboral", {}) or {})  # copia para no mutar el original
    s = datos.get("siniestro", {}) or {}
    c = datos.get("consulta", {}) or {}
    # Fallback: si laboral.cargo está vacío y empresa.cargo tiene valor, copiarlo
    if not l.get("cargo") and e.get("cargo"):
        l["cargo"] = e["cargo"]
    # Fallback: siniestro puede tener "tipo_evento" o "tipo"
    if not s.get("tipo") and s.get("tipo_evento"):
        s["tipo"] = s["tipo_evento"]
    if not s.get("tipo") and s.get("tipo_siniestro"):
        s["tipo"] = s["tipo_siniestro"]
    # Fallback: si paciente.afp está vacío pero hay dato en datos raíz
    if not p.get("afp") and datos.get("afp_notas"):
        p = dict(p)
        p["afp"] = datos["afp_notas"]
    # Fallback: si empresa.nit está vacío pero hay dato en datos raíz
    if not e.get("nit") and datos.get("nit_empresa_notas"):
        e = dict(e)
        e["nit"] = datos["nit_empresa_notas"]
    result = [p, e, l, s, c]

    if extra:
        for ex in extra:
            result.append(datos.get(ex, {}))
    return result


def generar_documento(tipo, datos, output_name=None):
    """Genera un documento clínico por tipo.

    Args:
        tipo: 'analisis', 'medidas', 'recomendaciones', 'cierre',
              'citacion', 'prueba', 'valoracion'
        datos: dict con los datos del paciente (formato JSON clínico)
        output_name: nombre del archivo de salida (sin extensión)

    Returns:
        str: ruta al archivo .docx generado
    """
    generadores = {
        "analisis": generar_analisis_exigencia,
        "medidas": generar_carta_medidas,
        "recomendaciones": generar_carta_recomendaciones,
        "cierre": generar_cierre_caso,
        "citacion": generar_citacion_empresas,
        "prueba": generar_prueba_trabajo,
        "valoracion": generar_valoracion_desempeno,
    }
    # Accept long-form names from json_validator
    ALIAS = {
        "analisis_exigencias": "analisis",
        "carta_medidas": "medidas",
        "carta_recomendaciones": "recomendaciones",
        "cierre_caso": "cierre",
        "citacion_empresas": "citacion",
        "prueba_trabajo": "prueba",
        "valoracion_desempeno": "valoracion",
    }
    tipo = ALIAS.get(tipo, tipo)

    if tipo not in generadores:
        raise ValueError(
            f"Tipo no soportado: {tipo}. Disponibles: {list(generadores.keys())}"
        )

    return generadores[tipo](datos, output_name)


# ============================================================
# TEST
# ============================================================
if __name__ == "__main__":
    datos_prueba = {
        "paciente": {
            "nombre": "MARIA PRUEBA",
            "nombre_original": "JUAN ORIGINAL",
            "documento": "12345678",
            "edad": 42,
            "dominancia": "Derecha",
            "estado_civil": "Soltera",
            "nivel_educativo": "Universitario",
            "telefono": "3001112233",
            "direccion": "Calle 10 #5-30, Neiva",
            "eps_ips": "Nueva EPS",
            "afp": "Porvenir",
        },
        "empresa": {
            "nombre": "EMPRESA DE PRUEBA SAS",
            "nombre_original": "EMPRESA ORIGINAL SAS",
            "nit": "900123456-1",
            "contacto": "Carlos Lopez",
            "contacto_original": "Juan Perez Original",
            "cargo_contacto": "Jefe de Talento Humano",
            "correo": "carlos@empresa.com",
            "telefono": "3112223344",
            "direccion": "Av. Siempre Viva 742, Neiva",
        },
        "laboral": {"cargo": "Operario", "area": "Producción"},
        "siniestro": {
            "id_siniestro": "987654321",
            "tipo": "Accidente de trabajo",
            "fecha_evento": "15/03/2026",
            "diagnosticos": "M545 - Lumbago no especificado",
            "tiempo_incapacidad": 15,
            "incapacitado": True,
            "ultimo_dia_incapacidad": "01/05/2026",
            "segmento_lesionado": "Columna lumbar",
        },
        "consulta": {
            "fecha": "2026-05-09",
            "motivo_consulta": "Dolor lumbar persistente",
            "enfermedad_actual": "Paciente refiere dolor en zona lumbar...",
        },
        "visita": {
            "tipo_estudio": "Prueba de trabajo",
            "tiempo_ejecucion": "3 horas",
            "objetivo": "Evaluar capacidad de desempeño",
            "fecha_hora": "15 mayo 2026, 9:00 AM",
            "fecha": "2026-05-15",
            "nombre_auditor": "Dr. Juan Perez",
        },
        "rehabilitacion": {
            "forma_integracion": "Reintegro con modificaciones",
            "tiempo_vigencia": "3 meses",
            "obstaculos": "Ninguno frente al caso.",
            "observaciones": "Ninguna frente al caso",
            "actividades_funcional": "Medicina laboral, Fisiatría",
            "actividades_ocupacional": "Valoración ocupacional, Seguimiento",
        },
        "recomendaciones": {
            "trabajador": ["Realizar pausas activas", "Evitar sobreesfuerzos"],
            "empresa": ["Adecuar puesto de trabajo"],
        },
    }

    import sys

    tipo = sys.argv[1] if len(sys.argv) > 1 else "citacion"
    print(f"🧪 Probando generador: {tipo}")
    try:
        r = generar_documento(tipo, datos_prueba)
        print(f"✅ DOCX generado: {r}")
        pdf_path = str(PDF / (Path(r).stem + ".pdf"))
        if Path(pdf_path).exists():
            print(f"✅ PDF generado: {pdf_path}")
        else:
            print(f"⚠️  PDF no encontrado (posible error en conversión)")
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
