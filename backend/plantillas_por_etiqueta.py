"""
Sistema de plantillas por etiquetas.
En vez de poner {{MARCADORES}} en el texto, identificamos cada celda
por la etiqueta de su fila/columna y reemplazamos el valor.
"""
import shutil, os, json
from docx import Document

BASE = "/root/fisioterapia/templates/formatos"
DOCS_DIR = "/root/fisioterapia/storage/docs"
PDF_DIR = "/root/fisioterapia/storage/pdfs"

# ====================================================================
# MAPEO DE ETIQUETAS → marcador (por tipo de documento)
# Cada entrada: (texto_etiqueta, columna_valor, marcador)
# columna_valor: 0 = misma celda, 1 = celda derecha, etc.
# ====================================================================

# FORMATO 1 — Análisis de Exigencias
ANALISIS_EXIGENCIA_MAP = [
    # Sección 1 — Identificación (tabla 0 y 1)
    ("NOMBRE DEL TRABAJADOR", 1, "{{NOMBRE_TRABAJADOR}}"),
    ("No. CÉDULA", 1, "{{NUMERO_DOCUMENTO}}"),
    ("ID del SINIESTRO", 1, "{{ID_SINIESTRO}}"),
    ("FECHA DE NACIMIENTO", 1, "{{FECHA_NACIMIENTO}}"),
    ("EDAD", 1, "{{EDAD}}"),
    ("DOMINANCIA", 1, "{{DOMINANCIA}}"),
    ("ESTADO CIVIL", 1, "{{ESTADO_CIVIL}}"),
    ("NIVEL EDUCATIVO", 1, "{{NIVEL_EDUCATIVO}}"),
    ("TELÉFONO", 1, "{{TELEFONO_TRABAJADOR}}"),
    ("DIRECCIÓN", 1, "{{DIRECCION_RESIDENCIA}}"),
    ("DIAGNÓSTICO", 1, "{{DIAGNOSTICOS}}"),
    ("FECHA DEL EVENTO", 1, "{{FECHA_EVENTO}}"),
    ("EPS", 1, "{{EPS_IPS}}"),
    ("AFP", 1, "{{AFP}}"),
    ("TIEMPO TOTAL DE INCAPACIDAD", 1, "{{TIEMPO_INCAPACIDAD}}"),
    ("NOMBRE DE LA EMPRESA", 1, "{{EMPRESA}}"),
    ("NIT", 1, "{{NIT_EMPRESA}}"),
    ("CARGO ACTUAL", 1, "{{CARGO_ACTUAL}}"),
    ("ÁREA", 1, "{{AREA}}"),
    ("FECHA DE INGRESO AL CARGO", 1, "{{FECHA_INGRESO_CARGO}}"),
    ("ANTIGÜEDAD EN EL CARGO", 1, "{{ANTIGUEDAD_CARGO}}"),
    ("FECHA DE INGRESO A LA EMPRESA", 1, "{{FECHA_INGRESO_EMPRESA}}"),
    ("ANTIGÜEDAD EN LA EMPRESA", 1, "{{ANTIGUEDAD_EMPRESA}}"),
    ("TIPO DE VINCULACION LABORAL", 1, "{{VINCULACION_LABORAL}}"),
    ("Contacto en empresa", 1, "{{CONTACTO_EMPRESA}}"),
    ("Correo(s) electrónico(s)", 1, "{{CORREO_EMPRESA}}"),
    ("Teléfonos de contacto empresa", 1, "{{TELEFONO_EMPRESA}}"),
    ("Dirección de empresa", 1, "{{DIRECCION_EMPRESA}}"),
]

# FORMATO 5 — Citación de Empresas (el más simple)
CITACION_EMPRESAS_MAP = [
    ("CARGO", 1, "{{CARGO_AFILIADO}}"),
    ("TIPO DE ESTUDIO A REALIZAR", 1, "{{TIPO_ESTUDIO}}"),
    ("TIEMPO DE EJECUCIÓN DEL ESTUDIO", 1, "{{TIEMPO_EJECUCION}}"),
    ("OBJETIVO DEL ESTUDIO", 1, "{{OBJETIVO_ESTUDIO}}"),
    ("FECHA Y HORA DE LA VISITA", 1, "{{FECHA_HORA_VISITA}}"),
    ("NOMBRE DEL AUDITOR DE REHABILITACIÓN", 1, "{{NOMBRE_AUDITOR}}"),
    ("DESCRIPCIÓN DEL ESTADO ACTUAL DEL CASO A", 1, "{{DESCRIPCION_ESTADO_CASO}}"),
]


def reemplazar_por_etiqueta(doc_path, mapping, output_path=None):
    """
    Busca etiquetas en las celdas y reemplaza el valor adyacente.
    Retorna cantidad de reemplazos hechos.
    """
    doc = Document(doc_path)
    cambios = 0
    
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = fila.cells
            for i, celda in enumerate(celdas):
                texto_celda = celda.text.strip()
                
                for etiqueta, col_offset, marcador in mapping:
                    if etiqueta.lower() in texto_celda.lower():
                        # Buscar la celda del valor
                        idx_valor = i + col_offset
                        if 0 <= idx_valor < len(celdas):
                            celda_valor = celdas[idx_valor]
                            valor_actual = celda_valor.text.strip()
                            
                            # Solo reemplazar si el valor no está ya marcado
                            if valor_actual and not valor_actual.startswith("{{"):
                                # Limpiar la celda valor y poner marcador
                                for p in celda_valor.paragraphs:
                                    for run in p.runs:
                                        run.text = marcador
                                cambios += 1
                        break  # Ya encontramos la etiqueta en esta celda
    
    if output_path:
        doc.save(output_path)
    return cambios


# ====================================================================
# PROCESAR DOCUMENTOS
# ====================================================================

MAPPINGS = {
    "ejemplo analisis de exigencia.docx": ANALISIS_EXIGENCIA_MAP,
    "ejemplo formato de citacion de empresas.docx": CITACION_EMPRESAS_MAP,
}

print("🏷️  Sistema de plantillas por etiquetas\n")

for archivo, mapping in MAPPINGS.items():
    path = os.path.join(BASE, archivo)
    bak = path + ".bak2"
    
    if not os.path.exists(path):
        print(f"  ❌ No encontrado: {archivo}")
        continue
    
    # Restaurar desde backup original si existe
    if os.path.exists(bak):
        shutil.copy(bak, path)
    
    # Crear backup
    if not os.path.exists(bak):
        shutil.copy(path, bak)
    
    cambios = reemplazar_por_etiqueta(path, mapping, path)
    print(f"📄 {archivo}: {cambios} campos reemplazados por etiqueta")

print(f"\n✅ Listo. Backups en .bak2")
